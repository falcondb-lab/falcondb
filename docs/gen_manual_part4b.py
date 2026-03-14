# -*- coding: utf-8 -*-
"""Part 4b - AI优化器/性能调优/故障排除/检查清单"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn; from docx.oxml import OxmlElement

OUTPUT = r"c:\falcondb\falcondb-main\docs\FalconDB_操作手册.docx"
doc = Document(OUTPUT)

def h1(t): return doc.add_heading(t, 1)
def h2(t): return doc.add_heading(t, 2)
def h3(t): return doc.add_heading(t, 3)
def np(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(11)
    return p
def bp(t):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t); r.font.size = Pt(11); return p
def cp(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(t); r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0F0F0')
    pPr.append(shd); return p
def bl(): doc.add_paragraph()
def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        if c.paragraphs[0].runs:
            run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = c._tc; tcPr = tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        r = t.rows[ri+1]
        for ci,v in enumerate(row):
            r.cells[ci].text = str(v)
            if r.cells[ci].paragraphs[0].runs:
                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return t
def pb(): doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第十六章  AI 优化器与 AIOps
# ═══════════════════════════════════════════════════
h1("第十六章  AI 查询优化器与 AIOps 引擎")

h2("16.1  AI 查询优化器设计理念")
np("传统数据库查询优化器依赖基于成本的规则（CBO），其代价模型参数通常是人工调优的固定值，无法适应不同工作负载下的运行时特征。FalconDB 内置的 AI 优化器（AiOptimizer）采用在线学习方法，将实际执行时间作为反馈信号，持续更新内部代价预测模型。")
np("与外部 AI 优化器（如独立 ML 服务）不同，FalconDB 的 AI 优化器完全在数据库进程内运行，无需额外网络通信和模型服务，模型参数仅占 184 字节（23 个 f64 权重）。")
bl()

h2("16.2  AI 优化器工作流程")
for line in [
    "SQL 文本 → 解析 → AST → 绑定分析 → 逻辑计划",
    "                              ↓",
    "                  extract_features()  ← 提取 15 维特征向量",
    "                              ↓",
    "         规则优化器生成候选计划集 [IndexScan, SeqScan, HashJoin, ...]",
    "                              ↓",
    "         AiOptimizer::select_plan()",
    "         ├─ samples_trained >= 50？",
    "         │   YES → 对每个候选计划打分，选择最低预测代价的",
    "         │   NO  → 使用规则优化器默认计划（零风险预热期）",
    "                              ↓",
    "         执行器执行选中的物理计划",
    "                              ↓",
    "         record_feedback(actual_execution_us)",
    "         在线 SGD 更新权重：w ← w - lr×∇loss + λ×w（L2 正则）",
]: cp(line)
bl()

h2("16.3  15 维特征向量详解")
tbl(["特征索引","特征名称","计算方法","含义"],
[["0","log2_estimated_rows","log₂(预期输出行数+1)","输出规模（对数，防大值主导）"],
 ["1","join_count","JOIN 节点数量","JOIN 越多计划越复杂"],
 ["2","filter_count","WHERE 谓词数量","过滤越多索引价值越高"],
 ["3","has_group_by","0.0 或 1.0","是否有 GROUP BY（影响聚合算子选择）"],
 ["4","has_order_by","0.0 或 1.0","是否有 ORDER BY（影响是否需要排序节点）"],
 ["5","has_limit","0.0 或 1.0","是否有 LIMIT（可提前终止扫描）"],
 ["6","index_available","0.0 或 1.0","主扫描表是否存在可用索引"],
 ["7","log2_max_table_rows","log₂(最大表行数+1)","最大表规模"],
 ["8","log2_second_table_rows","log₂(第二大表行数+1)","估计 JOIN 笛卡尔积规模"],
 ["9","selectivity","min(过滤后行数/总行数,1.0)","选择率（越低越受益于索引）"],
 ["10","projection_columns","SELECT 投影列数量","投影越多扫描宽度越大"],
 ["11","has_aggregate","0.0 或 1.0","是否有聚合函数"],
 ["12","has_distinct","0.0 或 1.0","是否有 DISTINCT"],
 ["13","subquery_depth","子查询嵌套最大深度","嵌套越深计划越复杂"],
 ["14","log2_total_bytes_estimate","log₂(估计处理总字节数+1)","内存压力估计"]])
bl()

h2("16.4  在线 SGD 训练机制")
np("AI 优化器使用随机梯度下降进行在线学习，每次查询完成后立即更新模型：")
bp("预测目标：log₂(实际执行时间_μs)，对数变换使模型对长尾查询更鲁棒")
bp("损失函数：均方误差 L = (predicted_log₂_us - actual_log₂_us)²")
bp("L2 正则化：λ = 0.001，防止过拟合，让不重要特征权重收缩到 0")
bp("线程安全：权重通过 RwLock 保护，预测时读锁，更新时写锁")
bp("预热阈值：50 个训练样本，未达到前始终使用规则优化器（零风险）")
for line in [
    "-- 验证 AI 优化器学习进度",
    "SHOW AI STATS;",
    "-- enabled:            true",
    "-- samples_trained:    15234   （已训练样本数）",
    "-- model_ready:        true    （>= 50 样本时为 true）",
    "-- ema_mae_log2:       0.42    （越低越准，≤ 0.5 为良好）",
    "-- query_fingerprints: 87      （已追踪 87 种不同查询形状）",
]: cp(line)
bl()

h2("16.5  AIOps 四大子系统")
tbl(["组件","职责","核心算法","查询接口"],
[["SlowQueryDetector","识别超过自适应阈值的慢查询","EMA p95 基线 × 2.0","SHOW AIOPS SLOW QUERIES"],
 ["AnomalyDetector","检测 TPS/延迟的统计异常","Welford 在线方差 3σ 规则","SHOW AIOPS ALERTS"],
 ["IndexAdvisor","分析全扫描模式生成索引建议","SeqScan 频率 + 过滤列分析","SHOW AIOPS INDEX ADVICE"],
 ["WorkloadProfiler","按 SQL 指纹聚合执行统计","蓄水池抽样 p95/p99","SHOW AIOPS WORKLOAD"]])
bl()

h3("16.5.1  慢查询检测器（自适应阈值）")
np("与固定阈值不同，FalconDB 的慢查询阈值随负载自动调整，负载高时阈值自动放宽，避免误报：")
for line in [
    "# 阈值计算逻辑",
    "EMA_ALPHA = 0.1    # 指数衰减系数",
    "",
    "# 每次查询完成后更新 EMA p95",
    "ema_p95 = ema_p95 * (1 - EMA_ALPHA) + p95_sample * EMA_ALPHA",
    "",
    "# 自适应阈值 = max(绝对下限 100ms, EMA_p95 × 2.0)",
    "threshold_us = max(100_000, ema_p95 * 2.0)",
    "",
    "# 效果：",
    "# 低负载（p95=1ms）：threshold = max(100ms, 2ms) = 100ms（使用绝对下限）",
    "# 高负载（p95=200ms）：threshold = max(100ms, 400ms) = 400ms（自动放宽）",
    "# 负载恢复：EMA 衰减，阈值逐渐收紧回正常水平",
]: cp(line)
bl()

h3("16.5.2  异常检测器（3σ 规则）")
np("使用 Welford 在线方差算法实时计算均值和标准差，当观测值超过 mean ± 3σ 时触发告警：")
tbl(["信号","WARNING 条件","CRITICAL 条件","冷启动保护"],
[["TPS（事务/秒）","TPS > mean+3σ（异常高）或 < mean-3σ（异常低）","N/A","< 30 个 1 秒样本时不触发"],
 ["查询延迟（μs）","延迟 > mean + 3σ","延迟 > mean × 5（超均值 5 倍）","< 30 个样本时不触发"],
 ["连续慢查询","连续 ≥ 5 次慢查询","连续 ≥ 20 次慢查询","无冷启动（基于计数）"]])
bl()

h3("16.5.3  索引推荐器工作原理")
for line in [
    "# 触发条件：某表 SeqScan 累计 >= 10 次 AND 平均耗时 >= 50ms",
    "",
    "# 分析过程：",
    "1. 提取该表所有 SeqScan 执行中 WHERE 子句的过滤列",
    "2. 统计每列出现频率，找出 Top-3 高频过滤列",
    "3. 生成建议：CREATE INDEX ON table_name(col1[, col2, ...])",
    "4. 计算影响分：impact_score = total_scans × avg_duration_ms",
    "5. 按影响分降序输出",
    "",
    "-- 查看建议",
    "SHOW AIOPS INDEX ADVICE;",
    "",
    "-- 示例输出：",
    "-- table  | scans | avg_us | impact | suggestion",
    "-- orders | 452   | 12500  | 5650   | CREATE INDEX ON orders(user_id)",
    "-- users  | 211   | 8300   | 1751   | CREATE INDEX ON users(email, status)",
]: cp(line)
bl()

h3("16.5.4  工作负载分析器")
np("WorkloadProfiler 按 SQL 指纹聚合执行统计，SQL 规范化后去掉字面值只保留结构（前 120 字符）：")
bp("每个指纹追踪：调用次数、均值/最大延迟、p95/p99（100 样本蓄水池抽样）、错误次数")
bp("最多追踪 1000 个不同查询指纹，超出时 LRU 淘汰最久未使用的")
bp("p95/p99 使用蓄水池抽样算法，内存开销固定（100 个样本/指纹）")
for line in [
    "-- Top-20 高耗时查询（含 p95/p99）",
    "SHOW AIOPS WORKLOAD;",
    "",
    "-- 示例输出：",
    "-- fingerprint                              | calls | avg_us | p95_us | p99_us | errors",
    "-- SELECT * FROM orders WHERE user_id = ?   | 45231 | 125    | 380    | 890    | 0",
    "-- UPDATE accounts SET balance = ? WHERE... | 8943  | 234    | 1200   | 4500   | 12",
]: cp(line)
pb()

# ═══════════════════════════════════════════════════
# 第十七章  性能调优
# ═══════════════════════════════════════════════════
h1("第十七章  性能调优指南")

h2("17.1  性能调优方法论")
np("高效的性能调优遵循「测量-分析-优化-验证」循环。在任何调优操作之前必须先建立基线，才能量化改进效果。调优优先级：")
bp("第一优先：应用层优化（SQL 写法、索引、事务粒度、连接池）——成本最低，效果最大")
bp("第二优先：系统配置调整（内存、WAL、GC 参数）——需要重启或热更新")
bp("第三优先：硬件升级（CPU、内存、SSD）——成本最高，应在前两步穷尽后再考虑")
bl()

h2("17.2  Rowstore 内存引擎调优")
h3("17.2.1  内存预算配置")
for line in [
    "[storage]",
    "memory_limit_bytes = 8589934592     # 8 GB 内存预算（0=不限制）",
    "",
    "# 推荐：为 OS 和其他进程预留 20% 物理内存",
    "# 4 GB 物理内存：memory_limit_bytes = 3221225472  (3 GB)",
    "# 16 GB 物理内存：memory_limit_bytes = 12884901888 (12 GB)",
    "# 64 GB 物理内存：memory_limit_bytes = 51539607552 (48 GB)",
]: cp(line)
bl()

h3("17.2.2  MVCC GC 调优")
tbl(["场景","interval_ms","batch_size","min_chain_length","说明"],
[["高写入（>10K 写/秒）","1000","5000","2","频繁 GC 防止内存积累"],
 ["标准生产","5000","1000","2","默认值，适合大多数场景"],
 ["读密集（<100 写/秒）","10000","500","5","减少 GC 对读性能的干扰"],
 ["批量导入期间","0（临时禁用）","—","—","导入完成后重新启用 GC"]])
bl()
np("长事务阻塞 GC 的诊断：")
for line in [
    "-- 1. 检查 GC 安全点是否停滞（long time = 有长事务阻塞）",
    "SHOW falcon.gc_safepoint;",
    "",
    "-- 2. 找出阻塞 GC 的长事务",
    "SHOW falcon.txn;   -- 查看 start_ts 很早的事务",
    "",
    "-- 3. 防止长事务积累",
    "SET statement_timeout = 30000;  -- 30 秒未完成自动取消",
    "SET idle_in_transaction_session_timeout = 60000;  -- 空闲事务 60s 超时",
]: cp(line)
bl()

h2("17.3  WAL 性能调优")
h3("17.3.1  组提交参数")
tbl(["参数","低延迟配置","高吞吐配置","说明"],
[["flush_interval_us","500 μs","5000 μs","刷新间隔越大批次越多，吞吐越高延迟略增"],
 ["sync_mode","fdatasync","fdatasync","比 fsync 快（不刷文件元数据）"],
 ["segment_size_bytes","64 MB","128 MB","大段减少段切换，但崩溃恢复时间略长"],
 ["group_commit","true","true","始终启用，几乎无副作用"]])
bl()

h3("17.3.2  磁盘 I/O 优化建议")
bp("WAL 目录使用独立 NVMe SSD（与数据文件分离），避免 I/O 竞争")
bp("Linux：对 WAL 分区添加 noatime 挂载选项，禁用访问时间更新")
bp("Windows：确保 WAL 目录磁盘的写缓存已启用（FalconDB 用 fsync 保证持久性）")
bp("内核参数（Linux）：vm.dirty_background_ratio=5，vm.dirty_ratio=10")
bl()

h2("17.4  SQL 查询优化最佳实践")
h3("17.4.1  高频 OLTP 优化")
for line in [
    "-- ✓ 主键点查（O(1) 哈希查找，最快路径）",
    "SELECT id, status FROM orders WHERE id = $1;",
    "",
    "-- ✓ 预处理语句（避免每次重复解析）",
    "PREPARE get_order (BIGINT) AS SELECT * FROM orders WHERE id = $1;",
    "EXECUTE get_order(12345);",
    "",
    "-- ✓ 单语句读-改-写（减少 RTT 次数）",
    "UPDATE accounts SET balance = balance - $1",
    "WHERE id = $2 AND balance >= $1",
    "RETURNING balance;   -- 扣款+验证+确认一次完成",
    "",
    "-- ✗ 多次 RTT 的读-改-写（避免）",
    "BEGIN;",
    "SELECT balance FROM accounts WHERE id = $1;   -- RTT 1",
    "-- 应用层验证...",
    "UPDATE accounts SET balance = balance - $2 WHERE id = $1;  -- RTT 2",
    "COMMIT;  -- RTT 3",
]: cp(line)
bl()

h3("17.4.2  批量操作优化")
for line in [
    "-- ✓ 单语句多值批量插入（比逐行快 5-10x）",
    "INSERT INTO events (ts, type, data) VALUES",
    "    (now(), 'click', $1), (now(), 'view', $2), ...;",
    "-- 建议每批 100-5000 行，超过 5000 行分批提交",
    "",
    "-- ✓ COPY 超大批量导入（比 INSERT 快 2-5x）",
    "COPY events FROM '/data/events.csv' WITH (FORMAT CSV);",
    "",
    "-- ✓ 事务内批量：共享一次提交开销",
    "BEGIN;",
    "INSERT INTO logs ...;   -- 1000 行",
    "UPDATE metrics ...;     -- 500 行",
    "COMMIT;   -- 只需一次 WAL fsync",
]: cp(line)
bl()

h3("17.4.3  JOIN 查询优化")
for line in [
    "-- ✓ 小表驱动大表（减少 JOIN 输入规模）",
    "SELECT u.username, o.total",
    "FROM orders o",
    "INNER JOIN users u ON o.user_id = u.id",
    "WHERE u.region = 'CN';   -- 过滤早发生，先缩小 users 再 JOIN",
    "",
    "-- ✓ 覆盖索引（避免回表查询）",
    "CREATE INDEX idx_orders_covering",
    "ON orders (user_id, status)   -- WHERE/JOIN 列",
    "INCLUDE (id, total);           -- SELECT 列（避免回表）",
    "",
    "-- ✓ EXISTS 替代 IN（大子查询时更高效）",
    "SELECT * FROM users u WHERE EXISTS (",
    "    SELECT 1 FROM orders o WHERE o.user_id = u.id AND o.status = 'pending'",
    ");",
]: cp(line)
bl()

h3("17.4.4  索引失效场景（避免）")
tbl(["反模式","原因","正确写法"],
[["WHERE YEAR(created_at) = 2026","对索引列使用函数","WHERE created_at >= '2026-01-01' AND created_at < '2027-01-01'"],
 ["WHERE name LIKE '%alice%'","前导通配符无法使用索引","WHERE name LIKE 'alice%' 或使用全文搜索"],
 ["WHERE id = '123'（id 为 INT）","隐式类型转换","WHERE id = 123（去掉引号）"],
 ["WHERE col1 = 1 OR col2 = 2","OR 跨多列难以使用复合索引","UNION ALL 分开查询"],
 ["对低基数列建单列索引","gender=M/F 两种值，全扫比索引快","与高基数列组成复合索引"]])
bl()

h2("17.5  连接池最佳实践")
tbl(["参数","推荐值","说明"],
[["最大连接数（pool_max_size）","CPU 核心数 × 2~4","避免过多线程上下文切换"],
 ["最小连接数（pool_min_idle）","CPU 核心数","保持热连接，避免高峰冷启动"],
 ["连接超时（connect_timeout）","5 秒","避免长时间阻塞"],
 ["空闲超时（idle_timeout）","30-60 秒","回收空闲连接"],
 ["最大生命周期（max_lifetime）","30 分钟","定期轮换，防止状态积累"],
 ["检活（keepalive_time）","30 秒，SELECT 1","检测并替换失效连接"]])
for line in [
    "# HikariCP（Java）配置示例",
    "config.setMaximumPoolSize(16)          # CPU×4（4核服务器）",
    "config.setMinimumIdle(4)               # CPU×1",
    "config.setConnectionTimeout(5000)      # 5s",
    "config.setIdleTimeout(60000)           # 60s",
    "config.setMaxLifetime(1800000)         # 30min",
    "config.setKeepaliveTime(30000)         # 30s",
]: cp(line)
bl()

h2("17.6  热点行竞争处理")
np("高并发下多事务同时修改同一行会导致大量 OCC 冲突（SQLSTATE 40001）。常见场景：库存扣减、余额更新、计数器。解决方案：")
for line in [
    "-- 方案一：写入队列化（应用层合并批量更新）",
    "-- 1000 个并发扣减合并为一条 SQL",
    "UPDATE inventory SET stock = stock - $batch_size WHERE product_id = 101;",
    "",
    "-- 方案二：乐观分桶（将热点行分散到 N 个分桶）",
    "UPDATE inventory_shards",
    "SET stock = stock - 1",
    "WHERE product_id = 101 AND shard_id = (user_id % 16)",
    "AND stock > 0;",
    "-- 再汇总：SELECT SUM(stock) FROM inventory_shards WHERE product_id = 101",
    "",
    "-- 方案三：异步处理（写消息队列，后台批处理）",
    "INSERT INTO deduction_queue (product_id, qty, user_id, ts) VALUES (101, 1, $uid, now());",
    "-- 后台批处理器定期合并队列更新库存",
]: cp(line)
pb()

# ═══════════════════════════════════════════════════
# 第十八章  常见问题与故障排除
# ═══════════════════════════════════════════════════
h1("第十八章  常见问题与故障排除")

h2("18.1  启动问题")
tbl(["症状","错误信息","根本原因","解决步骤"],
[["服务启动失败","address already in use (0.0.0.0:5433)","端口被占用",
  "1. netstat -ano | findstr 5433\n2. 终止占用进程或修改 pg_listen_addr\n3. 重启服务"],
 ["服务启动失败","config file not found","配置文件路径错误",
  "1. falcon config check\n2. 确认 --config 参数路径\n3. 用 --print-default-config 生成"],
 ["服务启动失败","permission denied: data dir","数据目录权限不足",
  "1. 检查 data_dir 路径\n2. icacls data_dir /grant 'FalconDBSvc:F'（Windows）"],
 ["服务启动失败","WAL recovery failed: corrupted segment","WAL 文件损坏",
  "1. 检查最后 WAL 段文件\n2. 从最近备份恢复\n3. 联系技术支持"],
 ["服务启动失败","config version mismatch","升级后配置格式变更",
  "1. falcon config check\n2. falcon config migrate"]])
bl()

h2("18.2  连接与认证问题")
tbl(["症状","诊断方法","可能原因","解决方案"],
[["连接被拒绝（Connection refused）","telnet host 5433","服务未启动/端口错误/防火墙",
  "falcon service status；开放防火墙端口；检查 pg_listen_addr"],
 ["FATAL: password authentication failed","SHOW falcon.security_audit","密码错误/认证方式不匹配",
  "确认密码；检查 auth.method；确认环境变量已设置"],
 ["IP 被锁定 300 秒","SHOW falcon.security_audit（auth_lockouts）","5 次认证失败触发速率限制",
  "等待 300 秒自动解锁；检查应用密码配置"],
 ["SSL 握手失败","openssl s_client -connect host:5433","证书过期/CA 不匹配",
  "检查证书有效期；确认 cert_path/key_path/ca_cert_path"],
 ["连接超时","ping/traceroute","网络问题/服务过载/连接数满",
  "检查网络；检查 max_connections；优化连接池大小"]])
bl()

h2("18.3  查询性能问题")
tbl(["症状","诊断命令","根本原因","解决方案"],
[["查询延迟突然升高（全局）","SHOW AIOPS ALERTS; SHOW falcon.txn_stats",
  "大量 OCC 冲突/内存压力/WAL 慢/长事务",
  "按告警指引处理；EXPLAIN ANALYZE 分析热点查询"],
 ["特定查询慢（单次）","EXPLAIN ANALYZE <sql>",
  "缺少索引/统计信息陈旧/执行计划退化",
  "按 AIOPS 建议建索引；ANALYZE TABLE 更新统计"],
 ["大量 OCC 冲突（40001）","SHOW falcon.txn_stats（occ_conflicts）",
  "热点行竞争（高频修改同一行）",
  "重新设计数据模型；使用批量更新；减少事务持有时间"],
 ["内存持续增长","SHOW falcon.gc_stats; SHOW falcon.txn",
  "长事务阻塞 GC/GC 间隔过大",
  "设置 statement_timeout；降低 gc.interval_ms"],
 ["全表扫描（SeqScan）","EXPLAIN ANALYZE; SHOW AIOPS INDEX ADVICE",
  "缺少索引/索引失效",
  "按 AIOPS 建议创建索引；检查索引失效场景"]])
bl()

h2("18.4  复制与集群问题")
tbl(["症状","诊断命令","原因","解决方案"],
[["副本复制延迟持续增大","SHOW falcon.replication_stats",
  "网络带宽不足/副本负载过高",
  "检查网络带宽；减少副本查询并发；优化 WAL 段大小"],
 ["副本断开后无法重连","检查服务器日志 falcon.log",
  "WAL 段被回收（副本落后太远）",
  "从最新备份重建副本；增加 WAL 保留时间"],
 ["脑裂告警（split-brain rejected）","检查日志 EPOCH FENCE",
  "网络分区后双主场景",
  "Epoch 隔离已自动防护；确认旧主节点日志显示写入被拒绝"],
 ["故障转移后数据不一致","SHOW falcon.replication_stats",
  "提交策略为 Local，副本有延迟",
  "改用 sync-replica 策略；定期验证副本一致性"]])
bl()

h2("18.5  WAL 和存储问题")
tbl(["症状","排查步骤","解决方案"],
[["WAL 占用磁盘过大","检查 data_dir 大小；SHOW falcon.checkpoint_stats",
  "执行 CHECKPOINT；确认 GC 正常；增大磁盘容量"],
 ["PITR 恢复失败","检查 archive_dir WAL 文件连续性",
  "确保归档目录完整无缺失；检查 retention_hours 是否覆盖目标时间"],
 ["启动时 WAL 恢复报错","检查日志 WAL replay；falcon doctor",
  "WAL 文件损坏则从备份恢复；磁盘不足则清理后重试"],
 ["数据目录权限拒绝","检查文件系统权限",
  "确保服务账户有 data_dir 读写权限（见 11.11 节）"]])
bl()

h2("18.6  常用诊断命令集合")
for line in [
    "# ── 服务状态 ──────────────────────────────────",
    "falcon service status",
    "falcon doctor",
    "",
    "# ── SQL 诊断（psql 连接后执行）─────────────",
    "SELECT version();                    -- 确认连接和版本",
    "SHOW falcon.status;                  -- 服务器状态",
    "SHOW falcon.connections;             -- 连接分布",
    "SHOW falcon.txn;                     -- 活跃事务（找长事务）",
    "SHOW falcon.txn_stats;               -- 事务统计（OCC 冲突率）",
    "SHOW falcon.gc_stats;                -- GC 状态",
    "SHOW falcon.gc_safepoint;            -- GC 安全点诊断",
    "SHOW falcon.replication_stats;       -- 复制状态",
    "SHOW falcon.table_stats;             -- 表大小统计",
    "SHOW AIOPS ALERTS;                   -- 近期异常告警",
    "SHOW AIOPS SLOW QUERIES;             -- 慢查询记录",
    "SHOW AIOPS INDEX ADVICE;             -- 索引建议",
    "SHOW AIOPS WORKLOAD;                 -- 工作负载热点",
    "SHOW AI STATS;                       -- AI 优化器状态",
    "SHOW falcon.security;                -- 安全态势",
    "SHOW falcon.audit_log;               -- 近期审计事件",
    "",
    "# ── 日志查看（Windows）──────────────────────",
    r"Get-Content C:\ProgramData\FalconDB\logs\falcon.log -Tail 200 -Wait",
    r'Select-String "ERROR|WARN|panic" C:\ProgramData\FalconDB\logs\falcon.log',
    "",
    "# ── 网络诊断 ─────────────────────────────────",
    "Test-NetConnection -ComputerName 127.0.0.1 -Port 5433",
]: cp(line)
pb()

# ═══════════════════════════════════════════════════
# 第十九章  生产运维检查清单
# ═══════════════════════════════════════════════════
h1("第十九章  生产运维检查清单")

h2("19.1  上线前检查清单")
tbl(["类别","检查项目","验证方法","状态"],
[["安全","认证方式为 scram-sha-256（非 trust）","检查 falcon.toml [server.auth]","□"],
 ["安全","TLS 启用，最低版本 TLS 1.2+","SHOW falcon.security","□"],
 ["安全","密码策略已启用（8+ 字符，复杂度）","SHOW falcon.security_audit","□"],
 ["安全","认证速率限制（5 次失败锁定）","SHOW falcon.security_audit","□"],
 ["安全","SQL 防火墙已启用","SHOW falcon.security_audit","□"],
 ["安全","审计日志已启用","SHOW falcon.audit_log","□"],
 ["安全","应用使用最小权限账户（非超级用户）","\\du 查看角色","□"],
 ["安全","密码通过环境变量传入（非硬编码）","检查配置和代码仓库","□"],
 ["持久化","WAL 已启用（wal_enabled = true）","检查 falcon.toml [storage]","□"],
 ["持久化","data_dir 在可靠存储（SSD）上","检查目录和磁盘类型","□"],
 ["持久化","持久化策略已配置（local-fsync/sync-replica）","检查 [wal] durability_policy","□"],
 ["持久化","DCG 保证已验证（写入后重启检验）","运行 DCG 演示脚本","□"],
 ["高可用","至少一个从节点已部署并同步","SHOW falcon.replication_stats","□"],
 ["高可用","故障转移已测试（演练过）","运行 e2e_failover 脚本","□"],
 ["高可用","客户端已配置重连和重试逻辑","检查应用代码","□"],
 ["备份","全量备份已创建并验证可恢复","运行备份验证流程","□"],
 ["备份","备份定时任务已设置","检查 Windows 任务计划/cron","□"],
 ["备份","PITR 已启用（如需精确恢复）","检查 falcon.toml [pitr]","□"],
 ["监控","Prometheus 指标已接入监控系统","访问 /metrics 端点","□"],
 ["监控","关键告警规则已配置（延迟/连接/内存）","检查 Prometheus/Grafana","□"],
 ["监控","日志已接入日志系统","检查日志格式和输出配置","□"],
 ["性能","关键查询已创建索引","SHOW AIOPS INDEX ADVICE","□"],
 ["性能","连接池已正确配置","检查应用连接池参数","□"],
 ["性能","内存限制已配置（Rowstore）","检查 [storage] memory_limit_bytes","□"],
 ["性能","GC 已启用并参数合理","检查 falcon.toml [gc]","□"],
 ["运维","服务已注册为系统服务","falcon service status","□"],
 ["运维","日志轮转已配置","检查 [logging] rotation/max_files","□"],
 ["运维","磁盘空间告警已设置（WAL 目录）","检查监控规则","□"],
 ["运维","文档已更新（连接信息/架构图）","检查 wiki/运维文档","□"],
 ["运维","on-call 值班人员已接受数据库培训","培训记录","□"]])
bl()

h2("19.2  日常运维检查（推荐每日执行）")
for line in [
    "-- 每日例行检查（约 10 分钟）",
    "",
    "-- 1. 检查服务健康状态",
    "SHOW falcon.status;",
    "SHOW falcon.connections;     -- 确认连接数正常",
    "",
    "-- 2. 检查是否有未处理的告警",
    "SHOW AIOPS ALERTS;",
    "",
    "-- 3. 检查慢查询趋势",
    "SHOW AIOPS SLOW QUERIES;",
    "",
    "-- 4. 检查复制状态（主节点执行）",
    "SHOW falcon.replication_stats;   -- lag_lsn 应接近 0",
    "",
    "-- 5. 检查 GC 健康",
    "SHOW falcon.gc_stats;",
    "SHOW falcon.gc_safepoint;        -- safepoint_ts 应持续增长",
    "",
    "-- 6. 检查事务冲突率",
    "SHOW falcon.txn_stats;           -- occ_conflicts 突增需关注",
    "",
    "-- 7. 查看近期审计事件（安全异常）",
    "SHOW falcon.audit_log;",
    "",
    "-- 8. 查看索引建议（优化机会）",
    "SHOW AIOPS INDEX ADVICE;",
]: cp(line)
bl()

h2("19.3  周期性维护任务")
tbl(["频率","任务","操作","说明"],
[["每日","验证备份可用性","运行备份验证脚本","确保备份文件未损坏且可恢复"],
 ["每日","检查磁盘空间（WAL/数据目录）","系统监控或 df -h","WAL 持续增长需扩容或调整保留策略"],
 ["每日","检查慢查询和索引建议","SHOW AIOPS INDEX ADVICE","及时处理新产生的慢查询"],
 ["每周","审查安全审计日志","SHOW falcon.audit_log","检查异常登录和权限变更"],
 ["每周","更新表统计信息","ANALYZE TABLE table_name","维护优化器统计准确性"],
 ["每月","验证故障转移流程","运行 e2e_failover 演练","确保高可用机制正常"],
 ["每月","审查和轮换密码","ALTER ROLE user PASSWORD '...'","遵循密码有效期策略"],
 ["每季度","检查 TLS 证书有效期","openssl x509 -in server.crt -noout -dates","证书到期前 30 天更新"],
 ["每季度","容量规划评估","分析历史趋势","预测未来 6 个月存储和内存需求"]])
bl()

h2("19.4  版本升级流程")
for line in [
    "# 步骤 1：检查配置版本兼容性",
    "falcon config check",
    "",
    "# 步骤 2：备份当前配置和数据",
    r'$ts = Get-Date -Format "yyyyMMdd_HHmmss"',
    r'Copy-Item "C:\ProgramData\FalconDB\conf\falcon.toml" "D:\backup\falcon_$ts.toml"',
    r'Copy-Item -Recurse "C:\ProgramData\FalconDB\data" "D:\backup\falcon_data_$ts"',
    "",
    "# 步骤 3：停止服务",
    "falcon service stop",
    "",
    "# 步骤 4：安装新版本（替换 falcon.exe）",
    "# 将新的 falcon.exe 复制到 C:\\Program Files\\FalconDB\\bin\\",
    "",
    "# 步骤 5：迁移配置（如有格式变更）",
    "falcon config migrate",
    "falcon config check",
    "",
    "# 步骤 6：启动服务并验证",
    "falcon service start",
    "falcon service status",
    "psql -h 127.0.0.1 -p 5433 -U falcon -c 'SELECT version();'",
    "SHOW falcon.status;",
    "",
    "# 步骤 7：验证关键功能",
    "# 运行应用层冒烟测试",
    "# 检查 SHOW AIOPS ALERTS 是否有异常",
]: cp(line)
bl()

h2("19.5  紧急故障响应手册")
tbl(["故障等级","触发条件","立即操作","后续处理"],
[["P0（服务中断）","主节点宕机，从节点未自动提升",
  "1. 手动检查故障转移日志\n2. 检查 Epoch 隔离是否生效\n3. 通知业务方切换连接端点",
  "事后根因分析；更新 HA 配置；补充监控告警"],
 ["P1（严重降级）","P99 延迟 > 1s 或 OCC 冲突 > 1000/min",
  "1. SHOW AIOPS ALERTS\n2. SHOW falcon.txn 找热点\n3. 临时降低应用并发",
  "索引优化；数据模型重设计；评估扩容"],
 ["P2（复制延迟）","副本延迟 > 10000 LSN",
  "1. SHOW falcon.replication_stats\n2. 检查副本负载\n3. 暂停副本查询流量",
  "优化副本资源；评估网络带宽；考虑硬件升级"],
 ["P3（磁盘告警）","data_dir 使用率 > 80%",
  "1. 检查 WAL 文件\n2. 执行 CHECKPOINT\n3. 清理旧备份",
  "扩容存储；调整 WAL 保留策略；容量规划"]])
pb()

# ═══════════════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════════════
h1("附  录")

h2("A.  SQLSTATE 错误代码参考")
tbl(["SQLSTATE","含义","常见触发场景","处理建议"],
[["40001","serialization_failure（序列化失败）","OCC 写-写冲突","应用层重试事务（指数退避，建议 3 次以上）"],
 ["23505","unique_violation（唯一键冲突）","INSERT 重复主键/唯一键","使用 ON CONFLICT DO NOTHING/UPDATE"],
 ["23502","not_null_violation（非空约束违反）","INSERT/UPDATE 给 NOT NULL 列设 NULL","补充默认值或应用层校验"],
 ["23514","check_violation（检查约束违反）","值不满足 CHECK 条件","检查业务逻辑和约束定义"],
 ["42P01","undefined_table（表不存在）","查询未建表或名称拼写错误","确认表名；检查所在 schema"],
 ["42703","undefined_column（列不存在）","查询不存在的列","检查列名拼写；表结构变更后更新 SQL"],
 ["0A000","feature_not_supported","使用了不支持的功能（触发器/物化视图等在早期版本）","查看版本支持情况；升级 FalconDB 版本"],
 ["08006","connection_failure（连接失败）","服务器宕机或网络中断","重试连接；检查服务器状态"],
 ["57014","query_canceled（查询取消）","超过 statement_timeout 或手动取消","优化查询性能；适当增大超时值"],
 ["53100","disk_full（磁盘满）","WAL 或数据目录磁盘空间不足","立即清理磁盘；扩容存储"],
 ["XX000","internal_error","数据库内部错误（Panic）","立即检查错误日志；联系技术支持"]])
bl()

h2("B.  FalconDB Crate 架构总览")
tbl(["Crate 名称","职责说明"],
[["falcon_storage","多引擎存储（Rowstore/LSM/RocksDB/redb）、MVCC VersionChain2、二级索引 ART-Tree、WAL、MVCC GC、TDE、CDC"],
 ["falcon_txn","事务生命周期管理、OCC 验证、时间戳分配、保存点、Savepoint 管理"],
 ["falcon_sql_frontend","SQL 解析（sqlparser-rs）、绑定分析、类型检查、名称解析"],
 ["falcon_planner","逻辑/物理计划生成、CBO 代价优化、AI 优化器（在线 SGD）、路由提示生成"],
 ["falcon_executor","算子执行、表达式求值、全文搜索引擎（tsvector/tsquery）、向量化列存聚合"],
 ["falcon_protocol_pg","PostgreSQL v3 线协议编解码、TCP 服务器、取消请求处理"],
 ["falcon_cluster","分片映射、WAL gRPC 复制、故障转移、scatter/gather、Epoch 双层隔离"],
 ["falcon_raft","Raft 共识实现（单节点默认/多节点 raft_member）"],
 ["falcon_observability","Prometheus 指标注册、结构化 tracing 日志、pg_stat_statements 视图"],
 ["falcon_enterprise","控制平面 HA、TLS 轮换、备份恢复编排、多租户、资源隔离、IP 白名单"],
 ["falcon_common","共享类型定义、错误码、配置结构、Schema 管理、RBAC 模型、Decimal 精度类型"],
 ["falcon_server","主二进制文件，组装所有 crate，服务入口，Windows 服务集成"],
 ["falcon_cli","交互式 REPL CLI，集群管理命令，数据导入导出工具"],
 ["falcon_bench","YCSB 风格基准测试工具，性能回归对比"]])
bl()

h2("C.  版本历史摘要")
tbl(["版本","主要新增功能"],
[["v0.1 – v0.9","OLTP 基础功能、WAL 持久化、主从复制（gRPC）、OCC 事务、SQL 解析器、认证安全、混沌测试加固"],
 ["v1.0","LSM 磁盘引擎、企业功能（RLS/TDE/PITR/CDC）、分布式加固、完整 SQL 覆盖、AIOps 基础"],
 ["v1.0.1 – v1.0.3","零 Panic 加固、故障转移×事务矩阵测试（47 项）、确定性与信任加固"],
 ["v1.1","USTM 分层内存（Hot/Warm/Cold）、融合流式聚合、接近 PG 查询对等性、pg_stat_statements"],
 ["v1.2（当前）","Standard 版（RocksDB 默认）、PL/pgSQL 触发器、物化视图、ENUM 类型、AI 优化器（在线 SGD + SHOW AI STATS）、分片自动再均衡（47 项混沌测试）、AIOps 四大子系统"]])
bl()

h2("D.  性能基准数据参考")
tbl(["测试场景","FalconDB","PostgreSQL 16","说明"],
[["INSERT 吞吐（百万行批量）","~340K 行/秒","~185K 行/秒","FalconDB 约 1.84x 更快"],
 ["INSERT 阶段耗时（100批×10K行）","~2.9 秒","~5.4 秒","FalconDB 约 1.86x 更快"],
 ["单分片点查延迟（P99）","< 1 ms","~2-5 ms","FalconDB 更快（绕过 2PC）"],
 ["跨分片查询（ORDER BY+COUNT）","~2.8 秒","~0.65 秒","PostgreSQL 更快（无分布式开销）"],
 ["总计（DDL+INSERT+查询）","~6.4 秒","~6.1 秒","大体相当"],
 ["写入+读取并发（pgbench -c 50）","待测","待测","参见 pocs/falcondb-poc-benchmarks/"]])
np("注：以上数据来自 FalconDB 内置基准测试（benchmarks/ 目录），硬件为标准云主机。实际性能取决于数据模型、查询复杂度和硬件配置。FalconDB 在单分片 OLTP 场景优势明显，跨分片分析查询场景 PostgreSQL 更有优势。")
bl()

h2("E.  不支持的功能清单（SQLSTATE 0A000）")
tbl(["不支持的功能","说明","替代方案"],
[["CREATE TRIGGER（v1.1 及以下）","v1.2+ 已支持 PL/pgSQL 触发器","升级到 v1.2+"],
 ["CREATE FUNCTION（v1.1 及以下）","v1.2+ 已支持 SQL/PL/pgSQL 函数","升级到 v1.2+"],
 ["CREATE MATERIALIZED VIEW（v1.1及以下）","v1.2+ 已支持","升级到 v1.2+"],
 ["CREATE EXTENSION","FalconDB 不支持 PostgreSQL 扩展机制","使用 FalconDB 内置功能"],
 ["FULL OUTER JOIN","不支持，返回 0A000","改写为 LEFT JOIN UNION ALL RIGHT JOIN"],
 ["LATERAL JOIN","不支持","改写为子查询或 CTE"],
 ["CREATE TABLESPACE","不支持表空间管理","通过 ENGINE= 选择不同引擎"],
 ["DOMAIN 类型","不支持","使用 CHECK 约束替代"],
 ["RANGE 类型","不支持","使用 BETWEEN 或两列替代"],
 ["COPY FROM STDIN（二进制格式）","只支持 CSV 格式","使用 CSV 格式 COPY"],
 ["VACUUM / ANALYZE（手动）","自动维护，不需要手动执行","FalconDB GC 自动处理"]])
bl()

doc.add_paragraph("─" * 80)
p = doc.add_paragraph()
r = p.add_run("文档结束  |  FalconDB v1.2 数据库操作手册  |  版权所有 © 2026 FalconDB Team")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
p2 = doc.add_paragraph()
r2 = p2.add_run("技术支持与源码：https://github.com/falcondb-lab/falcondb")
r2.font.size = Pt(10)
doc.add_paragraph("─" * 80)

print("Part 4b 完成，保存中...")
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
