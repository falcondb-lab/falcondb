# -*- coding: utf-8 -*-
"""FalconDB 操作手册生成器 Part 2 - DDL/DML/索引/安全/集群/故障转移"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\falcondb\falcondb-main\docs\FalconDB_操作手册.docx"
doc = Document(OUTPUT)

def h1(t): return doc.add_heading(t, 1)
def h2(t): return doc.add_heading(t, 2)
def h3(t): return doc.add_heading(t, 3)

def np(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(11)
    return p

def bp(t):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t); r.font.size = Pt(11)
    return p

def cp(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(t)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
    pPr.append(shd)
    return p

def bl(): doc.add_paragraph()

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        if c.paragraphs[0].runs:
            run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        r = t.rows[ri+1]
        for ci,v in enumerate(row):
            r.cells[ci].text = str(v)
            if r.cells[ci].paragraphs[0].runs:
                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return t

def pb(): doc.add_page_break()

# ══ 第八章 DDL ══
h1("第八章  数据定义语言（DDL）")
h2("8.1  数据库管理")
for line in [
    "-- 创建数据库",
    "CREATE DATABASE mydb;",
    "CREATE DATABASE mydb ENCODING 'UTF8';",
    "",
    "-- 删除数据库",
    "DROP DATABASE mydb;",
    "DROP DATABASE IF EXISTS mydb;",
    "",
    "-- 切换数据库（psql 客户端命令）",
    "\\c mydb",
]:
    cp(line)
bl()

h2("8.2  CREATE TABLE——完整语法")
np("FalconDB 支持完整的 PostgreSQL 建表语法，包括列约束、表约束、外键引用和存储引擎选项。")
for line in [
    "CREATE TABLE table_name (",
    "    col1  data_type  [NOT NULL] [DEFAULT expr] [PRIMARY KEY]",
    "                    [UNIQUE] [CHECK (condition)]",
    "                    [REFERENCES other_table(col)],",
    "    col2  data_type,",
    "    ...,",
    "    [PRIMARY KEY (col1, col2)],",
    "    [UNIQUE (col1, col2)],",
    "    [CHECK (condition)],",
    "    [FOREIGN KEY (col) REFERENCES other_table(col)]",
    ") [ENGINE=engine_name];",
    "",
    "-- 完整示例：订单表",
    "CREATE TABLE orders (",
    "    id           BIGSERIAL PRIMARY KEY,",
    "    user_id      BIGINT NOT NULL REFERENCES users(id),",
    "    order_no     VARCHAR(32) NOT NULL UNIQUE,",
    "    status       TEXT NOT NULL DEFAULT 'pending'",
    "                     CHECK (status IN ('pending','processing','shipped','done','cancelled')),",
    "    total        DECIMAL(12,2) NOT NULL CHECK (total >= 0),",
    "    discount     DECIMAL(5,2) DEFAULT 0.00,",
    "    note         TEXT,",
    "    created_at   TIMESTAMPTZ NOT NULL DEFAULT now(),",
    "    updated_at   TIMESTAMPTZ NOT NULL DEFAULT now()",
    ") ENGINE=rowstore;",
]:
    cp(line)
bl()

h2("8.3  ALTER TABLE——修改表结构")
for line in [
    "-- 添加列",
    "ALTER TABLE users ADD COLUMN phone VARCHAR(20);",
    "ALTER TABLE users ADD COLUMN score INT NOT NULL DEFAULT 0;",
    "",
    "-- 删除列",
    "ALTER TABLE users DROP COLUMN phone;",
    "ALTER TABLE users DROP COLUMN IF EXISTS legacy_field;",
    "",
    "-- 修改列类型",
    "ALTER TABLE users ALTER COLUMN age TYPE BIGINT;",
    "",
    "-- 设置/删除列默认值",
    "ALTER TABLE users ALTER COLUMN score SET DEFAULT 100;",
    "ALTER TABLE users ALTER COLUMN score DROP DEFAULT;",
    "",
    "-- 重命名表",
    "ALTER TABLE users RENAME TO customers;",
    "",
    "-- 重命名列",
    "ALTER TABLE users RENAME COLUMN username TO login_name;",
]:
    cp(line)
bl()

h2("8.4  DROP TABLE 和 TRUNCATE")
for line in [
    "-- 删除表（不可恢复）",
    "DROP TABLE users;",
    "DROP TABLE IF EXISTS users;",
    "",
    "-- 清空表数据（保留表结构，比逐行 DELETE 更快）",
    "TRUNCATE TABLE users;",
    "TRUNCATE TABLE users, orders;   -- 同时清空多张表",
]:
    cp(line)
bl()

h2("8.5  视图（VIEW）")
for line in [
    "-- 创建视图",
    "CREATE VIEW active_users AS",
    "    SELECT id, username, email, created_at",
    "    FROM users",
    "    WHERE is_active = true;",
    "",
    "-- 替换视图定义（原子操作）",
    "CREATE OR REPLACE VIEW active_users AS",
    "    SELECT id, username, email, created_at, last_login",
    "    FROM users",
    "    WHERE is_active = true;",
    "",
    "-- 删除视图",
    "DROP VIEW IF EXISTS active_users;",
    "",
    "-- 查询视图（与查询普通表相同）",
    "SELECT * FROM active_users WHERE username LIKE 'A%';",
]:
    cp(line)
bl()

h2("8.6  物化视图（MATERIALIZED VIEW）")
np("物化视图将查询结果持久化存储，适合计算开销大但数据不需要实时更新的场景。")
for line in [
    "-- 创建物化视图",
    "CREATE MATERIALIZED VIEW monthly_revenue AS",
    "    SELECT",
    "        date_trunc('month', created_at) AS month,",
    "        SUM(total) AS revenue,",
    "        COUNT(*) AS order_count",
    "    FROM orders",
    "    WHERE status = 'done'",
    "    GROUP BY 1",
    "    ORDER BY 1;",
    "",
    "-- 手动刷新物化视图（重新执行查询，替换所有数据）",
    "REFRESH MATERIALIZED VIEW monthly_revenue;",
    "",
    "-- 删除物化视图",
    "DROP MATERIALIZED VIEW IF EXISTS monthly_revenue;",
]:
    cp(line)
bl()

h2("8.7  序列（SEQUENCE）")
for line in [
    "-- 创建序列",
    "CREATE SEQUENCE order_seq",
    "    START WITH 1000",
    "    INCREMENT BY 1",
    "    MINVALUE 1000",
    "    NO MAXVALUE",
    "    CACHE 50;                -- 预取 50 个值，提升性能",
    "",
    "-- 获取下一个值（不可回滚，调用即消耗）",
    "SELECT nextval('order_seq');",
    "",
    "-- 获取当前会话最后一次 nextval 返回的值",
    "SELECT currval('order_seq');",
    "",
    "-- 设置序列当前值（用于数据迁移）",
    "SELECT setval('order_seq', 50000);",
    "",
    "-- 删除序列",
    "DROP SEQUENCE IF EXISTS order_seq;",
    "",
    "-- SERIAL/BIGSERIAL 语法糖（自动创建序列）",
    "CREATE TABLE invoices (id BIGSERIAL PRIMARY KEY, amount DECIMAL(12,2));",
    "-- 等价于：",
    "-- CREATE SEQUENCE invoices_id_seq;",
    "-- CREATE TABLE invoices (id BIGINT DEFAULT nextval('invoices_id_seq') PRIMARY KEY, ...);",
]:
    cp(line)
bl()

h2("8.8  枚举类型（ENUM）")
for line in [
    "-- 创建枚举类型",
    "CREATE TYPE order_status AS ENUM (",
    "    'pending', 'processing', 'shipped', 'delivered', 'cancelled', 'refunded'",
    ");",
    "",
    "-- 在表中使用枚举类型",
    "CREATE TABLE orders (",
    "    id      BIGSERIAL PRIMARY KEY,",
    "    status  order_status NOT NULL DEFAULT 'pending'",
    ");",
    "",
    "-- 插入枚举值",
    "INSERT INTO orders (status) VALUES ('processing');",
    "",
    "-- 枚举值比较",
    "SELECT * FROM orders WHERE status = 'pending';",
    "SELECT * FROM orders WHERE status IN ('pending', 'processing');",
    "",
    "-- 删除枚举类型（必须先删除使用该类型的列）",
    "DROP TYPE IF EXISTS order_status;",
]:
    cp(line)
bl()

h2("8.9  存储过程与用户自定义函数")
h3("8.9.1  SQL 语言函数")
for line in [
    "-- 创建简单 SQL 函数",
    "CREATE FUNCTION active_user_count() RETURNS BIGINT AS $$",
    "    SELECT COUNT(*) FROM users WHERE is_active = true;",
    "$$ LANGUAGE SQL STABLE;",
    "",
    "-- 带参数的 SQL 函数",
    "CREATE FUNCTION users_by_age(min_age INT, max_age INT)",
    "RETURNS TABLE(id BIGINT, username TEXT) AS $$",
    "    SELECT id, username FROM users",
    "    WHERE age BETWEEN min_age AND max_age;",
    "$$ LANGUAGE SQL;",
    "",
    "-- 调用函数",
    "SELECT active_user_count();",
    "SELECT * FROM users_by_age(18, 25);",
]:
    cp(line)
bl()

h3("8.9.2  PL/pgSQL 函数（支持控制流）")
for line in [
    "CREATE OR REPLACE FUNCTION transfer_funds(",
    "    from_account BIGINT,",
    "    to_account   BIGINT,",
    "    amount       DECIMAL(12,2)",
    ") RETURNS VOID AS $$",
    "DECLARE",
    "    from_balance DECIMAL(12,2);",
    "BEGIN",
    "    -- 检查余额",
    "    SELECT balance INTO from_balance FROM accounts WHERE id = from_account FOR UPDATE;",
    "    IF from_balance < amount THEN",
    "        RAISE EXCEPTION 'Insufficient funds: balance=%, need=%', from_balance, amount;",
    "    END IF;",
    "",
    "    -- 执行转账",
    "    UPDATE accounts SET balance = balance - amount WHERE id = from_account;",
    "    UPDATE accounts SET balance = balance + amount WHERE id = to_account;",
    "END;",
    "$$ LANGUAGE plpgsql;",
    "",
    "-- 调用存储过程",
    "CALL transfer_funds(1001, 1002, 500.00);",
    "",
    "-- 删除函数",
    "DROP FUNCTION IF EXISTS transfer_funds(BIGINT, BIGINT, DECIMAL);",
]:
    cp(line)
bl()

h2("8.10  触发器（TRIGGER）")
np("触发器需要在 falcon.toml 中启用（[triggers] enabled = true）。支持 BEFORE/AFTER INSERT/UPDATE/DELETE，FOR EACH ROW/STATEMENT。")
for line in [
    "-- 创建触发器函数（必须返回 trigger）",
    "CREATE OR REPLACE FUNCTION update_timestamp()",
    "RETURNS trigger AS $$",
    "BEGIN",
    "    NEW.updated_at = now();",
    "    RETURN NEW;",
    "END;",
    "$$ LANGUAGE plpgsql;",
    "",
    "-- 绑定触发器（每次 UPDATE 前自动更新时间戳）",
    "CREATE TRIGGER trg_update_timestamp",
    "BEFORE UPDATE ON orders",
    "FOR EACH ROW EXECUTE FUNCTION update_timestamp();",
    "",
    "-- 审计日志触发器",
    "CREATE OR REPLACE FUNCTION audit_changes() RETURNS trigger AS $$",
    "BEGIN",
    "    INSERT INTO audit_log (table_name, operation, changed_at, old_data, new_data)",
    "    VALUES (TG_TABLE_NAME, TG_OP, now(),",
    "            row_to_json(OLD)::JSONB, row_to_json(NEW)::JSONB);",
    "    RETURN NEW;",
    "END;",
    "$$ LANGUAGE plpgsql;",
    "",
    "CREATE TRIGGER trg_audit",
    "AFTER INSERT OR UPDATE OR DELETE ON orders",
    "FOR EACH ROW EXECUTE FUNCTION audit_changes();",
    "",
    "-- 删除触发器",
    "DROP TRIGGER IF EXISTS trg_audit ON orders;",
]:
    cp(line)
pb()

# ══ 第九章 DML ══
h1("第九章  数据操作语言（DML）")
h2("9.1  INSERT——数据插入")
for line in [
    "-- 单行插入（指定列名）",
    "INSERT INTO users (username, email, age, created_at)",
    "VALUES ('alice', 'alice@example.com', 28, now());",
    "",
    "-- 多行批量插入（单语句，原子操作）",
    "INSERT INTO products (name, price, category) VALUES",
    "    ('Laptop',   2999.00, 'electronics'),",
    "    ('Mouse',      29.99, 'electronics'),",
    "    ('Desk Chair', 499.00, 'furniture'),",
    "    ('Notebook',    4.99, 'stationery');",
    "",
    "-- INSERT ... SELECT（从查询结果插入，支持复杂子查询）",
    "INSERT INTO monthly_summary (month, total_orders, revenue)",
    "SELECT",
    "    date_trunc('month', created_at),",
    "    COUNT(*),",
    "    SUM(total)",
    "FROM orders",
    "WHERE status = 'done'",
    "GROUP BY 1;",
    "",
    "-- UPSERT：主键/唯一键冲突时更新",
    "INSERT INTO product_inventory (product_id, qty, updated_at)",
    "VALUES (101, 50, now())",
    "ON CONFLICT (product_id)",
    "DO UPDATE SET",
    "    qty = EXCLUDED.qty,",
    "    updated_at = EXCLUDED.updated_at;",
    "",
    "-- ON CONFLICT DO NOTHING：冲突时静默跳过",
    "INSERT INTO tags (name) VALUES ('featured'), ('new'), ('sale')",
    "ON CONFLICT (name) DO NOTHING;",
    "",
    "-- RETURNING：获取服务器生成的值",
    "INSERT INTO users (username, email)",
    "VALUES ('bob', 'bob@example.com')",
    "RETURNING id, created_at;",
]:
    cp(line)
bl()

h2("9.2  UPDATE——数据更新")
for line in [
    "-- 单列更新",
    "UPDATE users SET is_active = false WHERE last_login < now() - INTERVAL '90 days';",
    "",
    "-- 多列更新",
    "UPDATE users SET",
    "    email = 'newemail@example.com',",
    "    updated_at = now(),",
    "    login_count = login_count + 1",
    "WHERE id = 42;",
    "",
    "-- 使用子查询更新",
    "UPDATE orders SET status = 'cancelled'",
    "WHERE id IN (",
    "    SELECT order_id FROM payment_failures",
    "    WHERE failed_at < now() - INTERVAL '24 hours'",
    ");",
    "",
    "-- UPDATE ... FROM（关联表更新，PostgreSQL 语法）",
    "UPDATE orders o",
    "SET",
    "    status = 'shipped',",
    "    shipped_at = s.created_at",
    "FROM shipments s",
    "WHERE o.id = s.order_id",
    "  AND s.carrier_status = 'delivered'",
    "  AND o.status = 'processing';",
    "",
    "-- RETURNING（返回更新后的行数据）",
    "UPDATE accounts",
    "SET balance = balance - 100",
    "WHERE id = 1",
    "RETURNING id, balance;    -- 确认扣款后余额",
]:
    cp(line)
bl()

h2("9.3  DELETE——数据删除")
for line in [
    "-- 条件删除",
    "DELETE FROM sessions WHERE expires_at < now();",
    "",
    "-- 删除关联数据（DELETE ... USING）",
    "DELETE FROM order_items oi",
    "USING orders o",
    "WHERE oi.order_id = o.id",
    "  AND o.status = 'cancelled'",
    "  AND o.created_at < now() - INTERVAL '30 days';",
    "",
    "-- 使用子查询",
    "DELETE FROM users",
    "WHERE id IN (",
    "    SELECT user_id FROM flagged_accounts WHERE severity = 'critical'",
    ");",
    "",
    "-- RETURNING（返回被删除的行）",
    "DELETE FROM expired_tokens",
    "WHERE expires_at < now()",
    "RETURNING token, user_id;    -- 可用于日志记录",
]:
    cp(line)
bl()

h2("9.4  SELECT——查询数据")
h3("9.4.1  基础查询与过滤")
for line in [
    "-- 基本查询",
    "SELECT id, username, email FROM users WHERE is_active = true;",
    "",
    "-- 排序（多列）",
    "SELECT * FROM orders ORDER BY created_at DESC, total DESC LIMIT 20 OFFSET 0;",
    "",
    "-- DISTINCT 去重",
    "SELECT DISTINCT category FROM products ORDER BY category;",
    "",
    "-- 聚合 + 分组 + 过滤分组",
    "SELECT",
    "    category,",
    "    COUNT(*) AS product_count,",
    "    AVG(price) AS avg_price,",
    "    MAX(price) AS max_price",
    "FROM products",
    "GROUP BY category",
    "HAVING COUNT(*) >= 5",
    "ORDER BY avg_price DESC;",
]:
    cp(line)
bl()

h3("9.4.2  JOIN 连接查询")
for line in [
    "-- INNER JOIN：仅返回两表均有匹配的行",
    "SELECT u.username, o.id AS order_id, o.total",
    "FROM users u",
    "INNER JOIN orders o ON u.id = o.user_id",
    "WHERE o.status = 'done';",
    "",
    "-- LEFT JOIN：保留左表所有行",
    "SELECT u.username, COUNT(o.id) AS order_count",
    "FROM users u",
    "LEFT JOIN orders o ON u.id = o.user_id",
    "GROUP BY u.id, u.username",
    "ORDER BY order_count DESC;",
    "",
    "-- 多表 JOIN",
    "SELECT o.id, u.username, p.name AS product, oi.qty",
    "FROM orders o",
    "JOIN users u ON o.user_id = u.id",
    "JOIN order_items oi ON o.id = oi.order_id",
    "JOIN products p ON oi.product_id = p.id",
    "WHERE o.created_at > now() - INTERVAL '7 days';",
]:
    cp(line)
bl()

h3("9.4.3  子查询")
for line in [
    "-- IN 子查询",
    "SELECT * FROM users",
    "WHERE id IN (SELECT DISTINCT user_id FROM orders WHERE total > 1000);",
    "",
    "-- EXISTS 子查询（相关子查询）",
    "SELECT * FROM users u",
    "WHERE EXISTS (",
    "    SELECT 1 FROM orders o",
    "    WHERE o.user_id = u.id AND o.status = 'pending'",
    ");",
    "",
    "-- 标量子查询",
    "SELECT",
    "    username,",
    "    (SELECT COUNT(*) FROM orders WHERE user_id = u.id) AS total_orders",
    "FROM users u;",
]:
    cp(line)
bl()

h3("9.4.4  CTE（公用表表达式）")
for line in [
    "-- 普通 CTE",
    "WITH active_customers AS (",
    "    SELECT id, username FROM users WHERE is_active = true",
    "),",
    "recent_orders AS (",
    "    SELECT user_id, COUNT(*) AS cnt FROM orders",
    "    WHERE created_at > now() - INTERVAL '30 days'",
    "    GROUP BY user_id",
    ")",
    "SELECT ac.username, COALESCE(ro.cnt, 0) AS orders_last_30d",
    "FROM active_customers ac",
    "LEFT JOIN recent_orders ro ON ac.id = ro.user_id",
    "ORDER BY orders_last_30d DESC;",
    "",
    "-- 递归 CTE（树形结构遍历）",
    "WITH RECURSIVE employee_tree AS (",
    "    -- 锚定查询：根节点",
    "    SELECT id, name, manager_id, 0 AS depth",
    "    FROM employees WHERE manager_id IS NULL",
    "",
    "    UNION ALL",
    "",
    "    -- 递归查询：子节点",
    "    SELECT e.id, e.name, e.manager_id, et.depth + 1",
    "    FROM employees e",
    "    JOIN employee_tree et ON e.manager_id = et.id",
    ")",
    "SELECT id, name, depth, manager_id",
    "FROM employee_tree",
    "ORDER BY depth, name;",
]:
    cp(line)
bl()

h3("9.4.5  窗口函数")
for line in [
    "-- 排名窗口函数",
    "SELECT",
    "    username,",
    "    total_spent,",
    "    RANK() OVER (ORDER BY total_spent DESC) AS spending_rank,",
    "    DENSE_RANK() OVER (ORDER BY total_spent DESC) AS dense_rank,",
    "    ROW_NUMBER() OVER (ORDER BY total_spent DESC) AS row_num",
    "FROM user_stats;",
    "",
    "-- 分区窗口函数",
    "SELECT",
    "    category,",
    "    name,",
    "    price,",
    "    AVG(price) OVER (PARTITION BY category) AS category_avg,",
    "    price - AVG(price) OVER (PARTITION BY category) AS diff_from_avg",
    "FROM products;",
    "",
    "-- 滚动累计",
    "SELECT",
    "    order_date,",
    "    daily_revenue,",
    "    SUM(daily_revenue) OVER (",
    "        ORDER BY order_date",
    "        ROWS BETWEEN UNBOUNDED PRECEDING AND CURRENT ROW",
    "    ) AS cumulative_revenue",
    "FROM daily_sales;",
    "",
    "-- LEAD / LAG（前后行数据）",
    "SELECT",
    "    created_at,",
    "    amount,",
    "    LAG(amount) OVER (ORDER BY created_at) AS prev_amount,",
    "    amount - LAG(amount) OVER (ORDER BY created_at) AS change",
    "FROM transactions;",
]:
    cp(line)
bl()

h2("9.5  COPY——批量数据导入导出")
for line in [
    "-- 从 CSV 文件导入（服务器端路径）",
    "COPY users (username, email, age)",
    "FROM '/data/users.csv'",
    "WITH (FORMAT CSV, HEADER true, NULL '\\N', DELIMITER ',');",
    "",
    "-- 从客户端文件导入（psql \\copy，使用客户端路径）",
    "\\copy users (username, email) FROM 'C:/data/users.csv' WITH (FORMAT CSV, HEADER true)",
    "",
    "-- 导出整张表到 CSV",
    "COPY users TO '/data/users_backup.csv' WITH (FORMAT CSV, HEADER true);",
    "",
    "-- 导出查询结果",
    "COPY (",
    "    SELECT id, username, email, created_at",
    "    FROM users",
    "    WHERE is_active = true",
    "    ORDER BY id",
    ") TO '/data/active_users.csv'",
    "WITH (FORMAT CSV, HEADER true);",
]:
    cp(line)
pb()

# ══ 第十章 索引 ══
h1("第十章  索引管理")
h2("10.1  索引类型支持")
tbl(
    ["索引类型", "支持状态", "适用场景", "语法"],
    [
        ["B-tree（默认）", "✓ 完整支持", "等值查询、范围查询、ORDER BY", "CREATE INDEX ON t(col)"],
        ["GIN", "✓（全文搜索/JSONB/数组）", "全文搜索、JSONB、数组包含", "CREATE INDEX USING GIN"],
        ["哈希索引", "不支持", "—", "—"],
        ["GiST 索引", "不支持", "—", "—"],
        ["BRIN 索引", "不支持", "—", "—"],
    ]
)
bl()

h2("10.2  创建索引")
for line in [
    "-- 普通单列索引",
    "CREATE INDEX idx_users_email ON users (email);",
    "",
    "-- 唯一索引（同时创建唯一约束）",
    "CREATE UNIQUE INDEX idx_users_username ON users (username);",
    "",
    "-- 复合索引（列顺序很重要：高选择性列放前面）",
    "CREATE INDEX idx_orders_user_status ON orders (user_id, status);",
    "CREATE INDEX idx_orders_date ON orders (created_at DESC, user_id ASC);",
    "",
    "-- 条件索引（部分索引，只索引满足条件的行，节省存储）",
    "CREATE INDEX idx_pending_orders ON orders (created_at)",
    "WHERE status = 'pending';",
    "",
    "-- GIN 索引（JSONB 查询加速）",
    "CREATE INDEX idx_events_data ON events USING GIN(data);",
    "",
    "-- GIN 索引（全文搜索）",
    "CREATE INDEX idx_articles_body_fts",
    "ON articles USING GIN(to_tsvector('english', body));",
    "",
    "-- 函数索引（对表达式建索引）",
    "CREATE INDEX idx_users_email_lower ON users (lower(email));",
    "",
    "-- 查询时要匹配函数索引",
    "SELECT * FROM users WHERE lower(email) = 'alice@example.com';",
]:
    cp(line)
bl()

h2("10.3  删除索引")
for line in [
    "DROP INDEX idx_users_email;",
    "DROP INDEX IF EXISTS idx_users_email;",
]:
    cp(line)
bl()

h2("10.4  查看索引信息")
for line in [
    "-- 查看某表的所有索引",
    "SELECT indexname, indexdef FROM pg_indexes WHERE tablename = 'users';",
    "",
    "-- 查看表统计（包含估算行数）",
    "SHOW falcon.table_stats;",
    "",
    "-- 更新统计信息（优化器使用）",
    "ANALYZE TABLE users;",
]:
    cp(line)
bl()

h2("10.5  EXPLAIN 执行计划分析详解")
np("EXPLAIN 命令帮助理解查询的执行方式，是性能优化的核心工具。")
for line in [
    "-- 查看执行计划（不实际执行，速度快）",
    "EXPLAIN SELECT * FROM orders WHERE user_id = 100;",
    "",
    "-- 执行并显示实际统计（推荐用于生产调优）",
    "EXPLAIN ANALYZE SELECT * FROM orders WHERE user_id = 100 AND status = 'pending';",
    "",
    "-- 示例输出解读：",
    "-- IndexScan on orders  (cost=0.00..8.28 rows=1 width=72)",
    "--   actual time=0.015..0.018 rows=1 loops=1",
    "--   Index Cond: (user_id = 100)",
    "--   Filter: (status = 'pending')",
    "-- Planning Time: 0.12 ms",
    "-- Execution Time: 0.04 ms",
]:
    cp(line)
bl()
np("常见执行计划节点解读：")
tbl(
    ["节点类型", "说明", "适用场景", "优化建议"],
    [
        ["SeqScan", "全表顺序扫描", "小表或低选择性", "若大表出现，考虑建索引"],
        ["IndexScan", "B-tree 索引点查", "高选择性等值/范围查询", "理想情况，保持现状"],
        ["IndexRangeScan", "索引范围扫描", "BETWEEN、>、< 等", "确保索引覆盖查询条件"],
        ["NestedLoopJoin", "嵌套循环连接", "小外表驱动大内表（有索引）", "内表需有索引"],
        ["HashJoin", "哈希连接", "大表等值连接", "内存充足时高效"],
        ["MergeSortJoin", "归并排序连接", "两表已按连接键排序", "输入有序时最优"],
        ["Sort", "排序节点", "ORDER BY 无排序索引", "可建覆盖索引消除排序"],
        ["HashAggregate", "哈希聚合", "GROUP BY", "分组基数小时高效"],
    ]
)
bl()

h2("10.6  索引失效场景（避免误用）")
np("以下情况会导致索引无法被利用，应尽量避免：")
bp("对索引列使用函数或运算：WHERE YEAR(created_at) = 2024（应改为范围条件）")
bp("对索引列使用 NOT IN、!=、<>（部分情况优化器可利用，但选择性差时会全扫）")
bp("隐式类型转换：索引列是 INT，WHERE id = '123'（字符串与整数比较）")
bp("前导通配符：WHERE name LIKE '%alice%'（改为全文搜索或 LIKE 'alice%'）")
bp("OR 条件跨多列：WHERE col1 = 1 OR col2 = 2（考虑 UNION ALL 替代）")
bp("低选择性列：对 gender（只有 2 种值）建单列索引往往比全表扫描更慢")
pb()

# ══ 第十一章 安全 ══
h1("第十一章  安全管理")
h2("11.1  认证方式配置")
tbl(
    ["认证方式", "配置值", "传输安全", "适用环境"],
    [
        ["Trust（无认证）", "trust", "明文", "仅限本地开发，严禁生产使用"],
        ["Password（明文密码）", "password", "明文，需配合 TLS", "TLS 强制启用时可用"],
        ["MD5 散列", "md5", "散列（防嗅探）", "无 TLS 时的基本保护"],
        ["SCRAM-SHA-256", "scram-sha-256", "只传证明不传密码", "生产环境强烈推荐"],
    ]
)
bl()
for line in [
    "[server.auth]",
    'method = "scram-sha-256"             # 生产推荐',
    'username = "falcon"',
    'password = "${FALCON_DB_PASSWORD}"   # 通过环境变量传入，严禁硬编码',
    "",
    "# 多用户配置",
    "[[server.auth.users]]",
    'username = "app_user"',
    'password = "${APP_DB_PASSWORD}"',
    "",
    "[[server.auth.users]]",
    'username = "readonly_user"',
    'password = "${RO_DB_PASSWORD}"',
]:
    cp(line)
bl()

h2("11.2  SCRAM-SHA-256 认证流程")
np("SCRAM-SHA-256 是最安全的认证方式，整个握手过程不传输密码明文，仅传输密码的加密证明：")
tbl(
    ["步骤", "方向", "消息内容"],
    [
        ["1", "客户端 → 服务器", "StartupMessage（携带用户名）"],
        ["2", "服务器 → 客户端", "AuthenticationSASL（宣告 SCRAM-SHA-256 机制）"],
        ["3", "客户端 → 服务器", "SASLInitialResponse（client-first-message，含随机数）"],
        ["4", "服务器 → 客户端", "AuthenticationSASLContinue（server-first-message，含盐值和服务器随机数）"],
        ["5", "客户端 → 服务器", "SASLResponse（client-final-message，含客户端证明）"],
        ["6", "服务器 → 客户端", "AuthenticationSASLFinal（服务器签名验证）"],
        ["7", "服务器 → 客户端", "AuthenticationOk（认证成功）"],
    ]
)
bl()

h2("11.3  RBAC 权限管理")
for line in [
    "-- 创建角色（不可登录的权限组）",
    "CREATE ROLE readonly_group;",
    "CREATE ROLE readwrite_group;",
    "",
    "-- 创建可登录用户",
    "CREATE ROLE app_user LOGIN PASSWORD 'SecureP@ss123!';",
    "CREATE ROLE analyst LOGIN PASSWORD 'An@lyst456!';",
    "",
    "-- 授予表级权限",
    "GRANT SELECT ON ALL TABLES IN SCHEMA public TO readonly_group;",
    "GRANT SELECT, INSERT, UPDATE, DELETE ON orders TO readwrite_group;",
    "GRANT SELECT ON products, categories TO readwrite_group;",
    "",
    "-- 授予 Schema 级权限",
    "GRANT USAGE ON SCHEMA public TO readonly_group;",
    "GRANT ALL ON SCHEMA public TO admin_role;",
    "",
    "-- 将角色分配给用户（继承所有权限）",
    "GRANT readonly_group TO analyst;",
    "GRANT readwrite_group TO app_user;",
    "",
    "-- 撤销权限",
    "REVOKE DELETE ON orders FROM readwrite_group;",
    "REVOKE readwrite_group FROM app_user;",
    "",
    "-- 查看权限信息（psql 命令）",
    "\\dp orders          -- 查看 orders 表的权限",
    "\\du                 -- 查看所有角色",
]:
    cp(line)
bl()

h2("11.4  密码策略（默认值）")
tbl(
    ["参数", "默认值", "说明"],
    [
        ["min_length", "8", "密码最小长度"],
        ["require_uppercase", "true", "必须包含至少一个大写字母"],
        ["require_lowercase", "true", "必须包含至少一个小写字母"],
        ["require_digit", "true", "必须包含至少一个数字"],
        ["require_special", "false", "是否要求特殊字符（生产建议设为 true）"],
        ["max_age_days", "90", "密码有效期（天），0 = 永不过期"],
        ["history_count", "3", "记住最近 N 次密码，禁止重用"],
    ]
)
bl()

h2("11.5  认证暴力破解防护")
np("AuthRateLimiter 追踪每个源 IP 的认证失败次数，达到阈值后自动锁定：")
tbl(
    ["参数", "默认值", "说明"],
    [
        ["max_failures", "5", "锁定触发前的最大失败次数"],
        ["lockout_duration", "300 秒（5 分钟）", "IP 锁定持续时间"],
        ["failure_window", "600 秒（10 分钟）", "失败计数的时间窗口"],
        ["per_ip", "true", "按源 IP 独立追踪"],
    ]
)
bl()
np("成功认证后自动清零失败计数。锁定到期后自动重置。锁定期间的连接尝试会立即被拒绝，不消耗数据库资源。")
bl()

h2("11.6  TLS/mTLS 传输加密配置")
for line in [
    "[security.tls]",
    "enabled = true",
    'cert_path = "/etc/falcon/certs/server.crt"',
    'key_path  = "/etc/falcon/certs/server.key"',
    'ca_cert_path = "/etc/falcon/certs/ca.crt"   # 配置后启用 mTLS',
    "require_client_cert = false                   # true = 强制客户端证书（mTLS）",
    'min_version = "1.3"                           # 强烈建议 TLS 1.3',
]:
    cp(line)
bl()
tbl(
    ["环境", "TLS", "mTLS", "最低版本", "备注"],
    [
        ["开发环境", "可选", "不需要", "1.2", "本地环境可不启用"],
        ["测试/预发布", "必须", "可选", "1.2", "模拟生产配置"],
        ["生产环境", "必须", "推荐", "1.3", "应用/数据库间通信建议 mTLS"],
    ]
)
bl()

h2("11.7  透明数据加密（TDE）")
np("TDE 对磁盘上的所有 WAL 段文件和数据文件进行 AES-256-GCM 加密，使用 PBKDF2 密钥派生（per-scope DEK），支持密钥轮换和重新加密，对应用程序完全透明。")
for line in [
    "# 生产环境：将主密钥注入环境变量（不写入配置文件）",
    "export FALCON_TDE_KEY='$(openssl rand -base64 32)'",
    "",
    "[tde]",
    "enabled = true",
    'key_env_var = "FALCON_TDE_KEY"    # 主密钥从此环境变量读取',
    "encrypt_wal = true                # 加密所有 WAL 段",
    "encrypt_data = true               # 加密所有数据文件",
    "",
    "# 验证 TDE 状态",
    "SHOW falcon.security;",
]:
    cp(line)
bl()

h2("11.8  SQL 防火墙")
np("SQL 防火墙在每条 SQL 到达执行器之前进行安全检查，无法绕过（包括预处理语句参数注入场景）。")
tbl(
    ["防护层", "默认状态", "检测内容"],
    [
        ["注入检测", "启用", "永真条件（OR 1=1）、UNION 注入、注释截断（-- 后跟敏感模式）、pg_sleep() 时间盲注"],
        ["危险语句拦截", "启用", "非超级用户执行 DROP DATABASE、TRUNCATE 全表、ALTER SYSTEM"],
        ["语句堆叠检测", "启用", "包含分号的多语句注入（SELECT 1; DROP TABLE）"],
        ["长度限制", "1 MB", "单条 SQL 超过 1 MB 立即拒绝"],
        ["自定义规则", "空", "支持配置用户自定义阻断模式"],
    ]
)
bl()
np("超级用户豁免：超级用户可执行危险语句（如 DROP DATABASE），但注入检测对所有用户（包括超级用户）始终有效。")
bl()

h2("11.9  审计日志")
np("审计日志记录所有安全相关事件，采用环形缓冲区设计（4096 条事件，非阻塞写入）：")
for line in [
    "-- 查看最近审计事件（按时间倒序）",
    "SHOW falcon.audit_log;",
    "",
    "-- 字段说明：",
    "-- event_type: LOGIN / LOGOUT / AUTH_FAILURE / DDL / PRIVILEGE_CHANGE / ROLE_CHANGE / CONFIG_CHANGE",
    "-- ts_unix_ms: 事件时间戳（Unix 毫秒）",
    "-- username:   操作用户名",
    "-- client_ip:  客户端 IP 地址",
    "-- object:     操作对象（表名、角色名等）",
    "-- detail:     详细信息",
]:
    cp(line)
bl()

h2("11.10  生产安全检查清单")
tbl(
    ["序号", "检查项目", "验证命令/方法", "重要性"],
    [
        ["1", "认证方式为 scram-sha-256（非 trust）", "检查 falcon.toml [server.auth]", "★★★★★"],
        ["2", "TLS 已启用，版本 ≥ 1.2", "SHOW falcon.security", "★★★★★"],
        ["3", "密码复杂度策略已启用", "SHOW falcon.security_audit", "★★★★"],
        ["4", "认证速率限制已启用", "SHOW falcon.security_audit", "★★★★"],
        ["5", "SQL 防火墙已启用", "SHOW falcon.security_audit", "★★★★★"],
        ["6", "审计日志已启用", "SHOW falcon.audit_log", "★★★★"],
        ["7", "应用程序使用最小权限账户（非超级用户）", "检查角色和权限授予", "★★★★★"],
        ["8", "所有密码通过环境变量配置（非硬编码）", "检查 falcon.toml 和代码仓库", "★★★★★"],
        ["9", "IP 白名单已配置（如适用）", "SHOW falcon.security", "★★★"],
        ["10", "TDE 静态加密已启用（含 WAL 和数据文件）", "SHOW falcon.security", "★★★★"],
        ["11", "服务运行在最小权限账户下（非 root/Administrator）", "检查 Windows 服务账户", "★★★★"],
        ["12", "定期审查 SHOW falcon.audit_log 中的异常事件", "运维规程", "★★★"],
    ]
)
pb()

# ══ 第十二章 集群与复制 ══
h1("第十二章  集群与复制")
h2("12.1  集群架构概述")
np("FalconDB 分布式集群采用控制平面（Control Plane）与数据平面（Data Plane）分离的架构设计：")
bp("控制平面：3 节点 Raft 共识组，负责集群元数据管理（分片映射、节点注册、全局配置），可承受 1 节点故障")
bp("数据平面：N 个数据节点，每个节点承载多个分片（Shard）的主副本或从副本，通过 gRPC WAL 流式复制保持数据同步")
bp("Epoch 机制：每个分片维护单调递增的 Epoch 值，每次 Leader 切换时递增；所有写操作携带 Epoch 进行校验，防止脑裂")
bl()

h2("12.2  WAL 流式复制原理")
np("FalconDB 使用基于 gRPC 的 WAL 流式复制：")
tbl(
    ["组件", "说明"],
    [
        ["WalChunk", "WAL 传输单元，包含 LSN 范围（起始/结束）和 CRC32 校验码"],
        ["applied_lsn", "从节点向主节点报告已应用的最新 LSN，用于确认追踪"],
        ["ack_lsn", "主节点记录的从节点最新确认 LSN，断连后从 ack_lsn+1 恢复传输"],
        ["commit_ts", "提交时间戳，用于 MVCC 可见性判断和 GC 安全点计算"],
    ]
)
bl()

h2("12.3  提交确认策略")
tbl(
    ["策略名称", "ACK 要求", "典型提交延迟", "持久性级别", "适用场景"],
    [
        ["Local（本地）", "仅主节点本地 fsync", "100-500 μs", "单节点崩溃安全", "性能优先，可接受部分数据丢失风险"],
        ["Quorum（多数派）", "⌊N/2⌋+1 个节点 ACK", "1-5 ms（局域网）", "多数节点崩溃安全", "平衡性能与持久性"],
        ["All（全部）", "所有 N 个节点 ACK", "最高，等于最慢副本", "全节点崩溃安全", "金融/支付等零丢失场景"],
    ]
)
bl()
for line in [
    "[replication]",
    'commit_ack = "quorum"    # local / quorum / all',
    'role = "primary"         # primary / replica / standalone / raft_member',
]:
    cp(line)
bl()

h2("12.4  双层 Epoch 脑裂防护")
np("FalconDB 通过两层独立的 Epoch 校验防止脑裂（Split-Brain）场景，确保旧主节点（已被隔离的节点）无法继续写入数据：")
tbl(
    ["防护层", "组件", "位置", "行为"],
    [
        ["集群层", "SplitBrainDetector", "falcon_cluster", "拒绝 writer_epoch < current_epoch 的 RPC 请求"],
        ["存储层", "StorageEpochFence", "falcon_storage", "拒绝 writer_epoch < current_epoch 的 WAL 写入"],
    ]
)
bl()
for line in [
    "-- 被拒绝时日志记录：",
    "-- [ERROR] split-brain rejected: writer node 1 epoch 3 < current epoch 5",
    "-- [ERROR] EPOCH FENCE: stale-epoch WAL write rejected at storage barrier",
    "",
    "-- Prometheus 指标（epoch 隔离触发次数）",
    "-- falcon_epoch_fence_rejections_total",
]:
    cp(line)
bl()

h2("12.5  复制一致性不变量")
np("FalconDB 严格维护以下复制一致性不变量，任何违反均视为系统 Bug：")
tbl(
    ["不变量", "说明"],
    [
        ["前缀属性（Prefix Property）", "replica_committed_lsn ≤ primary_committed_lsn。从节点的已提交 LSN 绝不超过主节点"],
        ["无幻象提交（No Phantom Commits）", "replica_visible_lsn ≤ primary_durable_lsn。从节点不对主节点未持久化的事务可见"],
        ["可见性绑定（Visibility Binding）", "客户端可见性与提交策略绑定（Local/Quorum/All）"],
    ]
)
bl()

h2("12.6  复制监控命令")
for line in [
    "-- 主节点：查看所有副本状态和复制延迟",
    "SHOW falcon.replication_stats;",
    "",
    "-- 从节点：查看本节点复制进度",
    "SHOW falcon.replica_stats;",
    "",
    "-- 查看分布式查询 scatter/gather 统计",
    "SHOW falcon.scatter_stats;",
    "",
    "-- Prometheus 指标",
    "-- falcon_replication_lag_lsn        → 副本延迟（LSN 单位）",
    "-- falcon_replication_applied_lsn    → 已应用 LSN",
    "-- falcon_replication_reconnects_total → 重连次数",
]:
    cp(line)
pb()

# ══ 第十三章 故障转移 ══
h1("第十三章  故障转移与高可用")
h2("13.1  故障转移阶段详解")
tbl(
    ["阶段", "操作", "时间目标", "失败处理"],
    [
        ["detect_failure", "心跳超时触发故障检测（默认 3×heartbeat_interval_ms）", "< 3 秒", "超时后继续推进"],
        ["freeze_writes", "拒绝该分片所有新写请求", "< 100 ms", "等待所有飞行中写入完成"],
        ["seal_epoch", "递增 Epoch，向所有节点广播新 Epoch，隔离旧主节点", "< 100 ms", "广播失败则重试"],
        ["catch_up", "选出最优从节点，应用其与主节点间的 WAL 差距", "取决于复制延迟", "差距超阈值则告警"],
        ["promote", "原子切换从节点角色为主节点，更新分片路由表", "< 100 ms", "切换失败则回滚"],
        ["reopen", "新主节点开始接受写请求", "< 100 ms", "通知所有客户端重连"],
        ["verify", "验证所有一致性不变量（前缀、幻象、提升安全）", "< 500 ms", "不变量违反触发告警"],
        ["complete", "写入审计事件，更新 Prometheus 指标", "< 10 ms", ""],
    ]
)
bl()

h2("13.2  故障转移审计跟踪")
np("每次故障转移全过程均生成结构化审计事件，可通过以下命令查询：")
for line in [
    "-- 查看复制和故障转移指标（含最近一次故障转移持续时间）",
    "SHOW falcon.replication_stats;",
    "",
    "-- 查看审计日志中的故障转移事件",
    "SHOW falcon.audit_log;",
    "",
    "-- 每条审计事件包含：",
    "-- seq            → 单调递增序号",
    "-- stage          → 故障转移阶段（见上表）",
    "-- timestamp_ms   → 阶段开始的 Unix 毫秒时间戳",
    "-- elapsed_us     → 从故障转移开始到本阶段的累计微秒数",
    "-- success        → 本阶段是否成功",
    "-- description    → 人可读描述",
]:
    cp(line)
bl()

h2("13.3  RPO / RTO 参考指标")
tbl(
    ["部署模式", "RPO（最大数据丢失）", "RTO（恢复时间目标）", "适用场景"],
    [
        ["单节点 + local-fsync WAL", "0（本地持久化，DCG 保证）", "秒级（WAL 重放）", "开发/测试"],
        ["主从 + local-fsync", "最多丢失最后一批未复制的 WAL", "< 30 秒（自动故障转移）", "标准生产"],
        ["主从 + sync-replica", "≈ 0（主节点响应时已在副本持久化）", "< 30 秒", "金融/支付场景"],
        ["三节点 Raft + sync-replica", "0（Quorum 写入保证）", "< 10 秒（Raft 选举）", "最高可用性"],
        ["全量备份恢复（无 WAL 归档）", "到上次备份时刻", "分钟至小时（取决于数据量）", "灾难恢复兜底"],
        ["全量备份 + WAL 归档（PITR）", "秒级（最后一个 WAL 段）", "分钟级 + WAL 重放时间", "精确恢复场景"],
    ]
)
bl()

h2("13.4  故障转移演练")
for line in [
    "# 端到端故障转移测试（Linux/macOS）",
    "# 步骤：写入 1000 条记录 → kill -9 主节点 → 验证从节点提升 → 验证零数据丢失",
    "./scripts/e2e_two_node_failover.sh",
    "",
    "# Windows PowerShell",
    r".\scripts\e2e_two_node_failover.ps1",
    "",
    "# 主从复制演示",
    "./scripts/demo_replication.sh",
    "",
    "# 在负载下故障（持续写入期间 kill -9）",
    "# 参见 pocs/falcondb-poc-failover-under-load/",
]:
    cp(line)
bl()

h2("13.5  客户端重连最佳实践")
np("故障转移期间客户端会短暂收到连接拒绝或事务错误（40001），应用程序应实现自动重连逻辑：")
bp("使用支持 HA failover 的驱动（如 falcondb-jdbc、falcondb-go 内置 HA 故障转移）")
bp("对所有数据库操作包裹重试逻辑（指数退避，建议 3 次以上重试）")
bp("区分可重试错误（40001 OCC 冲突、连接断开）和不可重试错误（23505 唯一键冲突）")
bp("设置合理的连接池重连策略（最小 3 节点 seed list，支持动态路由更新）")
pb()

print("Part 2 完成，保存中...")
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
