# -*- coding: utf-8 -*-
"""Part 5 - 补充详细内容：集群运维/安全深度/SQL深度/事务深度/存储引擎深度"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn; from docx.oxml import OxmlElement

OUTPUT = r"c:\falcondb\falcondb-main\docs\FalconDB_操作手册.docx"
doc = Document(OUTPUT)

def h1(t): return doc.add_heading(t, 1)
def h2(t): return doc.add_heading(t, 2)
def h3(t): return doc.add_heading(t, 3)
def h4(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12)
    return p
def np(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(11)
    return p
def bp(t):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t); r.font.size = Pt(11); return p
def cp(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(t); r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0F0F0')
    pPr.append(shd); return p
def bl(): doc.add_paragraph()
def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        if c.paragraphs[0].runs:
            run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = c._tc; tcPr = tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        r = t.rows[ri+1]
        for ci,v in enumerate(row):
            r.cells[ci].text = str(v)
            if r.cells[ci].paragraphs[0].runs:
                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return t
def pb(): doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第二十章  集群管理深度指南
# ═══════════════════════════════════════════════════
h1("第二十章  集群管理深度指南")

h2("20.1  集群架构详解")
np("FalconDB 集群由控制平面和数据平面两层组成，各司其职：")
bp("控制平面：由 3 个或 5 个控制节点组成 Raft 共识组，负责管理集群元数据——分片映射表（Shard Map）、节点注册表、配置版本、Leader 选举。控制平面可以容忍 (N-1)/2 个节点故障，即 3 节点集群容忍 1 个节点故障。")
bp("数据平面：每个数据节点承载多个分片（Shard）的主副本（Primary）或从副本（Replica）。主副本接受读写请求，从副本通过 gRPC WAL 流式复制保持同步，处理只读查询。")
bp("客户端路由：客户端连接到任意节点，由节点内置的路由层（RoutingLayer）根据 Shard Map 将请求转发到正确的主节点。")
bl()

h2("20.2  分片（Shard）管理")
h3("20.2.1  分片概念")
np("分片是 FalconDB 中数据分布和复制的基本单位。每张表的数据按分片键（Shard Key）范围分布在不同的分片上。未指定分片键的表（默认）存储在单个分片（shard_id = 0）中。")
for line in [
    "-- 查看分片分布（集群模式）",
    "SHOW falcon.scatter_stats;",
    "",
    "-- 示例输出：",
    "-- scatter_count | gather_count | avg_scatter_us | total_scatter_bytes",
    "-- 45231         | 45231        | 234            | 892341234",
]: cp(line)
bl()

h3("20.2.2  分片自动再均衡")
np("FalconDB v1.2 新增了分片自动再均衡（Auto-Rebalancing）功能。当检测到分片负载不均衡时（如新增节点后），系统自动将部分分片迁移到负载较低的节点，迁移过程对应用透明：")
bp("触发条件：节点间分片数量差 > 20%，或节点间写入吞吐差 > 30%")
bp("迁移过程：后台逐步复制分片数据到目标节点（不影响正常读写），完成后原子切换路由")
bp("验证方式：已通过 47 项混沌测试，包括迁移期间故障注入、并发读写等极端场景")
bl()

h2("20.3  Raft 共识配置")
np("在多节点集群中，FalconDB 使用 Raft 共识算法保证控制平面的高可用性：")
for line in [
    "# 三节点 Raft 集群配置示例（每个节点的 falcon.toml）",
    "",
    "# 节点 1（Leader 候选）",
    "[raft]",
    "node_id = 1",
    "node_ids = [1, 2, 3]",
    'raft_listen_addr = "0.0.0.0:50052"',
    'peers = [',
    '    "2=http://192.168.1.2:50052",',
    '    "3=http://192.168.1.3:50052",',
    ']',
    "heartbeat_interval_ms = 50          # 心跳间隔（太小=网络压力大，太大=故障检测慢）",
    "election_timeout_min_ms = 150       # 选举超时最小值（应为 heartbeat × 3）",
    "election_timeout_max_ms = 300       # 选举超时最大值（随机化防止同时选举）",
    "",
    "# 节点 2（同样配置，node_id = 2）",
    "[raft]",
    "node_id = 2",
    "node_ids = [1, 2, 3]",
    'raft_listen_addr = "0.0.0.0:50052"',
    'peers = [',
    '    "1=http://192.168.1.1:50052",',
    '    "3=http://192.168.1.3:50052",',
    ']',
]: cp(line)
bl()

h2("20.4  成员变更操作")
np("向集群添加或移除节点需要通过成员变更（Membership Change）协议完成，以保证 Raft 安全性：")
h3("20.4.1  添加新节点")
for line in [
    "# 步骤 1：准备新节点（安装 FalconDB，配置 falcon.toml）",
    "# 新节点的 falcon.toml 中添加 join 配置",
    "[cluster]",
    'join_addr = "http://192.168.1.1:8080"   # 现有集群任意节点的管理地址',
    "",
    "# 步骤 2：启动新节点",
    "falcon service start",
    "",
    "# 步骤 3：通过 CLI 将新节点加入集群",
    "falconctl cluster add-node --addr http://192.168.1.4:50051 --id 4",
    "",
    "# 步骤 4：验证节点已加入",
    "falconctl cluster status",
    "SHOW falcon.replication_stats;",
]: cp(line)
bl()

h3("20.4.2  移除节点")
for line in [
    "# 步骤 1：先将节点上的分片迁移到其他节点",
    "falconctl cluster drain-node --id 3",
    "",
    "# 步骤 2：等待分片迁移完成（监控 scatter_stats）",
    "SHOW falcon.scatter_stats;",
    "",
    "# 步骤 3：从集群成员中移除节点",
    "falconctl cluster remove-node --id 3",
    "",
    "# 步骤 4：停止被移除节点的服务",
    "falcon service stop   # 在节点 3 上执行",
]: cp(line)
bl()

h2("20.5  跨数据中心部署")
np("FalconDB 支持跨数据中心（Multi-DC）部署，但需要权衡延迟和一致性：")
tbl(["部署模式","延迟要求","一致性","适用场景"],
[["单 DC 主从（推荐）","< 1 ms RTT","强一致","同城高可用，延迟最低"],
 ["双 DC 主从（Quorum）","< 5 ms RTT","强一致（多数派）","同城双中心，平衡延迟和可用性"],
 ["双 DC 主从（Async）","任意","最终一致","异地容灾，主 DC 故障后可能丢少量数据"],
 ["三 DC Raft（生产最高可用）","< 10 ms RTT","强一致（Raft Quorum）","异地三中心，任意单 DC 故障不影响可用性"]])
bl()

h2("20.6  Epoch 双层脑裂防护详解")
np("脑裂（Split-Brain）是分布式数据库中最危险的故障场景，可能导致多个节点同时认为自己是主节点并接受写入，最终造成数据分歧（Data Divergence）。FalconDB 通过双层独立 Epoch 机制在两个不同层次拦截过期写入：")
h4("第一层：集群层（SplitBrainDetector）")
np("位于网络 RPC 层，每个 Raft RPC 携带当前 Epoch。接收方对比本地已知最新 Epoch，如果请求的 Epoch 过期，立即拒绝并记录：")
for line in [
    "# 日志示例（集群层拦截）",
    "[WARN] split-brain rejected: writer node 1 epoch 3 < current epoch 5",
    "[INFO] SplitBrainDetector: rejected stale RPC from node=1 epoch=3 current=5",
]: cp(line)
bl()
h4("第二层：存储层（StorageEpochFence）")
np("位于 WAL 写入路径，每次写入 WAL 前校验 Epoch。即使网络层的拦截失效（如极端的时钟偏差或实现缺陷），存储层也会作为最后防线：")
for line in [
    "# 日志示例（存储层拦截）",
    "[ERROR] EPOCH FENCE: stale-epoch WAL write rejected at storage barrier",
    "[WARN] StorageEpochFence: write rejected epoch=old_epoch fence=new_epoch",
]: cp(line)
np("这种双层防护确保即使在最极端的网络分区和延迟场景下，也不会发生脑裂写入，从而保证数据一致性。")
bl()

h2("20.7  WAL 复制协议详解")
np("FalconDB 使用基于 gRPC 的流式 WAL 复制协议，主节点将 WAL 段文件（每段默认 64MB）实时流式发送给所有从节点：")
for line in [
    "# WAL 复制协议主要步骤",
    "",
    "1. 主节点写入 WAL（WalSegment）",
    "2. 主节点通过 gRPC ReplicateWal RPC 流式发送 WalChunk",
    "   每个 WalChunk = 64KB WAL 数据块 + CRC32 校验",
    "3. 从节点接收 WalChunk，验证 CRC32，写入本地 WAL",
    "4. 从节点重放 WAL 更新本地 StorageEngine（重放=反序列化操作并执行）",
    "5. 从节点发送 AppliedLsn 反馈给主节点",
    "6. 主节点收到 Quorum（多数派）ACK 后，返回客户端提交成功",
    "",
    "# 监控复制进度",
    "SHOW falcon.replication_stats;",
    "-- applied_lsn: 从节点已应用的 LSN（越接近主节点越好）",
    "-- lag_lsn: 主从 LSN 差距（0 = 完全同步）",
    "-- lag_bytes: 主从数据差距（字节）",
    "-- reconnects: 重连次数（非 0 说明网络不稳定）",
]: cp(line)
pb()

# ═══════════════════════════════════════════════════
# 第二十一章  安全管理深度指南
# ═══════════════════════════════════════════════════
h1("第二十一章  安全管理深度指南")

h2("21.1  认证系统深度解析")
h3("21.1.1  SCRAM-SHA-256 认证流程")
np("SCRAM-SHA-256（Salted Challenge Response Authentication Mechanism）是 FalconDB 生产环境推荐的认证机制。其核心优势是密码从不以明文传输，即使 TLS 被破解也无法获取密码：")
for line in [
    "# SCRAM-SHA-256 握手流程（6 步交换）",
    "",
    "Client → Server: client-first-message",
    "   包含: n,,n=username,r=client_nonce",
    "   （client_nonce = 32 字节随机值）",
    "",
    "Server → Client: server-first-message",
    "   包含: r=client_nonce+server_nonce, s=salt_base64, i=iterations",
    "   （salt = 16 字节随机值, iterations = 4096 次 PBKDF2 迭代）",
    "",
    "Client → Server: client-final-message",
    "   包含: c=channel_binding, r=combined_nonce, p=client_proof",
    "   client_proof = HMAC-SHA256(StoredKey, auth_message)",
    "   StoredKey = SHA256(HMAC-SHA256(salted_password, 'Client Key'))",
    "",
    "Server → Client: server-final-message",
    "   包含: v=server_signature（客户端验证服务端身份）",
    "",
    "# 安全特性：",
    "# 1. 密码永远不以明文传输（即使 TLS 被破解也安全）",
    "# 2. 每次认证使用不同的 nonce（防止重放攻击）",
    "# 3. 服务端身份也被验证（防止中间人攻击）",
]: cp(line)
bl()

h3("21.1.2  认证速率限制（防暴力破解）")
np("FalconDB 内置基于令牌桶（Token Bucket）算法的认证速率限制，自动防御暴力破解攻击：")
for line in [
    "[security.rate_limiting]",
    "enabled = true",
    "max_failures = 5           # 连续失败 5 次后锁定",
    "lockout_duration_secs = 300  # 锁定 300 秒（5 分钟）",
    "window_secs = 60           # 60 秒内的失败计数窗口",
    "",
    "# 监控速率限制状态",
    "SHOW falcon.security_audit;",
    "-- auth_checks:   总认证请求数",
    "-- auth_failures: 失败次数",
    "-- auth_lockouts: 触发锁定的 IP 数",
]: cp(line)
bl()

h2("21.2  行级安全（RLS）")
np("行级安全（Row-Level Security）允许基于当前用户的属性限制其可见的数据行，实现数据隔离，适用于多租户 SaaS 场景：")
for line in [
    "-- 创建策略：用户只能看到自己的订单",
    "ALTER TABLE orders ENABLE ROW LEVEL SECURITY;",
    "",
    "CREATE POLICY orders_user_isolation ON orders",
    "    FOR ALL",
    "    TO app_user_role",
    "    USING (user_id = current_user_id());",
    "",
    "-- 验证：以 alice 用户连接后",
    "-- SELECT * FROM orders; -- 只返回 alice 的订单",
    "",
    "-- 超级用户绕过 RLS（管理操作）",
    "SET row_security = off;   -- 需要超级用户权限",
    "SELECT * FROM orders;     -- 返回所有订单",
    "",
    "-- 创建更复杂的策略：部门隔离",
    "CREATE POLICY dept_isolation ON employees",
    "    USING (department_id = get_current_user_dept());",
    "",
    "-- 查看已创建的 RLS 策略",
    "SELECT tablename, policyname, cmd, qual",
    "FROM pg_policies;",
]: cp(line)
bl()

h2("21.3  TLS 证书管理")
h3("21.3.1  生成自签名证书（开发/测试）")
for line in [
    "# 生成 CA 根证书",
    "openssl genrsa -out ca.key 4096",
    "openssl req -new -x509 -key ca.key -out ca.crt -days 3650",
    '    -subj "/C=CN/ST=Beijing/O=FalconDB/CN=FalconDB-CA"',
    "",
    "# 生成服务器证书",
    "openssl genrsa -out server.key 4096",
    "openssl req -new -key server.key -out server.csr",
    '    -subj "/C=CN/ST=Beijing/O=FalconDB/CN=db.example.com"',
    "openssl x509 -req -in server.csr -CA ca.crt -CAkey ca.key",
    "    -CAcreateserial -out server.crt -days 365",
    "",
    "# 生成客户端证书（mTLS）",
    "openssl genrsa -out client.key 4096",
    "openssl req -new -key client.key -out client.csr",
    '    -subj "/C=CN/ST=Beijing/O=FalconDB/CN=app-client"',
    "openssl x509 -req -in client.csr -CA ca.crt -CAkey ca.key",
    "    -CAcreateserial -out client.crt -days 365",
    "",
    "# 将证书复制到 FalconDB 配置目录",
    "Copy-Item server.crt, server.key, ca.crt C:\\ProgramData\\FalconDB\\certs\\",
]: cp(line)
bl()

h3("21.3.2  证书轮换")
np("证书即将过期前（建议提前 30 天）执行证书轮换，FalconDB 支持热轮换（不中断服务）：")
for line in [
    "# 步骤 1：生成新证书（见上节）",
    "",
    "# 步骤 2：备份旧证书",
    r"Copy-Item C:\ProgramData\FalconDB\certs\server.crt",
    r"    C:\ProgramData\FalconDB\certs\server.crt.bak.$(Get-Date -Format yyyyMMdd)",
    "",
    "# 步骤 3：替换证书文件",
    r"Copy-Item new_server.crt C:\ProgramData\FalconDB\certs\server.crt",
    r"Copy-Item new_server.key C:\ProgramData\FalconDB\certs\server.key",
    "",
    "# 步骤 4：通过 falconctl 热更新 TLS（无需重启）",
    "falconctl tls reload",
    "",
    "# 步骤 5：验证新证书已生效",
    "openssl s_client -connect 127.0.0.1:5433 -CAfile ca.crt",
    "# 查看输出中的 certificate 部分，确认过期时间已更新",
]: cp(line)
bl()

h3("21.3.3  证书有效期监控")
for line in [
    "# 检查证书有效期（Windows PowerShell）",
    r'$cert = New-Object System.Security.Cryptography.X509Certificates.X509Certificate2',
    r'    "C:\ProgramData\FalconDB\certs\server.crt"',
    '$days = ($cert.NotAfter - (Get-Date)).Days',
    'Write-Host "证书还有 $days 天到期"',
    'if ($days -lt 30) { Write-Warning "证书即将到期，请尽快更新！" }',
    "",
    "# Linux：",
    "openssl x509 -in /etc/falcon/certs/server.crt -noout -dates",
    "# notAfter=Mar 12 00:00:00 2027 GMT",
]: cp(line)
bl()

h2("21.4  透明数据加密（TDE）详解")
np("TDE（Transparent Data Encryption）对 FalconDB 存储的所有数据文件和 WAL 进行 AES-256-GCM 加密。对应用层完全透明——SQL 查询正常执行，加解密在存储引擎内部自动完成。")
h3("21.4.1  加密算法细节")
bp("加密算法：AES-256-GCM（Galois/Counter Mode），同时提供机密性和完整性验证")
bp("密钥派生：PBKDF2-HMAC-SHA256，100,000 次迭代，随机盐值（防彩虹表攻击）")
bp("主密钥存储：仅通过环境变量传入，从不写入磁盘配置文件（防止密钥泄露）")
bp("加密粒度：每个 WAL 段文件和数据块独立加密，使用不同的随机 IV（防止 IV 重用攻击）")
for line in [
    "# 安全地生成和设置主密钥",
    '$key = [Convert]::ToBase64String((New-Object System.Security.Cryptography.RNGCryptoServiceProvider).GetBytes(32))',
    '$env:FALCON_TDE_KEY = $key',
    "",
    "# 在 Windows 服务中持久化环境变量",
    "[System.Environment]::SetEnvironmentVariable('FALCON_TDE_KEY', $key, 'Machine')",
    "",
    "# Linux：",
    "export FALCON_TDE_KEY=$(openssl rand -base64 32)",
    "# 将上述命令加入 /etc/falcondb/environment 文件",
    "",
    "# falcon.toml 配置",
    "[tde]",
    "enabled = true",
    'key_env_var = "FALCON_TDE_KEY"',
    "encrypt_wal = true",
    "encrypt_data = true",
    "",
    "# 验证 TDE 已启用",
    "SHOW falcon.security;",
    "-- encryption_enabled: true",
    "-- encryption_algorithm: AES-256-GCM",
]: cp(line)
bl()

h2("21.5  审计日志完整参考")
np("FalconDB 企业版审计日志记录所有安全相关事件，支持合规审计（如 SOC2、ISO 27001、PCI DSS）：")
tbl(["事件类别（event_type）","触发场景","记录内容"],
[["LOGIN_SUCCESS","用户成功登录","用户名、客户端 IP、认证方式、时间戳"],
 ["LOGIN_FAILURE","认证失败（密码错误、账户不存在）","用户名、客户端 IP、失败原因、时间戳"],
 ["LOGIN_LOCKOUT","IP 因多次失败被锁定","客户端 IP、失败次数、锁定时长"],
 ["DDL_EXECUTE","执行 DDL（CREATE/DROP/ALTER）","SQL 文本、执行用户、影响对象、时间戳"],
 ["PRIVILEGE_CHANGE","GRANT/REVOKE 权限变更","变更类型、被操作角色、权限对象"],
 ["PASSWORD_CHANGE","ALTER ROLE ... PASSWORD","操作用户、目标用户、时间戳"],
 ["ROLE_CREATE","CREATE ROLE","角色名、创建用户"],
 ["ROLE_DROP","DROP ROLE","角色名、操作用户"],
 ["SQL_INJECTION","SQL 防火墙检测到注入尝试","SQL 文本（前 200 字符）、客户端 IP、检测类型"],
 ["FIREWALL_BLOCKED","防火墙拦截危险操作","SQL 文本、操作类型、客户端 IP"],
 ["BACKUP_RESTORE","备份/恢复操作","操作类型、目标路径、字节数、耗时"],
 ["TLS_HANDSHAKE","TLS 握手（连接建立）","客户端 IP、TLS 版本、密码套件"],
 ["CONFIG_CHANGE","配置变更","变更键、旧值/新值（敏感值脱敏）、操作用户"]])
bl()
for line in [
    "-- 查询最近 100 条审计事件",
    "SHOW falcon.audit_log;",
    "",
    "-- 审计日志也写入磁盘文件（JSON 格式）",
    "# 路径：[logging.audit_file] 配置的路径",
    "# 每条记录格式：",
    '{',
    '  "ts_unix_ms": 1742000000000,',
    '  "event_type": "LOGIN_SUCCESS",',
    '  "username": "alice",',
    '  "client_ip": "192.168.1.100",',
    '  "auth_method": "scram-sha-256",',
    '  "session_id": 12345,',
    '  "detail": "authenticated successfully"',
    '}',
]: cp(line)
pb()

# ═══════════════════════════════════════════════════
# 第二十二章  SQL 语言深度参考
# ═══════════════════════════════════════════════════
h1("第二十二章  SQL 语言深度参考")

h2("22.1  窗口函数完整参考")
np("窗口函数（Window Functions）在保留原始行的同时对行集合执行计算，是 SQL 中最强大的分析功能之一：")
for line in [
    "-- 语法结构",
    "function_name(args) OVER (",
    "    [PARTITION BY partition_cols]  -- 分组（可选）",
    "    [ORDER BY sort_cols]           -- 排序（可选）",
    "    [frame_clause]                 -- 窗口帧（可选）",
    ")",
    "",
    "-- 排名函数",
    "SELECT",
    "    username,",
    "    total_spent,",
    "    ROW_NUMBER() OVER (ORDER BY total_spent DESC) AS row_num,    -- 唯一序号",
    "    RANK()       OVER (ORDER BY total_spent DESC) AS rank,       -- 并列排名（跳号）",
    "    DENSE_RANK() OVER (ORDER BY total_spent DESC) AS dense_rank, -- 并列排名（不跳号）",
    "    NTILE(4)     OVER (ORDER BY total_spent DESC) AS quartile,   -- 分四等分",
    "    PERCENT_RANK() OVER (ORDER BY total_spent DESC) AS pct_rank  -- 百分位排名",
    "FROM users;",
    "",
    "-- 分组内排名",
    "SELECT",
    "    department,",
    "    employee_name,",
    "    salary,",
    "    RANK() OVER (PARTITION BY department ORDER BY salary DESC) AS dept_rank",
    "FROM employees",
    "WHERE RANK() OVER (PARTITION BY department ORDER BY salary DESC) <= 3;",
    "-- 注：上面的写法无效（窗口函数不能在 WHERE 中），需用 CTE",
    "",
    "-- 正确写法：用 CTE 包装",
    "WITH ranked AS (",
    "    SELECT department, employee_name, salary,",
    "        RANK() OVER (PARTITION BY department ORDER BY salary DESC) AS dept_rank",
    "    FROM employees",
    ")",
    "SELECT * FROM ranked WHERE dept_rank <= 3;",
    "",
    "-- 聚合窗口函数",
    "SELECT",
    "    order_date,",
    "    daily_revenue,",
    "    SUM(daily_revenue) OVER (ORDER BY order_date) AS cumulative_revenue,  -- 累计",
    "    AVG(daily_revenue) OVER (ORDER BY order_date ROWS BETWEEN 6 PRECEDING AND CURRENT ROW)",
    "        AS rolling_7day_avg,   -- 7 日滚动平均",
    "    LAG(daily_revenue, 1) OVER (ORDER BY order_date) AS prev_day,         -- 上一行",
    "    LEAD(daily_revenue, 1) OVER (ORDER BY order_date) AS next_day          -- 下一行",
    "FROM daily_sales;",
]: cp(line)
bl()

h2("22.2  高级 CTE 使用场景")
for line in [
    "-- 递归 CTE：组织架构树（求员工所有上级）",
    "WITH RECURSIVE org_tree AS (",
    "    -- 基础情况：从指定员工开始",
    "    SELECT id, name, manager_id, 0 AS depth,",
    "           ARRAY[id] AS path",
    "    FROM employees WHERE id = 42",
    "    UNION ALL",
    "    -- 递归步骤：向上查找管理层",
    "    SELECT e.id, e.name, e.manager_id, t.depth + 1,",
    "           t.path || e.id",
    "    FROM employees e",
    "    INNER JOIN org_tree t ON e.id = t.manager_id",
    "    WHERE NOT e.id = ANY(t.path)  -- 防止循环引用",
    ")",
    "SELECT id, name, depth FROM org_tree ORDER BY depth;",
    "",
    "-- 多步 CTE 数据管道",
    "WITH",
    "active_orders AS (",
    "    SELECT * FROM orders WHERE status NOT IN ('cancelled', 'refunded')",
    "),",
    "order_totals AS (",
    "    SELECT user_id, SUM(total) AS lifetime_value",
    "    FROM active_orders",
    "    GROUP BY user_id",
    "),",
    "vip_users AS (",
    "    SELECT user_id FROM order_totals WHERE lifetime_value > 10000",
    ")",
    "UPDATE users SET tier = 'VIP'",
    "WHERE id IN (SELECT user_id FROM vip_users);",
]: cp(line)
bl()

h2("22.3  JSONB 操作完整参考")
for line in [
    "-- 创建带 JSONB 列的表",
    "CREATE TABLE events (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    ts TIMESTAMPTZ DEFAULT now(),",
    "    data JSONB NOT NULL",
    ");",
    "",
    "-- 插入 JSONB 数据",
    "INSERT INTO events (data) VALUES",
    '    (\'{"type": "click", "page": "/home", "user_id": 123, "tags": ["web", "mobile"]}\'),',
    '    (\'{"type": "purchase", "amount": 99.99, "items": [{"id": 1, "qty": 2}]}\');',
    "",
    "-- JSONB 查询操作符",
    "SELECT data->>'type' AS event_type          -- 提取字符串值",
    "FROM events WHERE data->>'type' = 'click';",
    "",
    "SELECT data->'items'->0->>'id' AS first_item  -- 嵌套访问",
    "FROM events WHERE data->>'type' = 'purchase';",
    "",
    "SELECT * FROM events WHERE data @> '{\"type\": \"click\"}';  -- 包含查询（利用 GIN 索引）",
    "",
    "-- JSONB 修改",
    "UPDATE events SET data = data || '{\"processed\": true}'  -- 合并（覆盖同名键）",
    "WHERE id = 1;",
    "",
    "UPDATE events SET data = data - 'tags'  -- 删除键",
    "WHERE id = 1;",
    "",
    "-- GIN 索引加速 @> 查询",
    "CREATE INDEX idx_events_data ON events USING GIN(data);",
]: cp(line)
bl()

h2("22.4  全文搜索深度使用")
for line in [
    "-- 中文全文搜索（需要分词器支持）",
    "-- FalconDB v1.2 内置 simple 和 english 分词器",
    "",
    "-- 创建包含 tsvector 列的表（预处理，性能更好）",
    "CREATE TABLE articles (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    title TEXT NOT NULL,",
    "    body TEXT NOT NULL,",
    "    search_vector tsvector GENERATED ALWAYS AS (",
    "        to_tsvector('english', title) ||",
    "        setweight(to_tsvector('english', body), 'B')  -- 标题权重更高",
    "    ) STORED",
    ");",
    "",
    "-- GIN 索引加速全文搜索",
    "CREATE INDEX idx_articles_fts ON articles USING GIN(search_vector);",
    "",
    "-- 全文搜索查询",
    "SELECT",
    "    id, title,",
    "    ts_rank(search_vector, query) AS relevance,",
    "    ts_headline('english', body, query, 'MaxWords=50, MinWords=20') AS excerpt",
    "FROM articles,",
    "     to_tsquery('english', 'database & performance & !slow') AS query",
    "WHERE search_vector @@ query",
    "ORDER BY relevance DESC",
    "LIMIT 20;",
    "",
    "-- 前缀搜索（用 tsquery 的 :* 操作符）",
    "SELECT title FROM articles",
    "WHERE search_vector @@ to_tsquery('english', 'perfor:*');  -- 匹配 performance, performs 等",
]: cp(line)
bl()

h2("22.5  高级数据类型使用")
h3("22.5.1  数组类型")
for line in [
    "-- 数组列定义和插入",
    "CREATE TABLE products (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    name TEXT,",
    "    tags TEXT[],",
    "    prices DECIMAL(10,2)[]",
    ");",
    "",
    "INSERT INTO products (name, tags, prices)",
    "VALUES ('Laptop', ARRAY['electronics', 'computer', 'portable'], ARRAY[999.99, 1199.99]);",
    "",
    "-- 数组查询",
    "SELECT * FROM products WHERE 'electronics' = ANY(tags);   -- 包含元素",
    "SELECT * FROM products WHERE tags @> ARRAY['electronics', 'computer'];  -- 包含子数组",
    "SELECT * FROM products WHERE tags && ARRAY['electronics', 'gaming'];    -- 有交集",
    "",
    "-- 数组函数",
    "SELECT array_length(tags, 1) AS tag_count FROM products;  -- 数组长度",
    "SELECT unnest(tags) AS tag FROM products WHERE id = 1;    -- 展开数组为行",
    "SELECT array_agg(name) AS names FROM products;            -- 聚合为数组",
    "SELECT string_agg(name, ', ') FROM products;              -- 聚合为字符串",
]: cp(line)
bl()

h3("22.5.2  枚举类型")
for line in [
    "-- 创建枚举类型",
    "CREATE TYPE order_status AS ENUM (",
    "    'pending', 'confirmed', 'processing',",
    "    'shipped', 'delivered', 'cancelled', 'refunded'",
    ");",
    "",
    "CREATE TABLE orders (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    status order_status NOT NULL DEFAULT 'pending'",
    ");",
    "",
    "-- 枚举值比较（保持定义时的顺序）",
    "SELECT * FROM orders WHERE status > 'pending';  -- 比 pending 更后面的状态",
    "SELECT * FROM orders ORDER BY status;           -- 按枚举定义顺序排序",
    "",
    "-- 添加新枚举值",
    "ALTER TYPE order_status ADD VALUE 'on_hold' BEFORE 'cancelled';",
    "",
    "-- 查看枚举类型的所有值",
    "SELECT enum_range(NULL::order_status);",
]: cp(line)
bl()

h2("22.6  存储过程和函数完整参考")
for line in [
    "-- 带异常处理的转账函数",
    "CREATE OR REPLACE FUNCTION transfer(",
    "    p_from_id BIGINT,",
    "    p_to_id   BIGINT,",
    "    p_amount  DECIMAL(12,2),",
    "    OUT result TEXT",
    ") LANGUAGE plpgsql AS $$",
    "DECLARE",
    "    v_from_balance DECIMAL(12,2);",
    "    v_to_exists    BOOLEAN;",
    "BEGIN",
    "    -- 验证目标账户存在",
    "    SELECT EXISTS(SELECT 1 FROM accounts WHERE id = p_to_id) INTO v_to_exists;",
    "    IF NOT v_to_exists THEN",
    "        result := 'ERROR: destination account not found';",
    "        RETURN;",
    "    END IF;",
    "",
    "    -- 加锁并检查余额",
    "    SELECT balance INTO v_from_balance",
    "    FROM accounts WHERE id = p_from_id FOR UPDATE;",
    "",
    "    IF v_from_balance < p_amount THEN",
    "        result := format('ERROR: insufficient balance %.2f < %.2f',",
    "                         v_from_balance, p_amount);",
    "        RETURN;",
    "    END IF;",
    "",
    "    -- 执行转账",
    "    UPDATE accounts SET balance = balance - p_amount,",
    "                        updated_at = now()",
    "    WHERE id = p_from_id;",
    "",
    "    UPDATE accounts SET balance = balance + p_amount,",
    "                        updated_at = now()",
    "    WHERE id = p_to_id;",
    "",
    "    -- 记录转账日志",
    "    INSERT INTO transfer_log (from_id, to_id, amount, ts)",
    "    VALUES (p_from_id, p_to_id, p_amount, now());",
    "",
    "    result := format('SUCCESS: transferred %.2f from %s to %s',",
    "                     p_amount, p_from_id, p_to_id);",
    "EXCEPTION",
    "    WHEN OTHERS THEN",
    "        result := 'ERROR: ' || SQLERRM;",
    "        RAISE WARNING 'Transfer failed: %', SQLERRM;",
    "END;",
    "$$;",
    "",
    "-- 调用函数",
    "SELECT transfer(1001, 1002, 500.00);",
]: cp(line)
bl()

h2("22.7  触发器使用指南（v1.2+）")
for line in [
    "-- 自动更新 updated_at 时间戳",
    "CREATE OR REPLACE FUNCTION set_updated_at() RETURNS trigger AS $$",
    "BEGIN",
    "    NEW.updated_at = now();",
    "    RETURN NEW;",
    "END;",
    "$$ LANGUAGE plpgsql;",
    "",
    "-- 为多个表注册触发器",
    "CREATE TRIGGER trg_users_updated_at",
    "    BEFORE UPDATE ON users",
    "    FOR EACH ROW EXECUTE FUNCTION set_updated_at();",
    "",
    "CREATE TRIGGER trg_orders_updated_at",
    "    BEFORE UPDATE ON orders",
    "    FOR EACH ROW EXECUTE FUNCTION set_updated_at();",
    "",
    "-- 审计触发器（记录数据变更历史）",
    "CREATE TABLE audit_trail (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    table_name TEXT, operation TEXT,",
    "    old_data JSONB, new_data JSONB,",
    "    changed_by TEXT DEFAULT current_user,",
    "    changed_at TIMESTAMPTZ DEFAULT now()",
    ");",
    "",
    "CREATE OR REPLACE FUNCTION audit_changes() RETURNS trigger AS $$",
    "BEGIN",
    "    INSERT INTO audit_trail (table_name, operation, old_data, new_data)",
    "    VALUES (",
    "        TG_TABLE_NAME,",
    "        TG_OP,",
    "        CASE WHEN TG_OP = 'INSERT' THEN NULL ELSE to_jsonb(OLD) END,",
    "        CASE WHEN TG_OP = 'DELETE' THEN NULL ELSE to_jsonb(NEW) END",
    "    );",
    "    RETURN COALESCE(NEW, OLD);",
    "END;",
    "$$ LANGUAGE plpgsql;",
    "",
    "CREATE TRIGGER trg_orders_audit",
    "    AFTER INSERT OR UPDATE OR DELETE ON orders",
    "    FOR EACH ROW EXECUTE FUNCTION audit_changes();",
]: cp(line)
pb()

# ═══════════════════════════════════════════════════
# 第二十三章  事务管理深度指南
# ═══════════════════════════════════════════════════
h1("第二十三章  事务管理深度指南")

h2("23.1  MVCC 实现机制")
np("FalconDB 使用多版本并发控制（Multi-Version Concurrency Control，MVCC）实现快照隔离。MVCC 的核心思想是每次写入不修改原数据，而是创建一个新版本，让不同事务看到数据在不同时间点的快照：")
for line in [
    "# MVCC 版本链示意（VersionChain2 数据结构）",
    "",
    "# 行 id=1 的版本链（链头为最新版本）：",
    "row_key=1 →",
    "  Version { ts=100, data='Alice', deleted=false }  ← 最新版本（ts=100）",
    "  Version { ts=50,  data='Alicia', deleted=false } ← 历史版本（ts=50）",
    "  Version { ts=10,  data='Ali',   deleted=false }  ← 原始版本（ts=10）",
    "",
    "# 读取规则：",
    "# 事务 T（start_ts=75）读取 id=1：",
    "#   扫描版本链，找到 ts <= 75 的最新版本 → ts=50 的版本 'Alicia'",
    "# 事务 T'（start_ts=110）读取 id=1：",
    "#   扫描版本链，找到 ts <= 110 的最新版本 → ts=100 的版本 'Alice'",
    "",
    "# OCC 写入验证（提交时）：",
    "# 检查写集中的所有行，确认 start_ts 之后没有其他事务提交了对同一行的写入",
    "# 如果有冲突 → ROLLBACK（SQLSTATE 40001），客户端需重试",
]: cp(line)
bl()

h2("23.2  隔离级别深度对比")
tbl(["隔离级别","FalconDB 实现","脏读","不可重复读","幻读","写偏斜","性能"],
[["Read Committed（默认）","每条 SQL 新建快照","不可能","可能","可能","可能","最高"],
 ["Repeatable Read（快照隔离）","事务开始时固定快照","不可能","不可能","不可能*","可能","高"],
 ["Serializable","快照隔离 + OCC 写冲突检测","不可能","不可能","不可能","不可能","较低（有重试）"]])
np("注：*FalconDB 的快照隔离（SI）使用 MVCC 实现，在标准 SI 定义下幻读不可能发生，因为快照固定了可见数据集。写偏斜（Write Skew）是 SI 下唯一可能的异常，需要 Serializable 级别来防止。")
bl()

h2("23.3  LocalTxn 快速路径详解")
np("LocalTxn 是 FalconDB 最重要的性能优化——单分片事务完全绕过两阶段提交（2PC），大幅降低提交延迟：")
for line in [
    "# LocalTxn 提交流程（单分片，OCC 路径）",
    "",
    "BEGIN",
    "  ↓",
    "操作 1: READ  key=user:1   → 记录读集：{user:1: observed_ts=50}",
    "操作 2: WRITE key=user:1   → 记录写集：{user:1: new_data}",
    "操作 3: WRITE key=order:42 → 记录写集：{order:42: new_data}",
    "  ↓",
    "COMMIT 快速路径（满足：所有操作在同一分片）：",
    "  1. get_commit_ts() → commit_ts = 101（单调递增）",
    "  2. OCC 验证：",
    "     对写集中每个 key，检查 start_ts 之后是否有其他提交",
    "     如果任意 key 有冲突 → ROLLBACK（40001）",
    "  3. 写入 WAL（原子地写入所有变更）",
    "  4. WAL fsync / group_commit（等待刷盘）",
    "  5. 更新 VersionChain（写入新版本，ts=101）",
    "  6. 返回客户端：COMMIT OK",
    "  总延迟：WAL fsync 时间（通常 < 1ms SSD）",
    "",
    "# GlobalTxn 提交流程（跨分片，2PC 路径）：",
    "  1. Prepare 阶段：发送 PREPARE 到所有参与分片",
    "  2. 等待所有分片 PREPARE ACK",
    "  3. Commit 阶段：发送 COMMIT 到所有参与分片",
    "  总延迟：2 × 网络 RTT + 各分片 WAL fsync",
]: cp(line)
bl()

h2("23.4  保存点（SAVEPOINT）使用指南")
for line in [
    "-- 保存点允许部分回滚，无需回滚整个事务",
    "BEGIN;",
    "",
    "-- 阶段 1：更新用户信息",
    "UPDATE users SET last_login = now() WHERE id = 1;",
    "SAVEPOINT sp_user_updated;  -- 保存点 1",
    "",
    "-- 阶段 2：尝试扣减库存",
    "UPDATE inventory SET stock = stock - 1 WHERE product_id = 101;",
    "",
    "-- 如果库存不足，回滚到保存点（保留用户更新）",
    "-- ROLLBACK TO SAVEPOINT sp_user_updated;",
    "",
    "-- 阶段 3：创建订单",
    "INSERT INTO orders (user_id, product_id, status) VALUES (1, 101, 'pending');",
    "SAVEPOINT sp_order_created;  -- 保存点 2",
    "",
    "-- 阶段 4：发送通知（可能失败）",
    "INSERT INTO notifications (user_id, message) VALUES (1, 'Order created');",
    "-- 如果通知失败，只回滚通知，保留订单",
    "-- ROLLBACK TO SAVEPOINT sp_order_created;",
    "",
    "-- 释放不再需要的保存点（释放内存）",
    "RELEASE SAVEPOINT sp_user_updated;",
    "",
    "COMMIT;  -- 提交所有未回滚的操作",
]: cp(line)
bl()

h2("23.5  并发控制最佳实践")
h3("23.5.1  SELECT FOR UPDATE 悲观锁")
np("虽然 FalconDB 主要使用 OCC（乐观并发控制），但支持 SELECT FOR UPDATE 进行悲观锁定：")
for line in [
    "-- 悲观锁定场景：资源预订（如机票、酒店房间）",
    "BEGIN;",
    "",
    "-- 加锁读取座位信息（其他事务无法同时修改）",
    "SELECT id, status FROM seats",
    "WHERE flight_id = 'CA123' AND status = 'available'",
    "LIMIT 1",
    "FOR UPDATE;  -- 锁定选中的行，直到事务提交或回滚",
    "",
    "-- 基于读取结果进行预订",
    "UPDATE seats SET status = 'reserved', user_id = $user_id",
    "WHERE id = $seat_id AND status = 'available';  -- 双重检查",
    "",
    "INSERT INTO bookings (user_id, seat_id, flight_id, ts) VALUES (...);",
    "",
    "COMMIT;",
    "",
    "-- SELECT FOR UPDATE SKIP LOCKED：跳过已锁定的行（任务队列模式）",
    "BEGIN;",
    "SELECT id, payload FROM job_queue",
    "WHERE status = 'pending'",
    "ORDER BY created_at",
    "LIMIT 1",
    "FOR UPDATE SKIP LOCKED;   -- 跳过其他 worker 已锁定的任务",
    "",
    "UPDATE job_queue SET status = 'processing' WHERE id = $job_id;",
    "COMMIT;",
]: cp(line)
bl()

h3("23.5.2  减少事务持有时间")
np("长事务是 OCC 冲突、GC 阻塞、连接占用的主要来源。减少事务持有时间的建议：")
bp("将所有数据验证逻辑移到事务外执行（只在事务内做写操作）")
bp("避免在事务内调用外部服务（HTTP API、文件操作等），可能长时间阻塞")
bp("使用 INSERT ... ON CONFLICT 替代 SELECT + INSERT 两步操作")
bp("批量操作分成适当大小的子批次提交（每次 1000-5000 行）")
bp("配置合理的 statement_timeout_ms 和 idle_in_transaction_session_timeout，自动清理遗忘的长事务")
pb()

# ═══════════════════════════════════════════════════
# 第二十四章  存储引擎深度指南
# ═══════════════════════════════════════════════════
h1("第二十四章  存储引擎深度指南")

h2("24.1  Rowstore 内存引擎内部架构")
np("Rowstore 是 FalconDB 的核心存储引擎，专为低延迟 OLTP 而设计。其内部架构由三个主要组件构成：")
h3("24.1.1  DashMap 哈希表层")
np("每张 Rowstore 表对应一个 DashMap<Key, VersionChain2>，其中 Key 是主键的字节序列化，VersionChain2 是 MVCC 版本链。DashMap 是一个基于分段锁的并发哈希表，支持高并发的无竞争读取。")
bp("时间复杂度：点查 O(1)，范围扫描 O(k)（k 为范围内行数）")
bp("哈希函数：FxHasher（固定速度哈希），比标准 SipHash 快约 3x")
bp("并发控制：64 个分片锁，不同键的操作可以完全并行")
bl()

h3("24.1.2  VersionChain2 MVCC 版本链")
np("VersionChain2 是 FalconDB v1.2 新设计的无锁 MVCC 版本链，针对高并发读写场景优化：")
bp("链结构：单向链表，链头为最新版本，链尾为最老版本")
bp("无锁读取：使用 Arc<AtomicPtr> 实现无锁版本链遍历，读操作完全无锁")
bp("写入追加：新版本追加到链头，不修改已有版本（写入安全，不影响并发读取）")
bp("惰性 GC：标记删除（设置 deleted=true），GC 线程异步清理")
for line in [
    "# 版本链内存布局（每个版本约 64 字节）",
    "struct Version {",
    "    commit_ts: u64,          // 8 字节：版本时间戳",
    "    data: Arc<[u8]>,         // 16 字节：行数据（共享引用）",
    "    deleted: bool,           // 1 字节：逻辑删除标记",
    "    next: AtomicPtr<Version> // 8 字节：指向旧版本的指针",
    "}",
    "",
    "# 版本链 GC 触发条件：",
    "# 1. GcRunner 定时扫描（interval_ms）",
    "# 2. 发现版本链长度 >= min_chain_length",
    "# 3. 该链最旧可能有效版本 > GC 安全点（safepoint_ts）",
    "# 4. 满足以上条件：从链尾移除过期版本，释放 Arc 引用（内存自动回收）",
]: cp(line)
bl()

h3("24.1.3  二级索引（ART-Tree）")
np("Rowstore 的二级索引使用自适应基数树（Adaptive Radix Tree，ART-Tree），相比 B-Tree 在内存使用和查询性能上更优：")
bp("节点类型自适应：根据子节点数量自动选择 Node4/Node16/Node48/Node256，最大化内存效率")
bp("前缀压缩：共享前缀存储在节点中，减少内存冗余")
bp("顺序扫描优化：ART-Tree 保持键的字典序，支持高效范围查询")
bp("并发安全：通过 RwLock 保护索引结构，允许并发读取")
bl()

h2("24.2  WAL 内部结构")
np("WAL（Write-Ahead Log）是 FalconDB 持久化和崩溃恢复的基础。了解 WAL 内部结构有助于更好地调优和排查问题：")
for line in [
    "# WAL 文件结构",
    "",
    "data_dir/",
    "└── wal/",
    "    ├── 00000001.wal  ← 第一个 WAL 段（64MB）",
    "    ├── 00000002.wal  ← 第二个 WAL 段",
    "    └── 00000003.wal  ← 当前活跃段（正在写入）",
    "",
    "# 每个 WAL 记录的结构：",
    "struct WalRecord {",
    "    lsn: u64,          // 8 字节：日志序号（全局单调递增）",
    "    record_type: u8,   // 1 字节：类型（Write/Commit/Rollback/Checkpoint）",
    "    txn_id: u64,       // 8 字节：关联的事务 ID",
    "    shard_id: u32,     // 4 字节：目标分片",
    "    payload_len: u32,  // 4 字节：数据长度",
    "    payload: [u8],     // 变长：序列化的操作数据",
    "    crc32: u32,        // 4 字节：CRC32 校验（检测磁盘位翻转）",
    "}",
    "",
    "# WAL 恢复（崩溃后重启）：",
    "# 1. 找到最后一个完整的 Checkpoint 记录",
    "# 2. 从 Checkpoint 位置开始重放后续所有 WAL 记录",
    "# 3. 对于 Commit 记录：应用对应事务的所有变更",
    "# 4. 对于孤立的 Write 记录（无对应 Commit）：忽略（回滚）",
    "# 5. 恢复完成，数据库回到崩溃前的一致状态",
]: cp(line)
bl()

h2("24.3  USTM 分层内存管理（v1.1+）")
np("USTM（Universal Storage Tiering Manager）是 FalconDB v1.1 引入的分层内存管理系统，将数据按访问频率分为三层：")
tbl(["层次","数据特征","存储位置","访问延迟","说明"],
[["Hot（热层）","最近 5 分钟内访问的行","主内存（DRAM）","纳秒至微秒","Rowstore 默认层，最快访问"],
 ["Warm（温层）","5分钟-1小时未访问","扩展内存（大内存页/PMEM）","微秒至毫秒","自动从 Hot 层降级"],
 ["Cold（冷层）","超过 1 小时未访问","磁盘（SSD/NVMe）","毫秒级","自动从 Warm 层降级，类似 LSM"]])
bp("数据在层间自动迁移：访问 Cold 层数据时自动提升到 Hot 层（缓存预热）")
bp("内存压力触发降级：当 Hot 层内存超过阈值，自动将最久未访问的数据降级到 Warm 层")
bp("对应用透明：无论数据在哪一层，SQL 查询语义完全一致")
bl()

h2("24.4  LSM 引擎调优")
np("LSM-Tree（Log-Structured Merge-tree）引擎适合大数据量写入密集场景。调优参数：")
for line in [
    "[storage.lsm]",
    "sync_writes = false       # 异步写入（更高吞吐，极端情况可能丢最后一条记录）",
    "block_cache_mb = 512      # 块缓存大小（热数据缓存在内存中）",
    "write_buffer_mb = 64      # 写缓冲区大小（越大批量压缩越多，但内存占用越高）",
    "max_write_buffers = 3     # 最大写缓冲数（超过触发 flush）",
    "level0_file_num_compaction_trigger = 4   # Level-0 文件数达到 4 时触发压缩",
    "target_file_size_mb = 64  # 目标 SSTable 文件大小",
    "compression = 'lz4'       # 压缩算法（snappy/lz4/zstd，lz4 速度最快）",
]: cp(line)
np("LSM 写放大（Write Amplification）：LSM 的数据可能被多次压缩合并。write_buffer_mb 越大，写放大越小（因为更多数据在一次压缩中合并），但内存占用越高。")
bl()

h2("24.5  存储引擎选择决策树")
for line in [
    "# 如何选择存储引擎？",
    "",
    "数据总量是否超过可用内存的 80%？",
    "  YES → 使用磁盘引擎",
    "    写入吞吐 > 10K 行/秒？",
    "      YES → 使用 RocksDB（ENGINE=rocksdb）",
    "      NO  → 使用 LSM（ENGINE=lsm）",
    "    需要零 C/C++ 依赖？",
    "      YES → 使用 redb（ENGINE=redb，纯 Rust）",
    "  NO  → 使用内存引擎（ENGINE=rowstore，默认）",
    "    需要最低延迟（< 1ms P99 提交）？",
    "      YES → Rowstore + local-fsync WAL",
    "    需要最高持久性（RPO≈0）？",
    "      YES → Rowstore + sync-replica WAL",
    "",
    "# 混合策略示例（同一数据库中使用不同引擎）",
    "CREATE TABLE hot_sessions (...) ENGINE=rowstore;      -- 热数据",
    "CREATE TABLE cold_archives (...) ENGINE=rocksdb;      -- 冷数据",
    "CREATE TABLE search_index (...) ENGINE=lsm;           -- 写密集",
]: cp(line)
pb()

print("Part 5 完成，保存中...")
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
