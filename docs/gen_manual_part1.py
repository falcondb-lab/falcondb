# -*- coding: utf-8 -*-
"""FalconDB 操作手册生成器 Part 1 - 文档框架和前半部分"""
import sys, subprocess

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

OUTPUT = r"c:\falcondb\falcondb-main\docs\FalconDB_操作手册.docx"
doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin = sec.bottom_margin = Cm(2.5)

# 修改内置样式
from docx.shared import Pt
for lvl, sz, color in [(1,18,(0x1F,0x38,0x64)),(2,14,(0x2E,0x74,0xB5)),(3,12,(0x26,0x26,0x26))]:
    try:
        st = doc.styles[f'Heading {lvl}']
        st.font.size = Pt(sz)
        st.font.color.rgb = RGBColor(*color)
        st.font.bold = True
    except: pass

def h1(t): return doc.add_heading(t, 1)
def h2(t): return doc.add_heading(t, 2)
def h3(t): return doc.add_heading(t, 3)

def np(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(11)
    return p

def bp(t):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t); r.font.size = Pt(11)
    return p

def cp(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(t)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
    pPr.append(shd)
    return p

def bl(): doc.add_paragraph()

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        r = t.rows[ri+1]
        for ci,v in enumerate(row):
            r.cells[ci].text = str(v)
            if r.cells[ci].paragraphs[0].runs:
                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return t

def pb(): doc.add_page_break()

# ══ 封面 ══
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FalconDB\n数据库操作手册\n")
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("版本 v1.2  |  兼容 PostgreSQL · 分布式 · 内存优先 · 确定性事务语义\n\n文档编制日期：2026年3月\n\n")
r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x59,0x59,0x59)
pb()

# ══ 目录 ══
h1("目  录")
toc_items = [
    "第一章   产品概述与核心特性",
    "第二章   安装与部署",
    "第三章   配置参考",
    "第四章   存储引擎",
    "第五章   数据库连接与客户端",
    "第六章   SQL 语言参考",
    "第七章   事务管理",
    "第八章   数据定义语言（DDL）",
    "第九章   数据操作语言（DML）",
    "第十章   索引管理",
    "第十一章  安全管理",
    "第十二章  集群与复制",
    "第十三章  故障转移与高可用",
    "第十四章  备份与恢复",
    "第十五章  监控与可观测性",
    "第十六章  AI 优化器与 AIOps 引擎",
    "第十七章  性能调优指南",
    "第十八章  常见问题与故障排除",
    "第十九章  生产运维检查清单",
]
for item in toc_items:
    bp(item)
pb()

# ══ 第一章 ══
h1("第一章  产品概述与核心特性")
h2("1.1  产品简介")
np("FalconDB 是一款兼容 PostgreSQL 协议的分布式内存优先 OLTP 数据库，具备确定性事务语义（Deterministic Commit Guarantee，DCG）。它融合了低延迟、高吞吐、强一致性和可证明的 ACID 保证，专为对延迟和可靠性要求极高的在线交易处理场景而设计。")
np("FalconDB 已与 PostgreSQL、VoltDB、SingleStore 进行基准测试对比。在百万行批量插入场景下，吞吐量约为 PostgreSQL 的 1.84 倍（340K 行/秒 vs 185K 行/秒）；单分片事务延迟通过绕过 2PC 显著降低；整体读写延迟达到微秒级别。")
bl()

h2("1.2  核心特性")
feats = [
    "低延迟：单分片快速路径提交完全绕过两阶段提交（2PC），显著降低事务延迟，读写延迟达微秒级",
    "确定性提交保证（DCG）：返回「已提交」后，事务在任何单节点崩溃、故障转移和恢复中均不丢失数据",
    "可证明的一致性：快照隔离下的 MVCC/OCC，持续集成自动化验证 ACID 属性（4,350+ 测试覆盖 18 个 crate）",
    "PostgreSQL 协议兼容：支持 psql、pgbench、JDBC、Go pgx、Python psycopg2 等所有标准 PostgreSQL 客户端",
    "多存储引擎：Rowstore（内存）、LSM（磁盘）、RocksDB、redb，同一数据库中不同表可独立选择引擎",
    "完整 SQL 覆盖：DDL/DML/子查询/CTE/窗口函数/全文搜索/存储过程（PL/pgSQL）/触发器/物化视图/枚举类型",
    "分布式高可用：Raft 共识复制、WAL gRPC 流式复制、自动故障转移、分片自动再均衡（47 项混沌测试验证）",
    "内置 AI 优化器：在线 SGD 代价模型、15 维特征向量、自适应查询计划选择、实时训练进度可观测",
    "AIOps 自治运维引擎：慢查询自适应检测、3σ 异常告警、自动索引建议、工作负载 p95/p99 分析",
    "企业级安全：SCRAM-SHA-256、TLS/mTLS、RBAC 权限继承、行级安全（RLS）、AES-256-GCM 透明数据加密、不可篡改审计日志",
    "丰富的可观测性：50+ SHOW 命令、Prometheus 指标端点、慢查询日志、pg_stat_statements 兼容视图",
    "变更数据捕获（CDC）：内置 CDC 引擎，支持逻辑复制协议，可对接下游数据管道",
]
for f in feats: bp(f)
bl()

h2("1.3  确定性提交保证（DCG）")
np("确定性提交保证（Deterministic Commit Guarantee，DCG）是 FalconDB 的核心工程特性。其核心承诺为：一旦 FalconDB 向客户端返回「已提交」（COMMIT OK）应答，该事务将在以下所有场景中存活，数据零丢失：")
for d in [
    "任何单节点崩溃，包括 kill -9 强制终止进程",
    "任何计划内或计划外的故障转移（主从切换）操作",
    "任何 WAL 崩溃恢复过程（包括断电、操作系统崩溃）",
]:
    bp(d)
np("DCG 在 LocalWalSync 提交策略下为默认行为，无需任何额外配置。生产环境必须启用 WAL（wal_enabled = true）并选择 local-fsync 或 sync-replica 持久化策略。DCG 不是配置选项，而是系统工程保证。")
bl()

h2("1.4  支持平台与最低要求")
tbl(
    ["平台", "构建支持", "测试支持", "备注"],
    [
        ["Linux (x86_64, Ubuntu 22.04+)", "✓", "✓", "主要 CI 目标，生产推荐"],
        ["Windows (x86_64, MSVC)", "✓", "✓", "CI 目标，Windows 服务支持"],
        ["macOS (x86_64 / aarch64)", "✓", "✓", "社区测试，开发环境"],
    ]
)
bl()
np("软件要求：Rust 工具链 1.75 或更高版本（通过 rustup 安装）；C/C++ 编译器（Linux/macOS 使用 gcc/clang，Windows 使用 MSVC Visual Studio Build Tools）")
np("硬件要求：生产环境建议最低 4 GB 内存、4 核 CPU；WAL 目录建议使用 NVMe SSD 以最小化刷盘延迟；Rowstore 内存引擎的数据容量受可用内存直接限制")
pb()

# ══ 第二章 ══
h1("第二章  安装与部署")
h2("2.1  从源码构建")
np("FalconDB 使用 Rust 语言编写，通过 Cargo 构建系统管理依赖和编译流程。支持多种存储引擎的特性编译开关（feature flags）。")
bl()

h3("2.1.1  标准构建（Rowstore + LSM 引擎）")
for line in [
    "# 调试版本（开发测试，编译快，运行慢）",
    "cargo build --workspace",
    "",
    "# 发布版本（生产推荐，全优化）",
    "cargo build --release --workspace",
    "",
    "# 运行完整测试套件（4,350+ 测试，跨 18 个 crate）",
    "cargo test --workspace",
    "",
    "# 代码静态检查",
    "cargo clippy --workspace",
]:
    cp(line)
bl()

h3("2.1.2  特定存储引擎构建")
for line in [
    "# RocksDB 引擎（Windows 需要 LLVM + MSVC）",
    "cargo build --release -p falcon_server --features rocksdb",
    "",
    "# redb 引擎（纯 Rust，零 C/C++ 依赖）",
    "cargo build --release -p falcon_server --features redb",
    "",
    "# Standard 版本（RocksDB 为默认存储引擎，面向磁盘优先场景）",
    "cargo build --release --features standard",
    "",
    "# 启用透明数据加密（TDE）",
    "cargo build --release -p falcon_server --features encryption_tde",
    "",
    "# 启用时间点恢复（PITR）",
    "cargo build --release -p falcon_server --features pitr",
]:
    cp(line)
bl()

h2("2.2  Windows 平台部署")
h3("2.2.1  一键开发环境配置")
for line in [
    r"# 以管理员身份运行 PowerShell",
    r".\scripts\setup_windows.ps1",
]:
    cp(line)
bl()

h3("2.2.2  Windows 服务管理")
for line in [
    "# 安装为 Windows 服务（需管理员权限）",
    "falcon service install",
    "",
    "# 启动服务",
    "falcon service start",
    "",
    "# 停止服务（优雅关闭，等待活跃事务完成）",
    "falcon service stop",
    "",
    "# 重启服务",
    "falcon service restart",
    "",
    "# 查看服务状态（显示 PID、端口、路径）",
    "falcon service status",
    "",
    "# 卸载服务",
    "falcon service uninstall",
]:
    cp(line)
bl()
np("Windows 服务自动故障恢复策略：")
tbl(
    ["故障次数", "处理动作", "等待延迟", "说明"],
    [
        ["第 1 次", "自动重启", "5 秒", "首次故障快速恢复"],
        ["第 2 次", "自动重启", "10 秒", "第二次适当延迟"],
        ["第 3 次及以后", "自动重启", "30 秒", "防止频繁重启消耗资源"],
    ]
)
bl()
np("注意：故障计数器在连续稳定运行 24 小时后自动重置。")
bl()

h2("2.3  单节点启动")
h3("2.3.1  无 WAL 内存模式（开发测试，最快，重启数据丢失）")
for line in [
    "cargo run -p falcon_server -- --no-wal",
    "",
    "# 或使用已构建的二进制文件",
    "./target/release/falcon --no-wal",
]:
    cp(line)
bl()

h3("2.3.2  带 WAL 持久化模式（生产推荐）")
for line in [
    "# 指定数据目录，自动启用 WAL",
    "cargo run -p falcon_server -- --data-dir ./falcon_data",
    "",
    "# 使用配置文件启动（推荐生产使用）",
    "cargo run -p falcon_server -- --config falcon.toml",
    "",
    "# 生成完整默认配置文件",
    "cargo run -p falcon_server -- --print-default-config > falcon.toml",
]:
    cp(line)
bl()

h2("2.4  主从复制部署")
h3("2.4.1  通过配置文件（推荐生产使用）")
for line in [
    "# 主节点——接受写入，通过 gRPC 向从节点流式传输 WAL",
    "cargo run -p falcon_server -- -c examples/primary.toml",
    "",
    "# 从节点——接收 WAL，仅服务只读查询",
    "cargo run -p falcon_server -- -c examples/replica.toml",
]:
    cp(line)
bl()

h3("2.4.2  通过命令行参数（快速测试）")
for line in [
    "# 主节点（PG 端口 5433，gRPC 流式复制端口 50051）",
    "cargo run -p falcon_server -- --role primary \\",
    "  --pg-addr 0.0.0.0:5433 --grpc-addr 0.0.0.0:50051 --data-dir ./node1",
    "",
    "# 从节点（PG 端口 5434，连接主节点 gRPC）",
    "cargo run -p falcon_server -- --role replica \\",
    "  --pg-addr 0.0.0.0:5434 --primary-endpoint http://127.0.0.1:50051 --data-dir ./node2",
]:
    cp(line)
bl()

h2("2.5  连接验证")
for line in [
    "# 使用 psql 连接（默认端口 5433）",
    "psql -h 127.0.0.1 -p 5433 -U falcon",
    "",
    "# 连接后验证服务状态",
    "SELECT version();",
    "SHOW falcon.status;",
    "SHOW falcon.connections;",
]:
    cp(line)
bl()

h2("2.6  Linux/macOS 开发环境")
for line in [
    "# 安装 Rust 工具链",
    "curl --proto '=https' --tlsv1.2 -sSf https://sh.rustup.rs | sh",
    "",
    "# 安装 PostgreSQL 客户端（Debian/Ubuntu）",
    "sudo apt install postgresql-client",
    "",
    "# 安装 PostgreSQL 客户端（macOS）",
    "brew install libpq",
    "",
    "# 构建",
    "cargo build --workspace",
]:
    cp(line)
bl()

h2("2.7  目录结构（Windows）")
for line in [
    "C:\\Program Files\\FalconDB\\",
    "  └── bin\\falcon.exe                     ← 程序文件（只读，MSI 管理）",
    "",
    "C:\\ProgramData\\FalconDB\\",
    "  ├── conf\\falcon.toml                   ← 配置文件",
    "  ├── data\\                              ← 数据文件 + WAL 段文件",
    "  │   ├── falcon.wal                     ← 写前日志",
    "  │   └── ...                            ← 表数据文件",
    "  ├── logs\\",
    "  │   ├── falcon.log                     ← 当前日志",
    "  │   └── falcon.log.YYYY-MM-DD          ← 历史日志（按日轮转）",
    "  ├── certs\\                             ← TLS 证书（可选）",
    "  │   ├── server.crt",
    "  │   ├── server.key",
    "  │   └── ca.crt",
    "  └── backup\\                            ← 升级自动备份",
]:
    cp(line)
pb()

# ══ 第三章 ══
h1("第三章  配置参考")
h2("3.1  配置文件概述")
np("FalconDB 使用 TOML 格式配置文件（默认名称 falcon.toml）。配置文件分为多个段落（Section），每个段落控制一个子系统的行为。配置变更支持版本管理，可分阶段发布（Canary → Rolling → Applied）并支持一键回滚。")
np("生成完整默认配置文件：cargo run -p falcon_server -- --print-default-config > falcon.toml")
bl()

h2("3.2  [server] 段——服务器基础配置")
for line in [
    "[server]",
    'pg_listen_addr = "0.0.0.0:5433"       # PostgreSQL 协议监听地址和端口',
    'admin_listen_addr = "0.0.0.0:8080"    # 管理 HTTP 端口（健康检查、Prometheus 指标）',
    "node_id = 1                             # 节点唯一标识符（集群内必须唯一）",
    "max_connections = 1024                  # 最大并发连接数（超出时新连接被拒绝）",
    "statement_timeout_ms = 0               # 单条 SQL 超时（毫秒，0 = 不限制）",
    "idle_timeout_ms = 0                    # 空闲连接超时（毫秒，0 = 不限制）",
    "shutdown_drain_timeout_secs = 30       # 优雅关闭等待活跃事务的最长时间（秒）",
    "slow_txn_threshold_us = 0             # 慢事务告警阈值（微秒，0 = 不记录）",
]:
    cp(line)
bl()

h2("3.3  [storage] 段——存储配置")
for line in [
    "[storage]",
    "memory_limit_bytes = 0          # Rowstore 内存预算（字节，0 = 不限制）",
    "wal_enabled = true              # 启用 WAL 持久化（生产环境必须为 true）",
    'data_dir = "./falcon_data"      # 数据目录路径（存放 WAL 段文件和表数据）',
    "lsm_sync_writes = false        # LSM 引擎是否同步写入（true 更安全但更慢）",
    'default_engine = "rowstore"    # 默认存储引擎（rowstore/lsm/rocksdb/redb）',
]:
    cp(line)
bl()

h2("3.4  [wal] 段——WAL 写前日志配置")
for line in [
    "[wal]",
    "group_commit = true                    # 启用组提交（多个事务共享一次 fsync，提升吞吐）",
    "flush_interval_us = 1000              # 组提交刷新间隔（微秒，降低可增加延迟）",
    'sync_mode = "fdatasync"               # 同步模式：fdatasync/fsync/none（none 仅用于测试）',
    'durability_policy = "local-fsync"     # 持久化策略：local-fsync 或 sync-replica',
    "segment_size_bytes = 67108864         # WAL 段文件大小（默认 64MB）",
    'wal_mode = "win_async"                # Windows 专用异步 WAL 模式',
]:
    cp(line)
bl()
np("持久化策略详解：")
bp("local-fsync（默认，推荐大多数场景）：主节点 WAL 写入本地磁盘并 fsync 后立即响应客户端。单节点崩溃下零数据丢失（DCG 保证），多节点同时崩溃时可能丢失未复制部分。RPO > 0，提交延迟最低。")
bp("sync-replica（RPO ≈ 0）：主节点等待至少一个从节点 WAL 确认后才响应客户端。即使主节点在响应后立即崩溃，数据已在从节点持久化。代价是提交延迟增加网络 RTT。")
bl()

h2("3.5  [gc] 段——MVCC 垃圾回收")
for line in [
    "[gc]",
    "enabled = true           # 启用自动 MVCC 垃圾回收（生产推荐 true）",
    "interval_ms = 5000       # GC 扫描间隔（毫秒）",
    "batch_size = 1000        # 每次 GC 扫描的最大版本链数量",
    "min_chain_length = 2     # 版本链长度达到此值才触发 GC（减少无效扫描）",
]:
    cp(line)
bl()

h2("3.6  [logging] 段——日志配置")
for line in [
    "[logging]",
    'level = "info"           # 日志级别：trace/debug/info/warn/error',
    "stderr = true            # 是否将日志输出到标准错误（前台模式有效）",
    'file = ""                # 日志文件路径（空字符串 = 不写文件）',
    'format = "text"          # 日志格式：text（人可读）或 json（机器可解析）',
    'rotation = "daily"       # 日志轮转周期：daily/hourly/never',
    "max_size_mb = 100        # 单个日志文件最大大小（MB）",
    "max_files = 7            # 保留的历史日志文件数量",
    "slow_query_ms = 100      # 慢查询阈值（毫秒，超过此值记录到慢查询日志）",
    'slow_query_log = ""      # 慢查询日志文件路径（空 = 不单独记录）',
]:
    cp(line)
bl()

h2("3.7  [tde] 段——透明数据加密（TDE）")
for line in [
    "[tde]",
    "enabled = false                      # 启用 AES-256-GCM 静态数据加密",
    'key_env_var = "FALCON_TDE_KEY"       # 主密钥环境变量名（禁止在配置文件中硬编码密钥）',
    "encrypt_wal = true                   # 加密 WAL 段文件",
    "encrypt_data = true                  # 加密数据文件",
]:
    cp(line)
bl()

h2("3.8  [pitr] 段——时间点恢复")
for line in [
    "[pitr]",
    "enabled = false          # 启用 WAL 归档和 PITR",
    'archive_dir = ""         # WAL 归档目录路径',
    "retention_hours = 168    # 归档 WAL 保留时间（小时，默认 7 天）",
]:
    cp(line)
bl()

h2("3.9  [raft] 段——Raft 共识（三节点集群示例）")
for line in [
    "[replication]",
    'role = "raft_member"     # 启用 Raft 模式',
    "",
    "[raft]",
    "node_id = 1              # 本节点唯一 ID（> 0，集群内唯一）",
    "node_ids = [1, 2, 3]    # 集群所有节点 ID 列表",
    'raft_listen_addr = "0.0.0.0:50052"  # Raft gRPC 传输监听地址',
    'peers = [                # 对等节点地址（本节点除外）',
    '    "2=http://node2:50052",',
    '    "3=http://node3:50052",',
    "]",
    "heartbeat_interval_ms = 50         # 心跳间隔",
    "election_timeout_min_ms = 150      # 选举超时下限",
    "election_timeout_max_ms = 300      # 选举超时上限",
    "failover_poll_interval_ms = 200    # 故障转移轮询间隔",
]:
    cp(line)
bl()

h2("3.10  [cdc] 段——变更数据捕获")
for line in [
    "[cdc]",
    "enabled = true           # 启用 CDC 事件流",
    "buffer_size = 100000     # CDC 内存缓冲区事件容量",
]:
    cp(line)
bl()

h2("3.11  配置版本管理")
np("FalconDB 的配置变更经过版本化管理，支持受控发布和回滚。每条配置记录包含：版本号（单调递增 u64）、校验和（djb2 散列）、发布状态。")
tbl(
    ["发布阶段", "说明", "操作"],
    [
        ["Stage（暂存）", "变更已记录，尚未生效", "falconctl config stage <key> <value>"],
        ["Canary（金丝雀）", "应用到金丝雀节点，观察 15 分钟", "falconctl config rollout <key> canary"],
        ["RollingOut（滚动）", "逐步应用到所有节点", "falconctl config rollout <key> rolling"],
        ["Applied（已应用）", "全集群生效", "自动完成"],
        ["RolledBack（已回滚）", "回滚到上一版本", "falconctl config rollback <key>"],
    ]
)
pb()

# ══ 第四章 ══
h1("第四章  存储引擎")
h2("4.1  存储引擎对比")
tbl(
    ["引擎", "存储方式", "持久化方式", "数据上限", "最适场景"],
    [
        ["Rowstore（默认）", "内存 MVCC 版本链", "WAL 崩溃安全恢复", "受可用内存限制", "低延迟 OLTP、热数据、会话存储"],
        ["LSM", "磁盘 LSM-Tree", "完整磁盘持久化", "受磁盘空间限制", "大数据集、日志、冷数据"],
        ["RocksDB", "磁盘 RocksDB", "完整磁盘持久化", "受磁盘空间限制", "生产级磁盘 OLTP、成熟压缩策略"],
        ["redb", "磁盘纯 Rust KV", "完整磁盘持久化", "受磁盘空间限制", "零 C 依赖嵌入式场景"],
    ]
)
bl()

h2("4.2  Rowstore 内存引擎（默认）")
np("Rowstore 是 FalconDB 的默认存储引擎，将所有数据存储在内存中，通过 DashMap 实现高并发键值访问，并使用 VersionChain2（无锁 MVCC 版本链）实现多版本并发控制。")
for line in [
    "-- 以下两种写法等价，均使用 Rowstore 引擎",
    "CREATE TABLE sessions (id INT PRIMARY KEY, token TEXT);",
    "CREATE TABLE sessions (id INT PRIMARY KEY, token TEXT) ENGINE=rowstore;",
]:
    cp(line)
bl()
np("Rowstore 关键特性：")
bp("读写延迟最低，单分片快速路径提交完全绕过 2PC，适合高频点查和单行更新")
bp("WAL 启用时崩溃后自动恢复；不启用 WAL 时服务重启后数据丢失（仅用于测试）")
bp("MVCC GC 后台线程（GcRunner）自动回收过期版本，防止内存泄漏")
bp("内存背压控制：[memory] 段可配置软限制（触发 GC）和硬限制（拒绝写入）")
bp("锁自由（Lock-Free）MVCC 版本链，高并发读写无互斥锁竞争")
bl()

h2("4.3  LSM 磁盘引擎")
np("LSM 引擎使用日志结构合并树（LSM-Tree）实现磁盘持久化，适合数据量超出内存容量的场景。LSM feature 默认包含在构建中，无需额外编译选项。")
for line in [
    "CREATE TABLE events (id BIGSERIAL PRIMARY KEY, payload JSONB) ENGINE=lsm;",
    "CREATE TABLE logs (ts TIMESTAMPTZ, level TEXT, msg TEXT) ENGINE=lsm;",
]:
    cp(line)
bl()

h2("4.4  RocksDB 引擎")
np("RocksDB 引擎基于 Facebook RocksDB 库，提供生产级的磁盘 OLTP 性能，具有成熟的 LSM 层级压缩策略和高吞吐批量写入能力。Standard 版以 RocksDB 为默认引擎。")
for line in [
    "# 编译时需要启用 rocksdb feature（Windows 需要 LLVM + MSVC）",
    "cargo build --release -p falcon_server --features rocksdb",
    "",
    "-- 建表时指定 RocksDB 引擎",
    "CREATE TABLE orders (id BIGINT PRIMARY KEY, status TEXT, total DECIMAL(12,2)) ENGINE=rocksdb;",
]:
    cp(line)
bl()

h2("4.5  redb 引擎")
np("redb 是一个纯 Rust 实现的嵌入式键值存储，零 C/C++ 依赖，适合需要简化部署链的场景。")
for line in [
    "cargo build --release -p falcon_server --features redb",
    "CREATE TABLE config (key TEXT PRIMARY KEY, value TEXT) ENGINE=redb;",
]:
    cp(line)
bl()

h2("4.6  混合引擎策略")
np("同一数据库中的不同表可以独立选择存储引擎，充分利用各引擎的优势：")
for line in [
    "-- 热数据：内存引擎，毫秒级延迟",
    "CREATE TABLE active_sessions (id UUID PRIMARY KEY, user_id BIGINT, expires_at TIMESTAMPTZ);",
    "",
    "-- 交易记录：RocksDB，高吞吐持久化",
    "CREATE TABLE transactions (id BIGSERIAL PRIMARY KEY, amount DECIMAL(12,2), created_at TIMESTAMPTZ) ENGINE=rocksdb;",
    "",
    "-- 审计日志：LSM，大容量追加写入优化",
    "CREATE TABLE audit_trail (id BIGSERIAL PRIMARY KEY, event JSONB, ts TIMESTAMPTZ) ENGINE=lsm;",
]:
    cp(line)
bl()
np("注意：跨不同引擎的表之间可以执行 JOIN 操作，但涉及多个存储引擎的事务会走全局慢速路径（XA-2PC）。对延迟敏感的业务路径建议单一引擎内完成事务。")
pb()

# ══ 第五章 ══
h1("第五章  数据库连接与客户端")
h2("5.1  PostgreSQL 协议兼容性详表")
tbl(
    ["功能", "状态", "备注"],
    [
        ["简单查询协议（Simple Query）", "✓", "单语句和多语句批处理"],
        ["扩展查询（Parse/Bind/Execute）", "✓", "预处理语句 + Portal 游标"],
        ["认证：Trust（无认证）", "✓", "开发环境，不建议生产"],
        ["认证：MD5 散列", "✓", "PG auth type 5"],
        ["认证：SCRAM-SHA-256", "✓", "PG 10+ 兼容，生产推荐"],
        ["认证：明文密码", "✓", "PG auth type 3"],
        ["TLS/SSL 连接加密", "✓", "SSLRequest 握手后升级"],
        ["COPY IN / COPY OUT", "✓", "Text 和 CSV 格式"],
        ["LISTEN / NOTIFY", "✓", "内存广播，跨会话通知"],
        ["取消请求（Cancel Request）", "✓", "AtomicBool 轮询，50ms 延迟"],
        ["逻辑复制协议", "✓", "IDENTIFY_SYSTEM、START_REPLICATION 等"],
        ["psql 12+", "✓", "完整测试兼容"],
        ["pgbench（init + run）", "✓", "内建脚本可用"],
        ["JDBC (pgjdbc 42.7+)", "✓", "完整测试兼容"],
        ["Go pgx / lib/pq", "✓", "通过 PG 线协议直连"],
        ["Python psycopg2 / psycopg3", "✓", "通过 PG 线协议直连"],
        ["Node.js node-postgres (pg)", "✓", "通过 PG 线协议直连"],
    ]
)
bl()

h2("5.2  psql 命令行连接")
for line in [
    "# 基本连接（默认端口 5433）",
    "psql -h 127.0.0.1 -p 5433 -U falcon",
    "",
    "# 指定数据库",
    "psql -h 127.0.0.1 -p 5433 -U falcon -d mydb",
    "",
    "# 使用连接字符串",
    "psql 'postgresql://falcon:password@127.0.0.1:5433/mydb'",
    "",
    "# SSL 加密连接",
    "psql 'postgresql://falcon@127.0.0.1:5433/mydb?sslmode=require'",
    "",
    "# 执行 SQL 文件",
    "psql -h 127.0.0.1 -p 5433 -U falcon -d mydb -f schema.sql",
]:
    cp(line)
bl()

h2("5.3  各语言客户端连接示例")
h3("5.3.1  Java JDBC（pgjdbc 42.7+）")
for line in [
    'String url = "jdbc:postgresql://127.0.0.1:5433/mydb";',
    "Properties props = new Properties();",
    'props.setProperty("user", "falcon");',
    'props.setProperty("password", "your_password");',
    'props.setProperty("ssl", "true");           // 启用 SSL',
    "Connection conn = DriverManager.getConnection(url, props);",
    "",
    "// 使用连接池（HikariCP 示例）",
    "HikariConfig config = new HikariConfig();",
    'config.setJdbcUrl("jdbc:postgresql://127.0.0.1:5433/mydb");',
    'config.setUsername("falcon");',
    'config.setPassword("your_password");',
    "HikariDataSource ds = new HikariDataSource(config);",
]:
    cp(line)
bl()

h3("5.3.2  Python psycopg2")
for line in [
    "import psycopg2",
    "conn = psycopg2.connect(",
    "    host='127.0.0.1',",
    "    port=5433,",
    "    dbname='mydb',",
    "    user='falcon',",
    "    password='your_password',",
    "    sslmode='require'    # 生产环境建议启用 SSL",
    ")",
    "cur = conn.cursor()",
    "cur.execute('SELECT version()')",
    "print(cur.fetchone())",
]:
    cp(line)
bl()

h3("5.3.3  Go pgx")
for line in [
    'import "github.com/jackc/pgx/v5"',
    "",
    'conn, err := pgx.Connect(ctx, "postgres://falcon:password@127.0.0.1:5433/mydb")',
    "if err != nil { log.Fatal(err) }",
    "defer conn.Close(ctx)",
    "",
    "// 连接池",
    'pool, err := pgxpool.New(ctx, "postgres://falcon:password@127.0.0.1:5433/mydb")',
]:
    cp(line)
bl()

h3("5.3.4  Node.js (node-postgres)")
for line in [
    "const { Pool } = require('pg');",
    "",
    "const pool = new Pool({",
    "  host: '127.0.0.1',",
    "  port: 5433,",
    "  database: 'mydb',",
    "  user: 'falcon',",
    "  password: 'your_password',",
    "  ssl: { rejectUnauthorized: false }  // 或配置 ca 证书",
    "});",
    "",
    "const { rows } = await pool.query('SELECT version()');",
    "console.log(rows[0]);",
]:
    cp(line)
bl()

h3("5.3.5  .NET (Npgsql)")
for line in [
    'var connStr = "Host=127.0.0.1;Port=5433;Database=mydb;Username=falcon;Password=your_password;";',
    "await using var conn = new NpgsqlConnection(connStr);",
    "await conn.OpenAsync();",
]:
    cp(line)
bl()

h2("5.4  健康检查 HTTP 接口")
np("通过管理端口（默认 8080）进行服务健康检查，适合负载均衡器探活和监控系统集成：")
for line in [
    "GET http://127.0.0.1:8080/health",
    "  → 200 OK        服务健康，可接受请求",
    "  → 503 Service Unavailable   服务异常或正在恢复",
    "",
    "GET http://127.0.0.1:8080/ready",
    "  → 200 OK        服务就绪，数据可访问",
    "  → 503 Service Unavailable   正在初始化或恢复",
    "",
    "GET http://127.0.0.1:8080/status",
    "  → JSON 对象，包含版本、角色、连接数、WAL LSN 等详情",
    "",
    "GET http://127.0.0.1:9090/metrics",
    "  → Prometheus 文本格式指标（所有 falcon_* 指标）",
]:
    cp(line)
pb()

# ══ 第六章 ══
h1("第六章  SQL 语言参考")
h2("6.1  数据类型完整表")
tbl(
    ["类型名称", "PG OID", "存储大小", "说明"],
    [
        ["BOOLEAN", "16", "1 字节", "布尔值（true/false/null）"],
        ["SMALLINT / INT2", "21", "2 字节", "小整数（-32768 到 32767）"],
        ["INT / INTEGER / INT4", "23", "4 字节", "32 位有符号整数"],
        ["BIGINT / INT8", "20", "8 字节", "64 位有符号整数"],
        ["REAL / FLOAT4", "700", "4 字节", "单精度 IEEE 754 浮点数"],
        ["FLOAT8 / DOUBLE PRECISION", "701", "8 字节", "双精度 IEEE 754 浮点数"],
        ["NUMERIC(p,s) / DECIMAL(p,s)", "1700", "可变", "任意精度定点小数"],
        ["TEXT", "25", "可变", "无长度限制变长文本"],
        ["VARCHAR(n)", "1043", "可变", "最长 n 字符的变长文本"],
        ["CHAR(n)", "1042", "n 字节", "定长文本，不足则空格填充"],
        ["BYTEA", "17", "可变", "二进制字节数组"],
        ["DATE", "1082", "4 字节", "日期（年月日）"],
        ["TIME", "1083", "8 字节", "时间（时分秒微秒）"],
        ["TIMESTAMP", "1114", "8 字节", "时间戳（不含时区）"],
        ["TIMESTAMPTZ", "1184", "8 字节", "时间戳（含时区，存储 UTC）"],
        ["INTERVAL", "1186", "16 字节", "时间间隔"],
        ["UUID", "2950", "16 字节", "128 位通用唯一标识符"],
        ["JSONB", "3802", "可变", "JSON 二进制存储，支持索引"],
        ["ARRAY（一维）", "varies", "可变", "BOOLEAN[]、INT[]、TEXT[] 等"],
        ["SERIAL", "—", "4 字节", "自增 INT（等价于 INT + SEQUENCE）"],
        ["BIGSERIAL", "—", "8 字节", "自增 BIGINT（等价于 BIGINT + SEQUENCE）"],
        ["ENUM（用户自定义）", "—", "可变", "CREATE TYPE ... AS ENUM 定义"],
        ["tsvector / tsquery", "3614/3615", "可变", "全文搜索专用类型"],
    ]
)
bl()

h2("6.2  运算符一览")
tbl(
    ["类别", "运算符", "示例"],
    [
        ["比较", "=、!=、<>、<、>、<=、>=", "age >= 18"],
        ["逻辑", "AND、OR、NOT", "is_active AND age > 18"],
        ["算术", "+、-、*、/、%（取余）", "price * 0.9"],
        ["字符串拼接", "||", "first_name || ' ' || last_name"],
        ["模式匹配", "LIKE、ILIKE（忽略大小写）", "name LIKE 'A%'"],
        ["正则匹配", "~（区分大小写）、~*（不区分）", "email ~* '@gmail\\.com$'"],
        ["NULL 判断", "IS NULL、IS NOT NULL", "phone IS NOT NULL"],
        ["NULL 合并", "COALESCE(a, b, c)", "COALESCE(nickname, username)"],
        ["条件替换", "NULLIF(a, b)", "NULLIF(count, 0)"],
        ["范围", "BETWEEN a AND b", "age BETWEEN 18 AND 65"],
        ["集合成员", "IN (list)、NOT IN (list)", "status IN ('active', 'pending')"],
        ["数组包含", "@>（包含）、<@（被包含）", "tags @> ARRAY['rust']"],
        ["数组交集", "&&（有交集）", "skills && ARRAY['sql']"],
        ["JSONB 字段", "->（返回 JSON）、->>（返回文本）", "data->>'name'"],
        ["JSONB 路径", "#>（路径返回 JSON）、#>>（路径返回文本）", "data#>>'{address,city}'"],
        ["全文搜索", "@@（文档匹配查询）", "to_tsvector('english',body) @@ to_tsquery('rust')"],
    ]
)
bl()

h2("6.3  内置函数分类参考")
h3("6.3.1  字符串处理函数（选录）")
tbl(
    ["函数", "说明", "示例"],
    [
        ["length(s)", "字符串长度", "length('hello') → 5"],
        ["upper(s) / lower(s)", "大/小写转换", "upper('hello') → 'HELLO'"],
        ["trim(s) / ltrim(s) / rtrim(s)", "去除首尾/左/右空格", "trim('  hi  ') → 'hi'"],
        ["substring(s, from, len)", "截取子字符串", "substring('hello',2,3) → 'ell'"],
        ["replace(s, from, to)", "替换子字符串", "replace('hello','l','r') → 'herro'"],
        ["concat(s1, s2, ...)", "拼接多个字符串", "concat('a','b','c') → 'abc'"],
        ["split_part(s, delim, n)", "按分隔符分割取第n部分", "split_part('a,b,c',',',2) → 'b'"],
        ["starts_with(s, prefix)", "判断是否以前缀开头", "starts_with('hello','he') → true"],
        ["md5(s)", "计算 MD5 散列（十六进制）", "md5('data') → '...'"],
        ["encode(bytes, 'hex')", "字节数组转十六进制字符串", "encode(E'\\xDEAD','hex') → 'dead'"],
    ]
)
bl()

h3("6.3.2  数学函数（选录）")
tbl(
    ["函数", "说明"],
    [
        ["abs(x)", "绝对值"],
        ["ceil(x) / floor(x)", "向上/向下取整"],
        ["round(x, d)", "四舍五入到 d 位小数"],
        ["sqrt(x)", "平方根"],
        ["power(x, y)", "x 的 y 次幂"],
        ["random()", "返回 [0,1) 的随机浮点数"],
        ["greatest(a, b, ...)", "返回最大值"],
        ["least(a, b, ...)", "返回最小值"],
        ["log(x) / ln(x) / exp(x)", "以 10 为底对数 / 自然对数 / e 的幂"],
    ]
)
bl()

h3("6.3.3  日期时间函数（选录）")
tbl(
    ["函数", "说明", "示例"],
    [
        ["now()", "当前时间戳（含时区）", "now() → 2026-03-12 14:23:00+08"],
        ["current_date", "当前日期", "current_date → 2026-03-12"],
        ["extract(field FROM ts)", "提取日期时间字段", "extract(year FROM now()) → 2026"],
        ["date_trunc(unit, ts)", "截断到指定精度", "date_trunc('month', now())"],
        ["age(ts1, ts2)", "两个时间戳的差值", "age(birthday) → '28 years 3 months'"],
        ["to_char(ts, format)", "格式化为字符串", "to_char(now(),'YYYY-MM-DD')"],
        ["to_timestamp(epoch)", "Unix 时间戳转时间", "to_timestamp(1710000000)"],
        ["now() + INTERVAL '7 days'", "时间运算", "到期时间计算"],
    ]
)
bl()

h3("6.3.4  聚合函数")
tbl(
    ["函数", "说明"],
    [
        ["count(*) / count(col)", "计数（* 含 NULL，col 不含 NULL）"],
        ["sum(col)", "求和"],
        ["avg(col)", "平均值"],
        ["min(col) / max(col)", "最小值 / 最大值"],
        ["string_agg(col, delim)", "字符串聚合（PostgreSQL 风格 GROUP_CONCAT）"],
        ["array_agg(col)", "将值聚合为数组"],
        ["bool_and(col) / bool_or(col)", "布尔值聚合"],
    ]
)
bl()

h2("6.4  全文搜索")
np("FalconDB 提供与 PostgreSQL 完全兼容的全文搜索功能，支持 tsvector/tsquery 类型和 14 个 FTS 函数。")
for line in [
    "-- 将文本转换为 tsvector（分词并标准化）",
    "SELECT to_tsvector('english', 'FalconDB is a fast database');",
    "-- 输出：'database':5 'falcondb':1 'fast':4",
    "",
    "-- 全文搜索查询",
    "SELECT title FROM articles",
    "WHERE to_tsvector('english', body) @@ to_tsquery('english', 'database & performance');",
    "",
    "-- 带相关性排序",
    "SELECT title,",
    "       ts_rank(to_tsvector('english', body), query) AS rank",
    "FROM articles, to_tsquery('english', 'performance') query",
    "WHERE to_tsvector('english', body) @@ query",
    "ORDER BY rank DESC LIMIT 10;",
    "",
    "-- 高亮匹配词",
    "SELECT ts_headline('english', body, to_tsquery('performance'))",
    "FROM articles WHERE id = 1;",
    "",
    "-- 创建 GIN 全文索引（提升搜索性能）",
    "CREATE INDEX idx_articles_fts ON articles USING GIN(to_tsvector('english', body));",
]:
    cp(line)
pb()

# ══ 第七章 ══
h1("第七章  事务管理")
h2("7.1  事务模型概述")
np("FalconDB 实现了完整的 MVCC（多版本并发控制）和 OCC（乐观并发控制）事务模型。每个事务获得一个唯一的事务 ID（TxnId）和读时间戳，通过版本链实现非阻塞读操作，确保读操作不会阻塞写，写操作不会阻塞读。")
bl()

h2("7.2  事务类型详解")
tbl(
    ["类型", "触发条件", "提交协议", "特点", "典型延迟"],
    [
        ["LocalTxn（本地快速路径）", "仅涉及单个分片", "OCC（乐观并发控制）", "无锁检测，零 2PC 开销", "微秒级"],
        ["GlobalTxn（全局慢速路径）", "跨多个分片或节点", "XA-2PC（两阶段提交）", "跨分片原子性保证", "毫秒级"],
    ]
)
bl()
np("FalconDB 在提交时硬性校验：LocalTxn 必须恰好只涉及 1 个分片且使用快速路径。如果应用程序写入了多个分片但错误地使用了 LocalTxn，系统会返回 InternalError。")
bl()

h2("7.3  隔离级别")
tbl(
    ["隔离级别", "SQL 设置", "FalconDB 实现", "脏读", "不可重复读", "幻读"],
    [
        ["Read Committed（读已提交）", "SET TRANSACTION ISOLATION LEVEL READ COMMITTED", "每个语句获取新快照", "不可能", "可能", "可能"],
        ["Snapshot Isolation（快照隔离）", "SET TRANSACTION ISOLATION LEVEL REPEATABLE READ", "事务开始时固定快照", "不可能", "不可能", "不可能*"],
        ["Serializable（可串行化）", "SERIALIZABLE（实验性）", "基于 SI + OCC 验证", "不可能", "不可能", "不可能"],
    ]
)
bl()
np("* FalconDB 的快照隔离通过 MVCC 版本可见性实现，防止不可重复读和幻读。Write-Write 冲突通过 OCC 在提交时检测，发生冲突时返回 SQLSTATE 40001，应用程序应重试事务。")
bl()

h2("7.4  事务操作语法")
for line in [
    "-- 开始事务（两种等价写法）",
    "BEGIN;",
    "START TRANSACTION;",
    "",
    "-- 开始只读事务（不会产生写冲突）",
    "BEGIN READ ONLY;",
    "",
    "-- 开始特定隔离级别的事务",
    "BEGIN ISOLATION LEVEL REPEATABLE READ;",
    "",
    "-- 提交事务",
    "COMMIT;",
    "",
    "-- 回滚事务（撤销所有未提交修改）",
    "ROLLBACK;",
    "",
    "-- 保存点（事务内的检查点）",
    "SAVEPOINT sp_before_update;",
    "UPDATE accounts SET balance = balance - 100 WHERE id = 1;",
    "-- 如果出错，回滚到保存点而不是回滚整个事务",
    "ROLLBACK TO SAVEPOINT sp_before_update;",
    "-- 明确释放保存点（释放其占用的资源）",
    "RELEASE SAVEPOINT sp_before_update;",
]:
    cp(line)
bl()

h2("7.5  OCC 冲突处理最佳实践")
np("FalconDB 使用乐观并发控制（OCC），在提交时而非操作时检测写-写冲突。这意味着应用程序需要处理 SQLSTATE 40001（serialization_failure）并重试事务。")
for line in [
    "-- Python 重试示例",
    "import psycopg2, time",
    "",
    "def execute_with_retry(conn, fn, max_retries=3):",
    "    for attempt in range(max_retries):",
    "        try:",
    "            with conn.cursor() as cur:",
    "                fn(cur)",
    "                conn.commit()",
    "                return",
    "        except psycopg2.errors.SerializationFailure:",
    "            conn.rollback()",
    "            if attempt == max_retries - 1: raise",
    "            time.sleep(0.01 * (2 ** attempt))  # 指数退避",
]:
    cp(line)
bl()

h2("7.6  MVCC 垃圾回收监控")
for line in [
    "-- 查看 GC 统计（回收版本数、字节数、安全点）",
    "SHOW falcon.gc_stats;",
    "",
    "-- 查看哪些事务阻塞了 GC 推进",
    "SHOW falcon.gc_safepoint;",
    "",
    "-- 查看活跃事务列表（包含 TxnId 和开始时间戳）",
    "SHOW falcon.txn;",
    "",
    "-- 查看事务统计汇总",
    "SHOW falcon.txn_stats;",
    "",
    "-- 查看近期事务历史",
    "SHOW falcon.txn_history;",
]:
    cp(line)
bl()
np("GC 安全点计算公式：gc_safepoint = min(最小活跃事务时间戳, 从节点已应用时间戳) − 1。如果存在长时间运行的事务，GC 安全点无法推进，可能导致内存或磁盘中积累大量历史版本。建议监控活跃事务的持续时间，设置合理的事务超时（statement_timeout_ms）。")
pb()

print(f"Part 1 完成，保存中...")
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
