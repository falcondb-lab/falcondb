# -*- coding: utf-8 -*-
"""Part 6 - 容量规划/迁移/CDC/多租户/操作案例/最佳实践"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn; from docx.oxml import OxmlElement

OUTPUT = r"c:\falcondb\falcondb-main\docs\FalconDB_操作手册.docx"
doc = Document(OUTPUT)

def h1(t): return doc.add_heading(t, 1)
def h2(t): return doc.add_heading(t, 2)
def h3(t): return doc.add_heading(t, 3)
def np(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(11)
    return p
def bp(t):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t); r.font.size = Pt(11); return p
def cp(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(t); r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0F0F0')
    pPr.append(shd); return p
def bl(): doc.add_paragraph()
def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        if c.paragraphs[0].runs:
            run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = c._tc; tcPr = tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        r = t.rows[ri+1]
        for ci,v in enumerate(row):
            r.cells[ci].text = str(v)
            if r.cells[ci].paragraphs[0].runs:
                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return t
def pb(): doc.add_page_break()

# ═══════════════════════════════════════════════
# 第二十五章  容量规划指南
# ═══════════════════════════════════════════════
h1("第二十五章  容量规划指南")

h2("25.1  容量规划方法论")
np("容量规划是在系统资源耗尽之前提前采取行动，避免突发性故障。FalconDB 的容量规划关注四个维度：内存（Rowstore 数据）、磁盘（WAL + 磁盘引擎 + 备份）、CPU（查询执行 + 复制）、网络（复制流量 + 客户端流量）。")
np("容量规划的基本原则：当前使用率不超过 70%（留 30% 缓冲），增长趋势分析基于至少 30 天的历史数据，提前 3 个月做出扩容决策（采购/审批/部署周期）。")
bl()

h2("25.2  内存容量规划（Rowstore）")
np("Rowstore 的内存需求计算公式：")
for line in [
    "# 内存需求估算（Rowstore 引擎）",
    "",
    "# 单行内存占用 = 数据大小 + 版本链开销 + 索引开销",
    "# 数据大小 = 每列平均字节数之和",
    "# 版本链开销 = 约 64 字节/版本 × 平均活跃版本数（通常 1-3）",
    "# 索引开销 = 每个索引约 32 字节/行（ART-Tree 节点）",
    "",
    "# 示例计算：",
    "# 表 orders：100M 行，平均行大小 200 字节，2 个二级索引",
    "# 数据大小 = 100M × 200 = 20 GB",
    "# 版本链（平均 1.5 版本）= 100M × 1.5 × 64 = 9.6 GB",
    "# 索引（2 个）= 100M × 2 × 32 = 6.4 GB",
    "# 合计 = 36 GB",
    "",
    "# 加上 OS 预留（20%）= 36 GB / 0.8 = 45 GB 物理内存",
    "",
    "# 每日增长率估算：",
    "# 今日数据量 = X GB，30 天前数据量 = Y GB",
    "# 日增长率 = (X - Y) / 30",
    "# 距离 80% 容量上限还有 N 天 = (limit * 0.8 - X) / 日增长率",
]: cp(line)
bl()

h2("25.3  磁盘容量规划")
h3("25.3.1  WAL 磁盘占用估算")
for line in [
    "# WAL 磁盘占用 = 活跃 WAL 段 + 归档 WAL 段（PITR）",
    "",
    "# 活跃 WAL 段（约 2-5 个 segment_size 大小，保留最近可回收）",
    "# 通常 = 5 × 64MB = 320 MB",
    "",
    "# 归档 WAL 段（PITR 保留期 × WAL 生成速率）",
    "# WAL 生成速率估算：",
    "#   每次写入操作（INSERT/UPDATE/DELETE）平均写入 200-500 字节 WAL",
    "#   1K 写/秒 → 约 500 KB/s WAL = 43 GB/天",
    "#   10K 写/秒 → 约 5 MB/s WAL = 432 GB/天",
    "",
    "# 保留 7 天的 PITR：",
    "#   1K 写/秒 场景：7 × 43 GB = 301 GB",
    "#   10K 写/秒 场景：7 × 432 GB = 3 TB",
    "",
    "# 建议：WAL 归档目录使用独立磁盘，容量 = 保留天数 × 日 WAL 量 × 1.2（20% 缓冲）",
]: cp(line)
bl()

h3("25.3.2  备份存储规划")
tbl(["备份策略","每次备份大小","保留周期","总存储需求"],
[["每日全量备份","= 数据库总大小","30 天","30 × 全量大小"],
 ["每日全量 + WAL 归档","全量 + 7天WAL","30 天全量 + 7 天WAL","30×全量 + 7×日WAL"],
 ["每周全量 + 每日增量","全量 + 每日WAL差量","4 周全量 + 每日增量","4×全量 + 28×增量（通常小很多）"]])
bl()

h2("25.4  CPU 和网络容量规划")
np("CPU 容量规划关注两个维度：查询处理和 WAL 复制。")
tbl(["指标","低负载","中等负载","高负载","扩容触发点"],
[["CPU 使用率（查询）","< 30%","30%-60%","60%-80%","持续 > 70%"],
 ["CPU 使用率（复制线程）","< 10%","10%-20%","20%-30%","持续 > 25%"],
 ["网络带宽（复制）","< 10 MB/s","10-100 MB/s","100 MB/s-1 GB/s","接近网络带宽 70%"],
 ["TPS","< 1K/s","1K-10K/s","10K-100K/s","接近设计上限 80%"],
 ["P99 提交延迟","< 1 ms","1-10 ms","10-50 ms","持续 > 20 ms"]])
bl()

h2("25.5  扩容决策矩阵")
tbl(["瓶颈类型","症状","短期缓解措施","长期解决方案"],
[["内存不足（Rowstore）","内存使用率 > 80%，GC 频繁","增大 GC 频率，清理过期数据","增加物理内存，或迁移冷数据到磁盘引擎"],
 ["WAL 磁盘不足","WAL 目录 > 80%","调短 retention_hours，手动清理旧归档","扩容 WAL 磁盘，或使用 S3 存储归档"],
 ["CPU 瓶颈（查询）","CPU > 70%，查询排队","优化热点 SQL，创建缺失索引","增加 CPU 核心数，或增加只读副本分担查询"],
 ["连接数饱和","connections_active 接近 max","提高 max_connections","部署连接池（PgBouncer），或增加节点"],
 ["复制带宽不足","复制延迟持续增大","降低 WAL 生成速率（批量写入）","升级网络带宽，或改为异步复制"],
 ["单分片热点","OCC 冲突率 > 10%","重新设计数据模型减少热点","启用分片自动再均衡，或手动分片分裂"]])
pb()

# ═══════════════════════════════════════════════
# 第二十六章  数据迁移指南
# ═══════════════════════════════════════════════
h1("第二十六章  数据迁移指南")

h2("26.1  从 PostgreSQL 迁移到 FalconDB")
np("FalconDB 兼容 PostgreSQL 协议，大多数 PostgreSQL 应用可以直接连接 FalconDB 而无需修改代码。迁移流程分为三个阶段：评估、迁移和验证。")

h3("26.1.1  兼容性评估")
for line in [
    "-- 在 PostgreSQL 中运行，识别不兼容的功能",
    "",
    "-- 检查是否使用了不支持的功能",
    "SELECT relname AS table_name",
    "FROM pg_class c",
    "JOIN pg_trigger t ON c.oid = t.tgrelid",
    "WHERE c.relkind = 'r'",
    "GROUP BY relname;  -- 查看有触发器的表（v1.2+ 支持触发器）",
    "",
    "-- 检查是否使用了不支持的扩展",
    "SELECT extname FROM pg_extension WHERE extname NOT IN ('plpgsql');",
    "",
    "-- 检查是否使用了 FULL OUTER JOIN",
    "-- 在应用代码中搜索：grep -ri 'FULL OUTER JOIN' ./sql/",
    "",
    "-- 检查是否使用了 LATERAL JOIN",
    "-- grep -ri 'LATERAL' ./sql/",
]: cp(line)
bl()

h3("26.1.2  数据导出（PostgreSQL 侧）")
for line in [
    "# 方法一：pg_dump（推荐，完整 DDL + 数据）",
    "pg_dump -h pg-host -U postgres -d mydb \\",
    "    --no-owner --no-acl \\",
    "    --format=plain \\",
    "    -f mydb_dump.sql",
    "",
    "# 方法二：COPY 导出为 CSV（大表分表导出）",
    "psql -h pg-host -U postgres -d mydb -c \\",
    '    "COPY orders TO STDOUT WITH (FORMAT CSV, HEADER)" > orders.csv',
    "",
    "psql -h pg-host -U postgres -d mydb -c \\",
    '    "COPY users TO STDOUT WITH (FORMAT CSV, HEADER)" > users.csv',
    "",
    "# 方法三：逻辑复制（零停机迁移，高级）",
    "# 在 PostgreSQL 端创建发布",
    "CREATE PUBLICATION falcondb_migration FOR ALL TABLES;",
    "# FalconDB 作为订阅方接收变更",
    "-- 详见 FalconDB CDC 订阅配置",
]: cp(line)
bl()

h3("26.1.3  数据导入（FalconDB 侧）")
for line in [
    "# 步骤 1：执行 DDL（建表、索引、约束）",
    "# 注意：移除 pg_dump 中 FalconDB 不支持的语句（如 EXTENSION、TABLESPACE）",
    "psql -h falcondb-host -p 5433 -U falcon -d mydb -f mydb_schema.sql",
    "",
    "# 步骤 2：导入数据（COPY 最快）",
    "psql -h falcondb-host -p 5433 -U falcon -d mydb -c \\",
    '    "COPY orders FROM STDIN WITH (FORMAT CSV, HEADER)" < orders.csv',
    "",
    "# 步骤 3：验证行数",
    "psql -h pg-host -p 5432 -U postgres -d mydb -c 'SELECT COUNT(*) FROM orders;'",
    "psql -h falcondb-host -p 5433 -U falcon -d mydb -c 'SELECT COUNT(*) FROM orders;'",
    "# 两者应相等",
    "",
    "# 步骤 4：验证关键数据",
    "psql -h falcondb-host -p 5433 -U falcon -d mydb -c \\",
    "    'SELECT SUM(total) AS total_revenue FROM orders;'",
    "# 与 PostgreSQL 对比验证",
]: cp(line)
bl()

h2("26.2  零停机切换策略")
np("对于不能停服的生产系统，推荐使用「双写」策略实现零停机迁移：")
for line in [
    "# 零停机迁移流程（约需 2-4 周）",
    "",
    "阶段 1：初始全量同步（1-2 天）",
    "  1. 全量导出 PostgreSQL 数据",
    "  2. 导入到 FalconDB",
    "  3. 验证数据一致性",
    "",
    "阶段 2：双写切流（1-2 周）",
    "  1. 应用层改造：所有写操作同时写入 PG 和 FalconDB",
    "  2. 读操作仍指向 PG（作为权威数据源）",
    "  3. 对比 FalconDB 和 PG 的读取结果，验证一致性",
    "  # 双写示例（Python）：",
    "  with pg_conn.cursor() as pg_cur, falcon_conn.cursor() as falcon_cur:",
    "      sql = 'INSERT INTO orders ... VALUES ...'",
    "      pg_cur.execute(sql, params)        # 主写入",
    "      falcon_cur.execute(sql, params)    # 影子写入",
    "  pg_conn.commit(); falcon_conn.commit()",
    "",
    "阶段 3：读流量切换（灰度，3-5 天）",
    "  1. 5% 读流量切换到 FalconDB，监控正确性",
    "  2. 逐步增加：10% → 25% → 50% → 100%",
    "  3. 每步观察 2-4 小时，确认无异常",
    "",
    "阶段 4：写流量切换（1-2 天）",
    "  1. 停止双写，写操作仅写入 FalconDB",
    "  2. PostgreSQL 进入只读模式（保留 1 周作为回滚保障）",
    "  3. 监控 FalconDB 的所有指标确认正常",
    "",
    "阶段 5：清理（1 天）",
    "  1. 确认 FalconDB 稳定运行 1 周以上",
    "  2. 停止 PostgreSQL 服务",
    "  3. 更新文档，通知所有相关团队",
]: cp(line)
pb()

# ═══════════════════════════════════════════════
# 第二十七章  CDC 变更数据捕获
# ═══════════════════════════════════════════════
h1("第二十七章  变更数据捕获（CDC）")

h2("27.1  CDC 概述")
np("变更数据捕获（Change Data Capture，CDC）是一种将数据库中的所有数据变更（INSERT/UPDATE/DELETE）实时推送到下游消费者的技术。典型用途包括：数据同步到数据仓库（Hive/ClickHouse）、实时搜索索引更新（Elasticsearch）、消息队列（Kafka）、数据湖（Delta Lake）。")
np("FalconDB 内置 CDC 引擎，基于 WAL 日志解析实现逻辑复制协议，兼容 PostgreSQL 的逻辑复制订阅协议，无需外部工具即可将变更流推送到下游。")
bl()

h2("27.2  CDC 配置与启动")
for line in [
    "# falcon.toml 中启用 CDC",
    "[cdc]",
    "enabled = true",
    "slot_name = 'falcon_cdc_slot'    # 复制槽名称（唯一）",
    'format = "json"                  # 输出格式：json / protobuf',
    "include_before = true            # 是否包含 UPDATE 前的旧值",
    "include_schema = true            # 是否包含表结构信息",
    "",
    "# 支持过滤特定表",
    'include_tables = ["orders", "payments", "users"]  # 只发布这些表',
    "# 或使用 exclude_tables 排除",
]: cp(line)
bl()

h2("27.3  CDC 消息格式")
for line in [
    "# INSERT 事件",
    "{",
    '  "op": "insert",',
    '  "ts_ms": 1742000000000,',
    '  "lsn": 12345678,',
    '  "table": "orders",',
    '  "after": {',
    '    "id": 42, "user_id": 1, "status": "pending",',
    '    "total": 99.99, "created_at": "2026-03-12T14:23:45Z"',
    '  }',
    "}",
    "",
    "# UPDATE 事件",
    "{",
    '  "op": "update",',
    '  "ts_ms": 1742000001000,',
    '  "lsn": 12345679,',
    '  "table": "orders",',
    '  "before": { "id": 42, "status": "pending" },',
    '  "after":  { "id": 42, "status": "shipped" }',
    "}",
    "",
    "# DELETE 事件",
    "{",
    '  "op": "delete",',
    '  "ts_ms": 1742000002000,',
    '  "lsn": 12345680,',
    '  "table": "orders",',
    '  "before": { "id": 42, ... }',
    "}",
    "",
    "# DDL 事件（表结构变更）",
    "{",
    '  "op": "schema_change",',
    '  "ts_ms": 1742000003000,',
    '  "table": "users",',
    '  "ddl": "ALTER TABLE users ADD COLUMN phone VARCHAR(20)"',
    "}",
]: cp(line)
bl()

h2("27.4  Kafka 集成示例")
for line in [
    "# 使用 Debezium PostgreSQL Connector 对接 FalconDB（兼容 PG 逻辑复制协议）",
    "{",
    '  "name": "falcondb-source",',
    '  "config": {',
    '    "connector.class": "io.debezium.connector.postgresql.PostgresConnector",',
    '    "database.hostname": "falcondb-host",',
    '    "database.port": "5433",',
    '    "database.user": "falcon",',
    '    "database.password": "${FALCON_PASSWORD}",',
    '    "database.dbname": "mydb",',
    '    "database.server.name": "falcondb",',
    '    "plugin.name": "pgoutput",',
    '    "slot.name": "falcon_cdc_slot",',
    '    "publication.name": "falcondb_publication",',
    '    "topic.prefix": "falcondb.",',
    '    "transforms": "unwrap",',
    '    "transforms.unwrap.type": "io.debezium.transforms.ExtractNewRecordState"',
    '  }',
    "}",
    "",
    "# 消费 Kafka 消息（Python）",
    "from kafka import KafkaConsumer",
    "import json",
    "",
    "consumer = KafkaConsumer(",
    "    'falcondb.mydb.orders',",
    "    bootstrap_servers=['kafka:9092'],",
    "    value_deserializer=lambda m: json.loads(m.decode('utf-8'))",
    ")",
    "",
    "for message in consumer:",
    "    event = message.value",
    "    print(f'op={event[\"op\"]}, table={event[\"source\"][\"table\"]}')",
    "    # 处理：同步到 ClickHouse、更新 Elasticsearch 索引等",
]: cp(line)
pb()

# ═══════════════════════════════════════════════
# 第二十八章  应用开发最佳实践
# ═══════════════════════════════════════════════
h1("第二十八章  应用开发最佳实践")

h2("28.1  连接管理最佳实践")
np("正确的连接管理对 FalconDB 的性能和稳定性至关重要：")
h3("28.1.1  Python（psycopg2 / asyncpg）")
for line in [
    "# psycopg2 连接池（同步）",
    "import psycopg2.pool",
    "",
    "pool = psycopg2.pool.ThreadedConnectionPool(",
    "    minconn=4, maxconn=16,",
    "    host='127.0.0.1', port=5433,",
    "    database='mydb', user='falcon',",
    "    password=os.environ['FALCON_DB_PASSWORD']  # 从环境变量读取",
    ")",
    "",
    "def execute_query(sql, params=None):",
    "    conn = pool.getconn()",
    "    try:",
    "        with conn.cursor() as cur:",
    "            cur.execute(sql, params)",
    "            return cur.fetchall()",
    "    finally:",
    "        pool.putconn(conn)  # 确保连接返回池",
    "",
    "# asyncpg 异步连接池（高性能异步场景）",
    "import asyncpg",
    "",
    "async def create_pool():",
    "    return await asyncpg.create_pool(",
    "        host='127.0.0.1', port=5433,",
    "        database='mydb', user='falcon',",
    "        password=os.environ['FALCON_DB_PASSWORD'],",
    "        min_size=4, max_size=16,",
    "        command_timeout=30  # 30 秒命令超时",
    "    )",
    "",
    "# OCC 重试装饰器",
    "import asyncio, functools",
    "from asyncpg.exceptions import SerializationError",
    "",
    "def retry_on_conflict(max_retries=3):",
    "    def decorator(func):",
    "        @functools.wraps(func)",
    "        async def wrapper(*args, **kwargs):",
    "            for attempt in range(max_retries):",
    "                try:",
    "                    return await func(*args, **kwargs)",
    "                except SerializationError:",
    "                    if attempt == max_retries - 1:",
    "                        raise",
    "                    await asyncio.sleep(0.01 * (2 ** attempt))  # 指数退避",
    "        return wrapper",
    "    return decorator",
    "",
    "@retry_on_conflict(max_retries=3)",
    "async def transfer_funds(pool, from_id, to_id, amount):",
    "    async with pool.acquire() as conn:",
    "        async with conn.transaction():",
    "            await conn.execute(",
    "                'UPDATE accounts SET balance=balance-$1 WHERE id=$2',",
    "                amount, from_id",
    "            )",
    "            await conn.execute(",
    "                'UPDATE accounts SET balance=balance+$1 WHERE id=$2',",
    "                amount, to_id",
    "            )",
]: cp(line)
bl()

h3("28.1.2  Go（pgx）")
for line in [
    'import "github.com/jackc/pgx/v5/pgxpool"',
    "",
    "// 创建连接池",
    "config, _ := pgxpool.ParseConfig(",
    '    "postgres://falcon:${FALCON_PASSWORD}@127.0.0.1:5433/mydb"',
    ")",
    "config.MaxConns = 16",
    "config.MinConns = 4",
    "config.MaxConnLifetime = 30 * time.Minute",
    "config.MaxConnIdleTime = 60 * time.Second",
    "",
    "pool, err := pgxpool.NewWithConfig(ctx, config)",
    "",
    "// OCC 重试函数",
    "func withRetry(ctx context.Context, pool *pgxpool.Pool,",
    "    fn func(pgx.Tx) error) error {",
    "    for attempt := 0; attempt < 3; attempt++ {",
    "        err := pgx.BeginTxFunc(ctx, pool, pgx.TxOptions{}, fn)",
    "        if err == nil { return nil }",
    "        // 检查是否为序列化失败（40001）",
    "        var pgErr *pgconn.PgError",
    '        if errors.As(err, &pgErr) && pgErr.Code == "40001" {',
    "            time.Sleep(time.Duration(10*(1<<attempt)) * time.Millisecond)",
    "            continue",
    "        }",
    "        return err",
    "    }",
    "    return fmt.Errorf(\"max retries exceeded\")",
    "}",
]: cp(line)
bl()

h2("28.2  Schema 设计最佳实践")
h3("28.2.1  主键设计")
np("主键设计直接影响 FalconDB 的哈希查找性能和存储效率：")
tbl(["主键类型","适用场景","优点","注意事项"],
[["BIGSERIAL（自增）","单分片表，新增场景","简单，紧凑，范围扫描好","多分片下可能有热点（全局序列）"],
 ["UUID（随机）","多分片分布式场景","均匀分布，无序列协调问题","范围扫描性能差，索引大"],
 ["UUID v7（有序时间）","兼顾分布和时序","有序（按时间），均匀分布","需要应用层生成"],
 ["复合主键（user_id, order_id）","多租户/分区场景","与分片键一致，本地化事务","查询必须携带 user_id"]])
for line in [
    "-- 推荐：UUID v7 主键（有序 + 分布均匀）",
    "CREATE TABLE orders (",
    "    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),  -- 随机 UUID",
    "    -- 或 id UUID PRIMARY KEY DEFAULT uuid_generate_v7(), -- 有序 UUID（需应用生成）",
    "    user_id BIGINT NOT NULL,",
    "    ...",
    ");",
    "",
    "-- 高性能场景：BIGINT 主键 + 应用层 Snowflake ID",
    "CREATE TABLE events (",
    "    id BIGINT PRIMARY KEY,  -- Snowflake ID（时间戳+机器ID+序号）",
    "    ...",
    ");",
]: cp(line)
bl()

h3("28.2.2  避免宽表（Wide Table）反模式")
np("宽表（单表有 50+ 列）在 FalconDB 中会导致内存占用高、索引膨胀、查询慢等问题。推荐垂直拆分：")
for line in [
    "-- 反模式：单表 100+ 列",
    "CREATE TABLE users (id BIGINT, name TEXT, email TEXT, phone TEXT,",
    "    address TEXT, city TEXT, country TEXT, zip TEXT,",
    "    pref_lang TEXT, pref_currency TEXT, pref_timezone TEXT,",
    "    bio TEXT, avatar_url TEXT,",
    "    created_at TIMESTAMPTZ, updated_at TIMESTAMPTZ,",
    "    -- ... 100 多列",
    ");",
    "",
    "-- 推荐：垂直拆分为核心表 + 扩展表",
    "CREATE TABLE users (  -- 核心表（高频访问，小行）",
    "    id BIGINT PRIMARY KEY,",
    "    email TEXT UNIQUE NOT NULL,",
    "    created_at TIMESTAMPTZ DEFAULT now()",
    ");",
    "",
    "CREATE TABLE user_profiles (  -- 扩展表（低频访问，允许宽）",
    "    user_id BIGINT PRIMARY KEY REFERENCES users(id),",
    "    name TEXT, phone TEXT, address TEXT,",
    "    bio TEXT, avatar_url TEXT",
    ");",
    "",
    "CREATE TABLE user_preferences (  -- 偏好表（独立变更）",
    "    user_id BIGINT PRIMARY KEY REFERENCES users(id),",
    "    lang TEXT DEFAULT 'zh-CN',",
    "    currency TEXT DEFAULT 'CNY',",
    "    timezone TEXT DEFAULT 'Asia/Shanghai'",
    ");",
]: cp(line)
bl()

h2("28.3  错误处理规范")
np("应用层应能够正确区分和处理不同类型的数据库错误：")
tbl(["错误类型","SQLSTATE","重试策略","示例"],
[["OCC 冲突（序列化失败）","40001","立即重试（指数退避，最多 3-5 次）","并发更新同一行"],
 ["死锁","40P01","立即重试（同 40001）","FalconDB 极少出现，但应处理"],
 ["唯一键冲突","23505","不重试，返回业务错误","重复注册相同邮箱"],
 ["非空约束","23502","不重试，修复应用逻辑","必填字段为空"],
 ["连接失败","08xxx","重试（带指数退避，共 3 次）","数据库重启/网络抖动"],
 ["查询超时","57014","不重试，优化查询或增大超时","慢查询超过 statement_timeout"],
 ["内部错误","XX000","记录日志，不重试，告警","数据库 Panic（罕见）"]])
for line in [
    "# Python 完整错误处理示例",
    "import psycopg2",
    "from psycopg2 import errorcodes",
    "",
    "def execute_with_retry(conn, sql, params, max_retries=3):",
    "    for attempt in range(max_retries):",
    "        try:",
    "            with conn.cursor() as cur:",
    "                cur.execute(sql, params)",
    "                conn.commit()",
    "                return cur.fetchall()",
    "        except psycopg2.Error as e:",
    "            conn.rollback()",
    "            if e.pgcode == errorcodes.SERIALIZATION_FAILURE:  # 40001",
    "                if attempt < max_retries - 1:",
    "                    time.sleep(0.01 * (2 ** attempt))",
    "                    continue",
    "            elif e.pgcode == errorcodes.UNIQUE_VIOLATION:  # 23505",
    "                raise BusinessError(f'Duplicate entry: {e.pgerror}')",
    "            elif e.pgcode and e.pgcode.startswith('08'):  # 连接错误",
    "                if attempt < max_retries - 1:",
    "                    time.sleep(1 * (attempt + 1))",
    "                    conn = reconnect()  # 重新建立连接",
    "                    continue",
    "            raise  # 其他错误直接抛出",
    "    raise MaxRetriesExceeded(f'Failed after {max_retries} attempts')",
]: cp(line)
pb()

# ═══════════════════════════════════════════════
# 第二十九章  典型应用场景案例
# ═══════════════════════════════════════════════
h1("第二十九章  典型应用场景案例")

h2("29.1  案例一：电商订单系统")
np("电商订单系统是 FalconDB 最典型的应用场景，具有高并发写入（下单）、强一致性要求（库存扣减不超卖）、低延迟点查（订单查询）等特征。")
h3("29.1.1  数据模型设计")
for line in [
    "-- 用户表（高频点查，Rowstore 内存引擎）",
    "CREATE TABLE users (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    username VARCHAR(50) UNIQUE NOT NULL,",
    "    email VARCHAR(200) UNIQUE NOT NULL,",
    "    balance DECIMAL(12,2) NOT NULL DEFAULT 0 CHECK (balance >= 0),",
    "    created_at TIMESTAMPTZ DEFAULT now()",
    ") ENGINE=rowstore;",
    "",
    "-- 商品表（含库存，高并发扣减）",
    "CREATE TABLE products (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    name TEXT NOT NULL,",
    "    price DECIMAL(10,2) NOT NULL,",
    "    stock INT NOT NULL DEFAULT 0 CHECK (stock >= 0),",
    "    category TEXT,",
    "    updated_at TIMESTAMPTZ DEFAULT now()",
    ") ENGINE=rowstore;",
    "",
    "-- 订单表（高写入，历史归档）",
    "CREATE TABLE orders (",
    "    id BIGSERIAL PRIMARY KEY,",
    "    user_id BIGINT NOT NULL REFERENCES users(id),",
    "    product_id BIGINT NOT NULL REFERENCES products(id),",
    "    quantity INT NOT NULL CHECK (quantity > 0),",
    "    unit_price DECIMAL(10,2) NOT NULL,",
    "    total DECIMAL(12,2) NOT NULL,",
    "    status VARCHAR(20) DEFAULT 'pending',",
    "    created_at TIMESTAMPTZ DEFAULT now()",
    ") ENGINE=rowstore;",
    "",
    "-- 关键索引",
    "CREATE INDEX idx_orders_user ON orders(user_id, created_at DESC);",
    "CREATE INDEX idx_orders_status ON orders(status) WHERE status != 'completed';",
    "CREATE INDEX idx_products_category ON products(category, price);",
]: cp(line)
bl()

h3("29.1.2  下单事务（原子扣减库存）")
for line in [
    "-- 原子性下单：库存扣减 + 订单创建（单分片快速路径）",
    "CREATE OR REPLACE FUNCTION place_order(",
    "    p_user_id BIGINT,",
    "    p_product_id BIGINT,",
    "    p_quantity INT,",
    "    OUT o_order_id BIGINT,",
    "    OUT o_result TEXT",
    ") LANGUAGE plpgsql AS $$",
    "DECLARE",
    "    v_price DECIMAL(10,2);",
    "    v_stock INT;",
    "    v_total DECIMAL(12,2);",
    "BEGIN",
    "    -- 加锁查询商品（防止并发超卖）",
    "    SELECT price, stock INTO v_price, v_stock",
    "    FROM products WHERE id = p_product_id FOR UPDATE;",
    "",
    "    IF v_stock < p_quantity THEN",
    "        o_result := 'ERROR: insufficient stock';",
    "        RETURN;",
    "    END IF;",
    "",
    "    v_total := v_price * p_quantity;",
    "",
    "    -- 扣减库存",
    "    UPDATE products",
    "    SET stock = stock - p_quantity, updated_at = now()",
    "    WHERE id = p_product_id;",
    "",
    "    -- 创建订单",
    "    INSERT INTO orders (user_id, product_id, quantity, unit_price, total)",
    "    VALUES (p_user_id, p_product_id, p_quantity, v_price, v_total)",
    "    RETURNING id INTO o_order_id;",
    "",
    "    o_result := 'SUCCESS';",
    "EXCEPTION",
    "    WHEN check_violation THEN",
    "        o_result := 'ERROR: stock went negative (concurrent conflict)';",
    "END;",
    "$$;",
    "",
    "-- 调用",
    "SELECT o_order_id, o_result FROM place_order(1, 101, 2);",
]: cp(line)
bl()

h2("29.2  案例二：金融支付系统")
np("金融支付系统对一致性和持久性要求极高，是 FalconDB 确定性提交保证（DCG）最重要的应用场景之一。")
h3("29.2.1  配置：最高持久性模式")
for line in [
    "[wal]",
    'durability_policy = "sync-replica"  # 等待从节点确认，RPO ≈ 0',
    'sync_mode = "fdatasync"',
    "group_commit = true",
    "flush_interval_us = 500              # 低延迟（500μs）",
    "",
    "[security]",
    "# 金融场景必须启用 TLS + TDE",
    "[security.tls]",
    "enabled = true",
    'min_version = "1.3"',
    "",
    "[tde]",
    "enabled = true",
    'key_env_var = "PAYMENT_DB_TDE_KEY"',
]: cp(line)
bl()

h3("29.2.2  账户余额管理")
for line in [
    "-- 账户表（每行代表一个账户）",
    "CREATE TABLE accounts (",
    "    id BIGINT PRIMARY KEY,",
    "    owner_id BIGINT NOT NULL,",
    "    currency CHAR(3) NOT NULL DEFAULT 'CNY',",
    "    balance DECIMAL(18,6) NOT NULL DEFAULT 0",
    "        CONSTRAINT positive_balance CHECK (balance >= 0),",
    "    version BIGINT NOT NULL DEFAULT 0,   -- 乐观锁版本号",
    "    updated_at TIMESTAMPTZ DEFAULT now()",
    ");",
    "",
    "-- 流水表（不可变，仅追加）",
    "CREATE TABLE transactions (",
    "    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),",
    "    from_account BIGINT REFERENCES accounts(id),",
    "    to_account BIGINT REFERENCES accounts(id),",
    "    amount DECIMAL(18,6) NOT NULL CHECK (amount > 0),",
    "    fee DECIMAL(18,6) NOT NULL DEFAULT 0,",
    "    txn_type VARCHAR(20) NOT NULL,",
    "    status VARCHAR(20) NOT NULL DEFAULT 'pending',",
    "    idempotency_key UUID UNIQUE,   -- 幂等键，防止重复提交",
    "    created_at TIMESTAMPTZ DEFAULT now()",
    ") ENGINE=rocksdb;  -- 历史流水用磁盘引擎",
    "",
    "-- 幂等转账（防止网络重试导致重复扣款）",
    "INSERT INTO transactions (",
    "    from_account, to_account, amount, txn_type, idempotency_key",
    ") VALUES ($1, $2, $3, 'transfer', $4)",
    "ON CONFLICT (idempotency_key) DO NOTHING  -- 幂等：相同 key 只执行一次",
    "RETURNING id, status;",
]: cp(line)
bl()

h2("29.3  案例三：IoT 时序数据")
np("IoT（物联网）场景的特点是高频写入（每秒百万级设备数据点）、时序查询（最近 N 小时/天的聚合）、数据老化（超过保留期的数据定期清理）。")
for line in [
    "-- 时序数据表（LSM 引擎，适合高吞吐追加写入）",
    "CREATE TABLE sensor_data (",
    "    device_id BIGINT NOT NULL,",
    "    ts TIMESTAMPTZ NOT NULL,",
    "    temperature FLOAT8,",
    "    humidity FLOAT8,",
    "    pressure FLOAT8,",
    "    -- 复合主键：时间 + 设备 ID（分区查询高效）",
    "    PRIMARY KEY (device_id, ts)",
    ") ENGINE=lsm;",
    "",
    "-- 时序索引（设备 + 时间范围）",
    "CREATE INDEX idx_sensor_time ON sensor_data (device_id, ts DESC);",
    "",
    "-- 批量写入（高吞吐，每批 1000 条）",
    "COPY sensor_data (device_id, ts, temperature, humidity, pressure)",
    "FROM '/data/sensor_batch.csv' WITH (FORMAT CSV);",
    "",
    "-- 典型时序聚合查询",
    "SELECT",
    "    date_trunc('hour', ts) AS hour,",
    "    AVG(temperature) AS avg_temp,",
    "    MAX(temperature) AS max_temp,",
    "    MIN(temperature) AS min_temp,",
    "    COUNT(*) AS sample_count",
    "FROM sensor_data",
    "WHERE device_id = 12345",
    "  AND ts >= now() - INTERVAL '24 hours'",
    "GROUP BY 1",
    "ORDER BY 1 DESC;",
    "",
    "-- 最新读数（设备最新状态）",
    "SELECT DISTINCT ON (device_id)",
    "    device_id, ts, temperature, humidity",
    "FROM sensor_data",
    "WHERE ts >= now() - INTERVAL '1 hour'",
    "ORDER BY device_id, ts DESC;",
    "",
    "-- 定期清理过期数据（保留最近 30 天）",
    "DELETE FROM sensor_data",
    "WHERE ts < now() - INTERVAL '30 days';",
    "-- 建议：在维护窗口执行，或使用分区表+DROP PARTITION",
]: cp(line)
pb()

# ═══════════════════════════════════════════════
# 第三十章  运维自动化脚本参考
# ═══════════════════════════════════════════════
h1("第三十章  运维自动化脚本参考")

h2("30.1  健康检查脚本")
for line in [
    "# Windows PowerShell 健康检查脚本",
    "# 保存为 C:\\Scripts\\falcon_health_check.ps1",
    "",
    "$server = '127.0.0.1'",
    "$port = 5433",
    "$user = 'falcon'",
    "$alerts = @()",
    "",
    "# 1. 检查服务状态",
    "$svc = Get-Service -Name 'FalconDB' -ErrorAction SilentlyContinue",
    "if ($svc.Status -ne 'Running') {",
    "    $alerts += '【严重】FalconDB 服务未运行'",
    "}",
    "",
    "# 2. 检查端口连通性",
    "$conn = Test-NetConnection -ComputerName $server -Port $port -WarningAction SilentlyContinue",
    "if (-not $conn.TcpTestSucceeded) {",
    "    $alerts += '【严重】无法连接 FalconDB 端口 5433'",
    "}",
    "",
    "# 3. 检查 Prometheus 指标",
    "try {",
    "    $metrics = Invoke-WebRequest -Uri 'http://127.0.0.1:9090/metrics' -TimeoutSec 5",
    "    if ($metrics.StatusCode -ne 200) {",
    "        $alerts += '【警告】Prometheus 指标端点响应异常'",
    "    }",
    "} catch {",
    "    $alerts += '【警告】无法访问 Prometheus 指标端点'",
    "}",
    "",
    "# 4. 检查日志中的 ERROR",
    "$logFile = 'C:\\ProgramData\\FalconDB\\logs\\falcon.log'",
    "$recentErrors = Select-String -Path $logFile -Pattern 'ERROR|panic' |",
    "    Where-Object { $_.Line -match (Get-Date -Format 'yyyy-MM-dd') }",
    "if ($recentErrors.Count -gt 0) {",
    "    $alerts += \"【警告】今日日志中发现 $($recentErrors.Count) 条错误\"",
    "}",
    "",
    "# 5. 输出结果",
    "if ($alerts.Count -eq 0) {",
    "    Write-Host '【正常】FalconDB 健康检查通过' -ForegroundColor Green",
    "} else {",
    "    $alerts | ForEach-Object { Write-Host $_ -ForegroundColor Red }",
    "    exit 1  # 非零退出码，供监控系统检测",
    "}",
]: cp(line)
bl()

h2("30.2  自动备份脚本")
for line in [
    "# Windows PowerShell 自动备份脚本",
    "# 保存为 C:\\Scripts\\falcon_backup.ps1",
    "# 通过 Windows 任务计划每日 02:00 执行",
    "",
    "$backupRoot = 'D:\\backup\\falcondb'",
    "$dataDir = 'C:\\ProgramData\\FalconDB\\data'",
    "$logFile = 'C:\\Scripts\\backup.log'",
    "$retentionDays = 30",
    "",
    "function Write-Log($msg) {",
    '    $ts = Get-Date -Format "yyyy-MM-dd HH:mm:ss"',
    '    "$ts $msg" | Tee-Object -FilePath $logFile -Append',
    "}",
    "",
    "Write-Log '=== 开始备份 FalconDB ==='",
    "",
    "# 1. 执行 CHECKPOINT",
    "Write-Log '执行 CHECKPOINT...'",
    "$env:PGPASSWORD = $env:FALCON_DB_PASSWORD",
    "& psql -h 127.0.0.1 -p 5433 -U falcon -c 'CHECKPOINT;' 2>&1 | Write-Log",
    "",
    "# 2. 创建备份目录",
    '$ts = Get-Date -Format "yyyyMMdd_HHmmss"',
    '$bdir = "$backupRoot\\$ts"',
    "New-Item -ItemType Directory -Path $bdir | Out-Null",
    "",
    "# 3. 复制数据",
    "Write-Log '复制数据目录...'",
    "Copy-Item -Recurse $dataDir '$bdir\\data' -ErrorAction Stop",
    "Copy-Item 'C:\\ProgramData\\FalconDB\\conf\\falcon.toml' $bdir",
    "",
    "# 4. 计算大小并记录",
    '$size = (Get-ChildItem -Recurse $bdir | Measure-Object Length -Sum).Sum',
    '$sizeMB = [math]::Round($size/1MB, 2)',
    "Write-Log \"备份完成，大小: ${sizeMB} MB，路径: $bdir\"",
    "",
    "# 5. 清理过期备份",
    '$cutoff = (Get-Date).AddDays(-$retentionDays)',
    "Get-ChildItem $backupRoot -Directory |",
    "    Where-Object { $_.CreationTime -lt $cutoff } |",
    "    ForEach-Object {",
    "        Remove-Item $_.FullName -Recurse -Force",
    "        Write-Log \"已删除过期备份: $($_.Name)\"",
    "    }",
    "",
    "Write-Log '=== 备份完成 ==='",
]: cp(line)
bl()

h2("30.3  性能监控脚本")
for line in [
    "# 每 5 分钟采集关键指标到 CSV 文件",
    "# 保存为 C:\\Scripts\\falcon_metrics.ps1",
    "",
    "$csvFile = 'C:\\Scripts\\falcon_metrics.csv'",
    "$env:PGPASSWORD = $env:FALCON_DB_PASSWORD",
    "",
    "# 初始化 CSV 标题",
    "if (-not (Test-Path $csvFile)) {",
    '    "timestamp,txn_committed,txn_conflicts,connections,gc_versions" | Set-Content $csvFile',
    "}",
    "",
    "# 采集指标",
    "while ($true) {",
    '    $ts = Get-Date -Format "yyyy-MM-dd HH:mm:ss"',
    "    ",
    "    # 通过 psql 查询 SHOW 命令",
    "    $stats = & psql -h 127.0.0.1 -p 5433 -U falcon",
    "        -c 'SHOW falcon.txn_stats;' -t 2>&1",
    "    ",
    "    # 解析输出（实际需要更完整的解析逻辑）",
    '    "$ts,$stats" | Add-Content $csvFile',
    "    ",
    "    Start-Sleep -Seconds 300  # 5 分钟采集一次",
    "}",
]: cp(line)
pb()

print("Part 6 完成，保存中...")
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
