# -*- coding: utf-8 -*-
"""FalconDB 操作手册生成器 Part 3 - 备份/监控/AI/性能/故障排除/检查清单"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\falcondb\falcondb-main\docs\FalconDB_操作手册.docx"
doc = Document(OUTPUT)

def h1(t): return doc.add_heading(t, 1)
def h2(t): return doc.add_heading(t, 2)
def h3(t): return doc.add_heading(t, 3)

def np(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(11)
    return p

def bp(t):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t); r.font.size = Pt(11)
    return p

def cp(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(t)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
    pPr.append(shd)
    return p

def bl(): doc.add_paragraph()

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        if c.paragraphs[0].runs:
            run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        r = t.rows[ri+1]
        for ci,v in enumerate(row):
            r.cells[ci].text = str(v)
            if r.cells[ci].paragraphs[0].runs:
                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return t

def pb(): doc.add_page_break()

# ══ 第十四章 备份与恢复 ══
h1("第十四章  备份与恢复")
h2("14.1  备份策略概述")
np("FalconDB 提供三种互补的备份策略，可根据 RPO/RTO 目标组合使用：")
tbl(
    ["备份类型", "内容", "RPO", "RTO", "典型使用"],
    [
        ["全量备份（Full Backup）", "所有数据文件 + 元数据 + WAL 历史", "N/A（作为基线）", "分钟至小时", "每日或每周基线备份"],
        ["增量备份（Incremental）", "上次备份以来的 WAL/变更日志", "分钟级", "全量恢复时间 + 增量重放", "每小时定期备份"],
        ["WAL 归档（PITR）", "持续实时 WAL 流式归档", "秒级", "分钟级 + WAL 重放", "金融/支付等精确恢复"],
    ]
)
bl()

h2("14.2  备份存储目标")
for line in [
    "-- 本地文件系统备份",
    'BackupTarget::Local { path: "/backup/falcondb" }',
    "",
    "-- Amazon S3 备份",
    "BackupTarget::S3 {",
    '    bucket: "my-prod-backups",',
    '    prefix: "falcondb/daily/",',
    '    region: "ap-east-1",',
    "}",
    "",
    "-- 阿里云 OSS 备份",
    "BackupTarget::Oss {",
    '    bucket: "my-prod-backups",',
    '    prefix: "falcondb/daily/",',
    '    endpoint: "oss-cn-hangzhou.aliyuncs.com",',
    "}",
]:
    cp(line)
bl()

h2("14.3  Windows 手动一致性备份")
np("对于 Rowstore 内存引擎，热备份需要先停止服务或执行 CHECKPOINT 以确保数据一致性：")
for line in [
    "# 方法一：停服备份（强一致性，推荐生产）",
    "falcon service stop",
    "",
    '$ts = Get-Date -Format "yyyyMMdd_HHmmss"',
    r'$backup_dir = "D:\backup\falcon_$ts"',
    r'New-Item -ItemType Directory -Path $backup_dir',
    r'Copy-Item -Recurse "C:\ProgramData\FalconDB\data" "$backup_dir\data"',
    r'Copy-Item "C:\ProgramData\FalconDB\conf\falcon.toml" "$backup_dir\falcon.toml"',
    "",
    "falcon service start",
    "",
    "# 方法二：CHECKPOINT 后热备份（服务不停止）",
    "psql -h 127.0.0.1 -p 5433 -U falcon -c 'CHECKPOINT;'",
    r'Copy-Item -Recurse "C:\ProgramData\FalconDB\data" "D:\backup\falcon_hot_$ts"',
]:
    cp(line)
bl()

h2("14.4  PITR 配置与使用")
np("时间点恢复（PITR）允许将数据库恢复到过去任意指定时刻，是应对误操作（错误 DELETE/UPDATE）的最有效手段。")
for line in [
    "# 1. 配置 WAL 归档（falcon.toml）",
    "[pitr]",
    "enabled = true",
    'archive_dir = "/wal_archive"   # WAL 归档目录，建议独立存储',
    "retention_hours = 168          # 保留 7 天（168 小时）",
    "",
    "# 2. 触发恢复（指定目标时间戳）",
    "# 先停止 FalconDB 服务",
    "falcon service stop",
    "",
    "# 从基础备份恢复数据目录",
    r'Copy-Item -Recurse "D:\backup\falcon_20260312_020000\data" "C:\ProgramData\FalconDB\data_restore"',
    "",
    "# 配置 PITR 恢复目标（在 falcon.toml 中添加）",
    "[recovery]",
    'target_time = "2026-03-12 14:30:00"   # 精确到秒的恢复目标',
    'archive_dir = "/wal_archive"',
    "",
    "# 启动数据库，自动重放 WAL 至目标时间",
    "falcon service start",
]:
    cp(line)
bl()

h2("14.5  恢复类型")
tbl(
    ["恢复类型", "描述", "使用场景"],
    [
        ["Latest（最新状态）", "恢复到最新可用状态，重放所有 WAL", "磁盘故障、主节点彻底损坏"],
        ["ToLsn（指定 LSN）", "精确到 WAL 日志序号的恢复", "已知 LSN 的精确恢复点"],
        ["ToTimestamp（指定时间）", "恢复到指定时间戳，重放 WAL 至该时刻", "误操作后的 PITR 恢复"],
        ["NewCluster（新集群）", "将备份恢复到全新集群（用于迁移/克隆）", "数据库迁移、环境克隆"],
    ]
)
bl()

h2("14.6  备份验证")
np("备份后应进行验证，确保备份文件可以成功恢复：")
for line in [
    "# 验证备份完整性（检查 WAL CRC 和元数据）",
    "falcon verify-backup --backup-dir D:\\backup\\falcon_20260312_020000",
    "",
    "# 在独立测试环境进行恢复验证（生产环境定期执行）",
    "falcon service stop   # 停止测试环境的 FalconDB",
    r"Copy-Item -Recurse D:\backup\falcon_20260312_020000\data C:\falcondb_test\data",
    "falcon --data-dir C:\\falcondb_test\\data --no-wal   # 启动验证实例",
    "psql -h 127.0.0.1 -p 5434 -U falcon -c 'SELECT COUNT(*) FROM orders;'",
]:
    cp(line)
bl()

h2("14.7  备份监控指标")
tbl(
    ["指标", "说明"],
    [
        ["backups_started", "已启动的备份任务总数"],
        ["backups_completed", "成功完成的备份任务数"],
        ["backups_failed", "失败的备份任务数"],
        ["restores_started / restores_completed", "恢复任务启动/完成数"],
        ["total_bytes_backed_up", "累计备份字节数"],
    ]
)
np("所有备份恢复操作均记录到企业审计日志（类别：BACKUP_RESTORE），包含：任务 ID、类型、目标、字节数、耗时。")
pb()

# ══ 第十五章 监控与可观测性 ══
h1("第十五章  监控与可观测性")
h2("15.1  SHOW 命令完整参考")
tbl(
    ["命令", "说明", "输出关键字段"],
    [
        ["SHOW falcon.status", "服务器整体状态", "版本、角色、运行时间、WAL LSN"],
        ["SHOW falcon.connections", "当前连接详情", "连接数、每个会话的 IP 和查询"],
        ["SHOW falcon.txn", "活跃事务列表", "TxnId、开始时间戳、类型（Local/Global）"],
        ["SHOW falcon.txn_stats", "事务统计汇总", "提交数、回滚数、OCC 冲突数、快/慢路径比例"],
        ["SHOW falcon.txn_history", "近期事务历史", "最近事务的类型、耗时、状态"],
        ["SHOW falcon.gc_stats", "MVCC GC 统计", "已回收版本数、字节数、安全点时间戳"],
        ["SHOW falcon.gc_safepoint", "GC 安全点诊断", "阻塞 GC 的最老活跃事务 ID 和开始时间"],
        ["SHOW falcon.replication_stats", "复制统计", "副本延迟 LSN、WAL 传输速率、故障转移次数"],
        ["SHOW falcon.replica_stats", "从节点状态（从节点执行）", "applied_lsn、复制延迟毫秒数"],
        ["SHOW falcon.slow_queries", "慢查询日志", "SQL、耗时、计划类型、表名"],
        ["SHOW falcon.table_stats", "表统计信息", "行数估计、存储大小、引擎类型"],
        ["SHOW falcon.checkpoint_stats", "检查点统计", "上次检查点时间、WAL 落后量"],
        ["SHOW falcon.plan_cache", "查询计划缓存", "缓存命中率、缓存条目数"],
        ["SHOW falcon.security", "安全态势", "TLS 状态、加密、IP 白名单"],
        ["SHOW falcon.security_audit", "安全组件指标", "速率限制、密码策略、防火墙拦截数"],
        ["SHOW falcon.audit_log", "近期审计事件", "事件类型、时间、用户、对象"],
        ["SHOW falcon.scatter_stats", "分布式查询统计", "scatter/gather 操作数、延迟"],
        ["SHOW AI STATS", "AI 优化器状态", "是否启用、训练样本数、模型就绪状态、MAE"],
        ["SHOW AIOPS STATS", "AIOps 引擎汇总", "慢查询阈值、异常计数、索引建议数、指纹数"],
        ["SHOW AIOPS ALERTS", "异常告警列表", "告警 ID、时间、严重级别、来源、描述"],
        ["SHOW AIOPS INDEX ADVICE", "索引建议", "表名、扫描次数、耗时、建议 SQL"],
        ["SHOW AIOPS WORKLOAD", "工作负载 Top-N", "SQL 指纹、调用次数、p95/p99 延迟、错误数"],
        ["SHOW AIOPS SLOW QUERIES", "慢查询记录", "时间戳、耗时、阈值、计划类型、SQL 文本"],
        ["SELECT * FROM pg_stat_statements", "语句级性能统计", "兼容 PostgreSQL pg_stat_statements 视图"],
    ]
)
bl()

h2("15.2  Prometheus 指标详表")
h3("15.2.1  事务指标")
tbl(
    ["指标名称", "类型", "说明"],
    [
        ["falcon_txn_committed_total", "Counter", "总已提交事务数（累计）"],
        ["falcon_txn_aborted_total", "Counter", "总回滚/中止事务数"],
        ["falcon_txn_fast_path_total", "Counter", "本地快速路径（OCC）提交数"],
        ["falcon_txn_slow_path_total", "Counter", "全局慢速路径（2PC）提交数"],
        ["falcon_txn_occ_conflicts_total", "Counter", "OCC 写-写冲突次数（需重试）"],
        ["falcon_txn_active", "Gauge", "当前进行中的事务数"],
        ["falcon_txn_commit_latency_us", "Histogram", "提交延迟分布（微秒，含 P50/P95/P99）"],
    ]
)
bl()

h3("15.2.2  WAL 和存储指标")
tbl(
    ["指标名称", "类型", "说明"],
    [
        ["falcon_wal_bytes_written_total", "Counter", "WAL 累计写入字节数"],
        ["falcon_wal_segments_total", "Counter", "WAL 段文件总创建数"],
        ["falcon_wal_flush_latency_us", "Histogram", "WAL fdatasync 延迟（关键性能指标）"],
        ["falcon_wal_backlog_bytes", "Gauge", "当前 WAL 未刷写积压字节数"],
        ["falcon_memory_used_bytes", "Gauge", "Rowstore 当前内存占用"],
        ["falcon_memory_limit_bytes", "Gauge", "配置的内存预算上限"],
        ["falcon_memory_pressure_total", "Counter", "触发内存背压事件的次数"],
    ]
)
bl()

h3("15.2.3  查询和连接指标")
tbl(
    ["指标名称", "类型", "说明"],
    [
        ["falcon_query_total", "Counter", "执行的查询总数"],
        ["falcon_query_latency_us", "Histogram", "端到端查询延迟（含 P50/P95/P99）"],
        ["falcon_slow_query_total", "Counter", "超过慢查询阈值的查询数"],
        ["falcon_query_errors_total", "Counter", "按错误类型分组（user/retryable/transient/bug）"],
        ["falcon_connections_active", "Gauge", "当前活跃连接数"],
    ]
)
bl()

h3("15.2.4  GC 和复制指标")
tbl(
    ["指标名称", "类型", "说明"],
    [
        ["falcon_gc_versions_reclaimed_total", "Counter", "GC 回收的 MVCC 版本总数"],
        ["falcon_gc_bytes_reclaimed_total", "Counter", "GC 回收的总字节数"],
        ["falcon_gc_safepoint_ts", "Gauge", "当前 GC 安全点时间戳"],
        ["falcon_replication_lag_lsn", "Gauge", "副本与主节点的 LSN 差距"],
        ["falcon_replication_applied_lsn", "Gauge", "副本最新应用的 LSN"],
        ["falcon_replication_reconnects_total", "Counter", "副本重连次数（非零说明网络不稳定）"],
    ]
)
bl()

h2("15.3  关键性能指标告警阈值（建议）")
tbl(
    ["指标", "警告阈值", "严重阈值", "说明"],
    [
        ["falcon_txn_commit_latency_us P99", "> 10 ms", "> 50 ms", "提交延迟过高，检查 WAL 刷盘和复制延迟"],
        ["falcon_replication_lag_lsn", "> 1000 LSN", "> 10000 LSN", "复制延迟过大，检查网络和从节点性能"],
        ["falcon_txn_occ_conflicts_total（增速）", "> 100/min", "> 1000/min", "大量 OCC 冲突，检查热点行争用"],
        ["falcon_connections_active", "> 800", "> 950", "接近 max_connections 限制（1024）"],
        ["falcon_memory_used_bytes / limit", "> 70%", "> 90%", "内存使用率高，可能触发背压"],
        ["falcon_wal_flush_latency_us P99", "> 5 ms", "> 20 ms", "WAL 刷盘慢，检查磁盘 I/O"],
        ["falcon_gc_safepoint_ts 停滞", "停滞 > 5 min", "停滞 > 30 min", "长事务阻塞 GC，检查 falcon.txn"],
        ["falcon_crash_domain_panic_count", "> 0", "> 5", "出现 Panic，立即检查错误日志"],
    ]
)
bl()

h2("15.4  结构化日志配置")
for line in [
    "# 文本格式（人可读，推荐开发环境）",
    '[logging]',
    'level = "info"',
    'format = "text"',
    'stderr = true',
    "",
    "# JSON 格式（机器可解析，推荐生产环境对接 ELK/Loki）",
    '[logging]',
    'level = "info"',
    'format = "json"',
    'file = "/var/log/falcondb/falcon.log"',
    'rotation = "daily"',
    'max_size_mb = 500',
    'max_files = 30',
    'slow_query_ms = 100    # 记录执行超过 100ms 的 SQL',
]:
    cp(line)
bl()

h2("15.5  查询阶段延迟分析")
np("每条 SQL 查询均记录各阶段耗时（微秒），帮助定位性能瓶颈：")
tbl(
    ["阶段字段", "说明", "过高时的排查方向"],
    [
        ["parse_us", "SQL 文本 → AST 解析时间", "极复杂 SQL 或超长 SQL"],
        ["bind_us", "名称解析和类型检查", "大量表/列元数据查找"],
        ["plan_us", "生成物理执行计划", "AI 优化器候选计划过多"],
        ["execute_us", "算子执行时间", "全表扫描、大量数据排序、JOIN 优化"],
        ["commit_us", "WAL 刷盘 + 复制确认", "磁盘慢、复制延迟、同步复制模式"],
        ["wal_flush_us", "专项 WAL fdatasync 耗时", "磁盘 I/O 瓶颈，考虑 SSD 或 NVMe"],
        ["replication_ack_us", "等待副本 ACK 耗时", "网络延迟、从节点负载过高"],
        ["total_us", "端到端总延迟", "所有上述阶段之和"],
    ]
)
bl()

h2("15.6  日志查看与分析")
for line in [
    "# Windows：实时查看最新日志",
    r"Get-Content C:\ProgramData\FalconDB\logs\falcon.log -Tail 100 -Wait",
    "",
    r"# 搜索 ERROR 级别日志",
    r'Select-String "ERROR" C:\ProgramData\FalconDB\logs\falcon.log',
    "",
    r"# 搜索慢查询记录",
    r'Select-String "slow_query" C:\ProgramData\FalconDB\logs\falcon.log | Select-Object -Last 20',
    "",
    "# Linux：查看最新错误",
    "tail -f /var/log/falcondb/falcon.log | grep -E 'ERROR|WARN'",
    "",
    "# JSON 格式日志解析（使用 jq）",
    "cat /var/log/falcondb/falcon.log | jq 'select(.level==\"ERROR\") | {ts: .timestamp, msg: .message}'",
]:
    cp(line)
pb()

# ══ 第十六章 AI 优化器 ══
h1("第十六章  AI 查询优化器与 AIOps 引擎")
h2("16.1  AI 查询优化器")
np("FalconDB 内置在线学习查询优化器（AI Optimizer），使用随机梯度下降（SGD）训练的线性代价模型，在规则优化器之上叠加智能计划选择层。无需外部服务，训练完全在进程内进行。")
bl()

h3("16.1.1  AI 优化器工作流程")
for line in [
    "SQL 查询到达",
    "  → SQL 解析 → 逻辑计划",
    "  → extract_features()         提取 15 维特征向量",
    "  → generate_candidates()      规则优化器生成候选计划（如 IndexScan vs SeqScan）",
    "  → AiOptimizer::select_plan() 代价模型对每个候选计划打分，选择预测代价最低的",
    "  → 执行器执行选中的计划",
    "  → record_feedback(actual_us) 将实际执行时间反馈给模型，在线 SGD 更新权重",
]:
    cp(line)
bl()

h3("16.1.2  15 维特征向量")
tbl(
    ["索引", "特征名称", "说明"],
    [
        ["0", "log₂(estimated_rows + 1)", "估计输出行数的对数（防止大数值主导）"],
        ["1", "join_count", "查询中 JOIN 操作的数量"],
        ["2", "filter_count", "WHERE 子句中谓词的数量"],
        ["3", "has_group_by (0/1)", "是否包含 GROUP BY 操作"],
        ["4", "has_order_by (0/1)", "是否包含 ORDER BY 操作"],
        ["5", "has_limit (0/1)", "是否包含 LIMIT 子句"],
        ["6", "index_available (0/1)", "主扫描表是否有可用索引"],
        ["7", "log₂(max_table_rows + 1)", "最大表的估计行数（对数）"],
        ["8", "log₂(second_table_rows + 1)", "第二大表的估计行数（对数）"],
        ["9", "selectivity (0..1)", "跨表选择率估计"],
        ["10", "projection_columns", "SELECT 投影列数量"],
        ["11", "has_aggregate (0/1)", "是否包含聚合函数"],
        ["12", "has_distinct (0/1)", "是否包含 DISTINCT"],
        ["13", "subquery_depth", "子查询嵌套深度"],
        ["14", "log₂(total_bytes_estimate + 1)", "总处理数据量估计（字节，对数）"],
    ]
)
bl()

h3("16.1.3  模型架构")
bp("算法：在线 SGD（随机梯度下降）+ L2 正则化，防止过拟合")
bp("预测目标：log₂(执行时间_μs)，对数变换使模型对延迟分布更鲁棒")
bp("计划类型偏置：7 维 one-hot 编码（SeqScan / IndexScan / IndexRangeScan / NestedLoopJoin / HashJoin / MergeSortJoin / Other）")
bp("模型总维度：23 权重 = 15 特征 + 7 计划类型 + 1 偏置项")
bp("预热阈值：50 个训练样本，未达到前始终使用规则优化器计划（零风险预热期）")
bl()

h2("16.2  AI 优化器监控")
for line in [
    "-- 查看 AI 优化器训练状态",
    "SHOW AI STATS;",
    "",
    "-- 输出字段说明：",
    "-- enabled            → AI 计划选择是否启用（始终为 true）",
    "-- samples_trained    → 已完成的 SGD 权重更新次数",
    "-- model_ready        → 是否达到预热阈值（samples_trained >= 50）",
    "-- ema_mae_log2       → 指数加权平均绝对误差（log₂ 代价单位，越低越准）",
    "-- query_fingerprints → 已追踪的不同查询形状数量",
]:
    cp(line)
bl()

h2("16.3  AIOps 四大子系统详解")
h3("16.3.1  慢查询检测器（SlowQueryDetector）")
np("使用自适应 EMA p95 基线（而非固定阈值），随系统负载自动调整：")
tbl(
    ["参数", "值", "说明"],
    [
        ["初始阈值下限", "100 ms", "冷启动期间的绝对最低阈值"],
        ["自适应阈值公式", "max(100ms, EMA_p95 × 2.0)", "负载轻时更严格，负载重时自动放宽"],
        ["预热期", "前 100 次查询", "使用绝对阈值下限"],
        ["滚动缓冲区大小", "200 条", "保留最近 200 条慢查询记录"],
        ["SQL 记录长度", "前 200 字符", "避免超长 SQL 占用过多内存"],
    ]
)
bl()

h3("16.3.2  异常检测器（AnomalyDetector）")
np("使用 Welford 在线方差算法的 3σ 规则，监控 TPS（每秒事务数）和延迟两个信号：")
tbl(
    ["触发条件", "告警级别", "说明"],
    [
        ["TPS 或延迟超过均值 + 3σ", "WARNING", "统计异常，可能是负载突增或性能退化"],
        ["延迟超过均值 5 倍", "CRITICAL", "严重性能问题，需立即介入"],
        ["连续 ≥ 5 次慢查询", "WARNING", "持续慢查询，可能是索引失效或数据热点"],
        ["冷启动保护（< 30 样本）", "不触发告警", "防止冷启动误报，需积累足够样本"],
    ]
)
bl()

h3("16.3.3  索引推荐器（IndexAdvisor）")
np("自动分析 SeqScan 执行计划的统计数据，生成可直接执行的索引建议：")
bp("触发条件：某表累计 ≥ 10 次全表扫描 AND 平均扫描耗时 ≥ 50ms")
bp("分析内容：识别 WHERE 子句中使用频率最高的 Top-3 过滤列")
bp("建议输出：生成具体的 CREATE INDEX 语句，包含表名和列名")
bp("排序方式：按影响分（全扫描次数 × 平均耗时）降序排列，优先显示影响最大的")
bl()

h3("16.3.4  工作负载分析器（WorkloadProfiler）")
np("按 SQL 指纹聚合执行统计，识别系统中的热点查询和低效 SQL：")
bp("SQL 规范化：移除字面值和参数，只保留查询结构（前 120 字符）")
bp("每个指纹追踪：调用次数、总/均/最大延迟、p95/p99（100 样本蓄水池抽样）、错误次数")
bp("容量：最多追踪 1000 个不同查询指纹，LRU 淘汰最久未使用的")
bl()

h2("16.4  AIOps SQL 查询接口")
for line in [
    "-- AIOps 引擎整体状态（12 项核心指标）",
    "SHOW AIOPS STATS;",
    "",
    "-- 异常告警列表（最近 50 条，最新在前）",
    "SHOW AIOPS ALERTS;",
    "",
    "-- 索引建议（按影响分降序排列）",
    "SHOW AIOPS INDEX ADVICE;",
    "",
    "-- Top 20 高耗时查询指纹（含 p95/p99）",
    "SHOW AIOPS WORKLOAD;",
    "",
    "-- 最近 100 条慢查询记录（含 SQL 文本）",
    "SHOW AIOPS SLOW QUERIES;",
]:
    cp(line)
bl()

h2("16.5  AIOps 指标字段说明")
tbl(
    ["SHOW AIOPS STATS 字段", "说明"],
    [
        ["slow_query_samples", "慢查询检测器处理的总查询样本数"],
        ["slow_query_threshold_us", "当前自适应慢查询阈值（微秒）"],
        ["slow_query_ema_baseline_us", "用于阈值计算的 EMA p95 基线（微秒）"],
        ["slow_query_logged", "滚动缓冲区中的慢查询记录数"],
        ["anomaly_tps_mean", "TPS 滚动均值（用于 3σ 基线）"],
        ["anomaly_tps_stddev", "TPS 滚动标准差"],
        ["anomaly_latency_mean_us", "延迟滚动均值（微秒）"],
        ["anomaly_latency_stddev_us", "延迟滚动标准差（微秒）"],
        ["anomaly_alerts_total", "累计触发的告警总次数"],
        ["index_advisor_tables_tracked", "有全扫描历史的表数量"],
        ["index_advisor_advice_count", "当前有效索引建议数量"],
        ["workload_fingerprints", "已追踪的不同查询指纹数量"],
    ]
)
pb()

# ══ 第十七章 性能调优 ══
h1("第十七章  性能调优指南")
h2("17.1  性能调优原则")
np("FalconDB 性能调优应遵循以下优先级顺序：")
bp("首先使用监控工具定位瓶颈（SHOW AIOPS WORKLOAD、EXPLAIN ANALYZE）")
bp("优先解决架构层面问题（索引缺失、数据模型不合理）")
bp("再考虑配置调整（内存、WAL、GC 参数）")
bp("最后才是硬件升级")
bl()

h2("17.2  Rowstore 内存引擎调优")
h3("17.2.1  内存配置")
for line in [
    "[storage]",
    "memory_limit_bytes = 4294967296    # 4 GB 内存预算",
    "",
    "[memory]",
    "# 配置内存软硬限制（示例：4GB 系统）",
    "# 软限制 = 80%：触发 GC 加速",
    "# 硬限制 = 90%：拒绝新写入",
]:
    cp(line)
bl()

h3("17.2.2  GC 调优")
for line in [
    "[gc]",
    "enabled = true",
    "interval_ms = 2000       # 更频繁 GC（默认 5000ms），适合高写入场景",
    "batch_size = 2000        # 每次扫描更多链（默认 1000）",
    "min_chain_length = 3     # 只 GC 版本链长 >= 3 的条目，减少不必要扫描",
]:
    cp(line)
bl()
np("GC 调优原则：")
bp("高写入场景：降低 interval_ms（更频繁 GC），防止内存累积")
bp("读密集场景：适当增大 min_chain_length，减少 GC 对读性能的干扰")
bp("监控 SHOW falcon.gc_stats 中的 gc_safepoint，确保安全点持续推进")
bp("发现 GC 安全点停滞 → 检查 SHOW falcon.txn 找出长事务，设置合理 statement_timeout_ms")
bl()

h2("17.3  WAL 性能调优")
h3("17.3.1  组提交配置（提升写入吞吐）")
for line in [
    "[wal]",
    "group_commit = true",
    "flush_interval_us = 2000    # 每 2ms 刷盘一次（增大间隔=更多批次聚合=更高吞吐）",
    "sync_mode = \"fdatasync\"    # fdatasync 比 fsync 快（不刷文件元数据）",
    "",
    "# 权衡：flush_interval_us 越大 → 吞吐越高，但单笔提交延迟可能增加",
    "# 对延迟敏感的 OLTP：flush_interval_us = 500~1000",
    "# 对吞吐优先的批量场景：flush_interval_us = 5000~10000",
]:
    cp(line)
bl()

h3("17.3.2  持久化策略选择")
tbl(
    ["场景", "推荐策略", "说明"],
    [
        ["最低延迟（可接受单点故障风险）", "local-fsync + group_commit", "微秒级提交延迟，单节点崩溃零丢失"],
        ["平衡（推荐大多数生产）", "local-fsync + sync-replica（quorum）", "毫秒级延迟，多数节点崩溃安全"],
        ["最高持久性（金融）", "sync-replica（all）", "所有副本确认，最高延迟但 RPO≈0"],
        ["开发测试（速度优先）", "no-wal（--no-wal）或 sync_mode=none", "零持久性，最快速度"],
    ]
)
bl()

h2("17.4  查询性能调优")
h3("17.4.1  索引优化策略")
np("以下是常见的索引优化场景和建议：")
tbl(
    ["问题场景", "诊断方法", "解决方案"],
    [
        ["高频全表扫描", "SHOW AIOPS INDEX ADVICE", "按建议创建索引"],
        ["ORDER BY 慢（有排序节点）", "EXPLAIN ANALYZE", "创建包含 ORDER BY 列的排序索引"],
        ["JOIN 慢（大表 NestedLoop）", "EXPLAIN ANALYZE", "在 JOIN 条件列上创建索引"],
        ["WHERE IN 子查询慢", "EXPLAIN ANALYZE", "考虑改写为 JOIN 或 EXISTS"],
        ["索引存在但不使用", "EXPLAIN 查看计划", "检查是否类型不匹配或选择性太低"],
    ]
)
bl()

h3("17.4.2  查询改写优化")
for line in [
    "-- 低效：对索引列使用函数（索引失效）",
    "SELECT * FROM orders WHERE YEAR(created_at) = 2026;",
    "",
    "-- 高效：改写为范围条件（利用索引）",
    "SELECT * FROM orders",
    "WHERE created_at >= '2026-01-01' AND created_at < '2027-01-01';",
    "",
    "-- 低效：前导通配符（索引失效）",
    "SELECT * FROM users WHERE email LIKE '%@gmail.com';",
    "",
    "-- 高效：全文搜索替代",
    "SELECT * FROM users WHERE email ~* '@gmail\\.com$';",
    "-- 或创建反转索引（应用层翻转字符串后插入）",
    "",
    "-- 低效：NOT IN 子查询（扫描大结果集）",
    "SELECT * FROM users WHERE id NOT IN (SELECT user_id FROM orders);",
    "",
    "-- 高效：LEFT JOIN IS NULL 替代",
    "SELECT u.* FROM users u",
    "LEFT JOIN orders o ON u.id = o.user_id",
    "WHERE o.user_id IS NULL;",
]:
    cp(line)
bl()

h2("17.5  连接池调优")
np("FalconDB 的每个连接都是独立的 OS 线程，连接过多会增加上下文切换开销。建议使用连接池控制连接数：")
tbl(
    ["参数", "建议值", "说明"],
    [
        ["连接池最大连接数", "CPU 核心数 × 2~4", "避免过多线程竞争"],
        ["空闲连接超时", "30-60 秒", "及时回收空闲连接"],
        ["连接检活", "每 30 秒 SELECT 1", "检测并替换失效连接"],
        ["最小连接数", "CPU 核心数", "保持热连接，避免冷启动延迟"],
    ]
)
for line in [
    "[server]",
    "max_connections = 1024          # 服务器端最大连接数",
    "idle_timeout_ms = 60000         # 空闲连接 60 秒后关闭",
    "statement_timeout_ms = 30000    # 单条 SQL 最长 30 秒",
]:
    cp(line)
bl()

h2("17.6  批量操作优化")
for line in [
    "-- 批量插入：单语句多值，比逐行插入快 10x+",
    "INSERT INTO events (ts, type, data) VALUES",
    "    (now(), 'click', '{\"page\": \"/home\"}'),",
    "    (now(), 'view',  '{\"page\": \"/product\"}'),",
    "    ...;    -- 建议每批 1000~5000 行",
    "",
    "-- COPY 导入：适合超大批量（百万行级别）",
    "COPY events FROM '/data/events.csv' WITH (FORMAT CSV);",
    "",
    "-- 批量更新：使用 UPDATE ... FROM 或 CASE WHEN",
    "UPDATE products SET price = v.price",
    "FROM (VALUES (1, 99.99), (2, 149.99), (3, 79.99)) AS v(id, price)",
    "WHERE products.id = v.id;",
    "",
    "-- 禁用事务自动提交（Python 示例）",
    "conn.autocommit = False",
    "for batch in batches:",
    "    cur.executemany(sql, batch)",
    "conn.commit()  # 每批一次提交，而非每行一次",
]:
    cp(line)
pb()

# ══ 第十八章 故障排除 ══
h1("第十八章  常见问题与故障排除")
h2("18.1  连接问题")
tbl(
    ["症状", "可能原因", "排查步骤", "解决方案"],
    [
        ["连接被拒绝（Connection refused）",
         "服务未启动，或端口号错误",
         "falcon service status；netstat -an | grep 5433",
         "falcon service start；检查 pg_listen_addr 配置"],
        ["认证失败（password authentication failed）",
         "密码错误、认证方式不匹配",
         "SHOW falcon.security；检查 falcon.toml 中的 auth 配置",
         "确认密码和认证方式（trust/md5/scram）；检查环境变量是否已设置"],
        ["连接超时（Connection timed out）",
         "防火墙规则阻断、服务过载",
         "ping 主机；telnet host 5433；查看 max_connections",
         "开放防火墙端口；增加 max_connections；优化连接池"],
        ["SSL 握手失败",
         "证书问题或 TLS 配置不匹配",
         "openssl s_client -connect host:5433",
         "检查证书有效期、路径配置；确认 min_version 兼容性"],
        ["IP 被锁定（auth rate limit）",
         "短时间内多次认证失败",
         "SHOW falcon.security_audit",
         "等待 300 秒自动解锁；或检查应用密码配置"],
    ]
)
bl()

h2("18.2  性能问题")
tbl(
    ["症状", "诊断命令", "常见原因", "解决方案"],
    [
        ["查询延迟突然升高",
         "SHOW AIOPS ALERTS; SHOW AIOPS SLOW QUERIES",
         "索引失效、大量 OCC 冲突、内存压力",
         "EXPLAIN ANALYZE 分析热点查询；检查索引；SHOW falcon.txn_stats 查看冲突率"],
        ["内存使用持续增长",
         "SHOW falcon.gc_stats; SHOW falcon.txn",
         "长事务阻塞 GC、GC 间隔过长",
         "设置 statement_timeout_ms；检查长事务；调低 gc.interval_ms"],
        ["写入吞吐低",
         "SHOW falcon.txn_stats",
         "WAL 刷盘慢、组提交未启用",
         "启用 group_commit；增大 flush_interval_us；升级存储到 SSD/NVMe"],
        ["OCC 冲突频繁（SQLSTATE 40001）",
         "SHOW falcon.txn_stats; EXPLAIN ANALYZE",
         "热点行争用（如计数器、状态更新）",
         "重新设计数据模型；使用批量更新；避免长事务持有版本"],
        ["全表扫描（SeqScan）",
         "EXPLAIN ANALYZE; SHOW AIOPS INDEX ADVICE",
         "缺少索引、统计信息陈旧",
         "按 AIOPS 建议创建索引；ANALYZE TABLE 更新统计信息"],
    ]
)
bl()

h2("18.3  复制问题")
tbl(
    ["症状", "诊断命令", "原因", "解决方案"],
    [
        ["从节点复制延迟持续增大",
         "SHOW falcon.replication_stats",
         "网络带宽不足、从节点负载过高",
         "检查网络带宽；降低从节点查询并发；优化 WAL 段大小"],
        ["从节点断开后无法重连",
         "检查服务器日志 falcon.log",
         "WAL 段被回收（从节点落后太远）",
         "从最新备份重建从节点；增加 WAL 保留时间"],
        ["脑裂告警（split-brain rejected）",
         "检查日志中的 EPOCH FENCE",
         "网络分区后双主场景",
         "确认 Epoch 隔离正常工作；旧主节点日志显示拒绝写入"],
        ["故障转移后数据不一致",
         "SHOW falcon.replication_stats",
         "复制策略配置为 Local，副本有延迟",
         "改用 sync-replica 策略；定期验证副本一致性"],
    ]
)
bl()

h2("18.4  WAL 和恢复问题")
tbl(
    ["症状", "排查步骤", "解决方案"],
    [
        ["启动时 WAL 恢复报错",
         "检查日志中的 WAL replay 错误；falcon doctor",
         "WAL 文件损坏则从备份恢复；磁盘空间不足则清理后重试"],
        ["数据目录权限拒绝",
         "检查 data_dir 的文件系统权限",
         "确保 FalconDB 服务账户有 data_dir 的读写权限"],
        ["WAL 段文件占用磁盘过大",
         "检查 data_dir 大小；SHOW falcon.checkpoint_stats",
         "执行 CHECKPOINT；确认 GC 正常运行；缩短 segment_size_bytes"],
        ["PITR 恢复失败",
         "检查 archive_dir 中的 WAL 文件连续性",
         "确保归档目录完整无缺失；检查 retention_hours 是否覆盖目标时间点"],
    ]
)
bl()

h2("18.5  诊断工具")
for line in [
    "# falcon doctor：一键诊断常见配置问题",
    r"falcon doctor --config C:\ProgramData\FalconDB\conf\falcon.toml",
    "",
    "# 检查：配置文件存在且可读；端口是否被占用；数据目录是否可写；服务注册状态",
    "",
    "# 配置版本验证",
    "falcon config check",
    "",
    "# 查看服务日志（Windows）",
    r"Get-Content C:\ProgramData\FalconDB\logs\falcon.log -Tail 200 | Select-String 'ERROR|WARN|panic'",
    "",
    "# 网络诊断",
    "Test-NetConnection -ComputerName 127.0.0.1 -Port 5433   # Windows",
    "nc -zv 127.0.0.1 5433                                   # Linux/macOS",
]:
    cp(line)
pb()

# ══ 第十九章 检查清单 ══
h1("第十九章  生产运维检查清单")
h2("19.1  上线前检查清单")
tbl(
    ["类别", "检查项目", "验证方法", "状态"],
    [
        ["安全", "认证方式为 scram-sha-256（非 trust）", "检查 falcon.toml [server.auth]", "□"],
        ["安全", "TLS 启用，版本 ≥ 1.3", "SHOW falcon.security", "□"],
        ["安全", "密码策略已启用（8+ 字符，复杂度）", "SHOW falcon.security_audit", "□"],
        ["安全", "认证速率限制已启用", "SHOW falcon.security_audit", "□"],
        ["安全", "SQL 防火墙已启用", "SHOW falcon.security_audit", "□"],
        ["安全", "审计日志已启用", "SHOW falcon.audit_log", "□"],
        ["安全", "应用使用最小权限账户（非超级用户）", "\\du 查看角色权限", "□"],
        ["安全", "密码通过环境变量传入（非硬编码）", "检查配置文件和代码", "□"],
        ["持久化", "WAL 已启用（wal_enabled = true）", "检查 falcon.toml [storage]", "□"],
        ["持久化", "data_dir 在独立高可用存储上", "检查目录路径和磁盘", "□"],
        ["持久化", "持久化策略已选择（local-fsync/sync-replica）", "检查 falcon.toml [wal]", "□"],
        ["持久化", "已验证 DCG 保证（写入后重启检验）", "运行 DCG 演示脚本", "□"],
        ["高可用", "至少一个从节点已部署并同步", "SHOW falcon.replication_stats", "□"],
        ["高可用", "故障转移已测试（演练）", "运行 e2e_failover 脚本", "□"],
        ["高可用", "客户端已配置重连和重试逻辑", "检查应用代码", "□"],
        ["备份", "全量备份已创建并验证可恢复", "运行备份验证", "□"],
        ["备份", "备份定时任务已设置", "检查定时任务配置", "□"],
        ["备份", "PITR 已启用（如需精确恢复）", "检查 falcon.toml [pitr]", "□"],
        ["监控", "Prometheus 指标已接入监控系统", "访问 /metrics 端点", "□"],
        ["监控", "告警规则已配置（延迟/连接数/内存）", "检查 Prometheus/Grafana 配置", "□"],
        ["监控", "日志已接入日志系统（ELK/Loki）", "检查日志配置", "□"],
        ["性能", "关键查询已创建索引（AIOPS 建议已处理）", "SHOW AIOPS INDEX ADVICE", "□"],
        ["性能", "连接池已正确配置", "检查应用连接池参数", "□"],
        ["性能", "内存限制已配置（Rowstore）", "检查 falcon.toml [storage]", "□"],
        ["性能", "GC 已启用并配置合理参数", "检查 falcon.toml [gc]", "□"],
        ["运维", "服务已注册为系统服务（Windows/systemd）", "falcon service status", "□"],
        ["运维", "日志轮转已配置", "检查 falcon.toml [logging]", "□"],
        ["运维", "磁盘空间告警已设置", "检查监控规则", "□"],
        ["运维", "文档已更新（连接信息、架构图）", "检查 wiki/文档", "□"],
    ]
)
bl()

h2("19.2  日常运维检查（推荐每日执行）")
for line in [
    "-- 1. 检查服务状态和活跃连接",
    "SHOW falcon.status;",
    "SHOW falcon.connections;",
    "",
    "-- 2. 检查是否有异常告警",
    "SHOW AIOPS ALERTS;",
    "",
    "-- 3. 检查慢查询",
    "SHOW AIOPS SLOW QUERIES;",
    "",
    "-- 4. 检查复制状态（主节点执行）",
    "SHOW falcon.replication_stats;",
    "",
    "-- 5. 检查 GC 状态（确保安全点在推进）",
    "SHOW falcon.gc_stats;",
    "SHOW falcon.gc_safepoint;",
    "",
    "-- 6. 检查事务统计（OCC 冲突率）",
    "SHOW falcon.txn_stats;",
    "",
    "-- 7. 检查审计日志中的异常事件",
    "SHOW falcon.audit_log;",
]:
    cp(line)
bl()

h2("19.3  周期性维护任务")
tbl(
    ["频率", "任务", "操作命令", "说明"],
    [
        ["每日", "验证备份可用性", "运行备份验证脚本", "确保备份文件未损坏且可恢复"],
        ["每日", "检查慢查询和索引建议", "SHOW AIOPS INDEX ADVICE", "及时处理新产生的慢查询"],
        ["每日", "检查磁盘空间（WAL 目录）", "系统监控或 df -h", "WAL 持续增长需扩容或清理"],
        ["每周", "审查安全审计日志", "SHOW falcon.audit_log", "检查异常登录和权限变更"],
        ["每周", "更新表统计信息", "ANALYZE TABLE table_name", "维护优化器统计数据准确性"],
        ["每月", "验证故障转移流程", "运行 e2e_failover 演练", "确保高可用机制正常工作"],
        ["每月", "审查和轮换密码", "ALTER ROLE user PASSWORD '...'", "遵循密码有效期策略"],
        ["每季度", "检查 TLS 证书有效期", "openssl x509 -in server.crt -noout -dates", "证书到期前 30 天更新"],
        ["每季度", "容量规划评估", "分析趋势数据", "预测未来 6 个月的存储和内存需求"],
    ]
)
bl()

h2("19.4  配置升级流程")
for line in [
    "# 步骤 1：检查配置版本兼容性",
    "falcon config check",
    "",
    "# 步骤 2：（如需要）迁移配置到新版本格式",
    "falcon config migrate",
    "",
    "# 步骤 3：备份当前配置和数据",
    r'Copy-Item "C:\ProgramData\FalconDB\conf\falcon.toml" "D:\backup\falcon.toml.bak"',
    "",
    "# 步骤 4：停止服务，安装新版本",
    "falcon service stop",
    "# 安装新版本 falcon.exe（覆盖 Program Files 中的二进制）",
    "",
    "# 步骤 5：验证新配置",
    "falcon doctor",
    "",
    "# 步骤 6：启动服务",
    "falcon service start",
    "",
    "# 步骤 7：验证服务正常运行",
    "falcon service status",
    "psql -h 127.0.0.1 -p 5433 -U falcon -c 'SELECT version();'",
]:
    cp(line)
bl()

h2("19.5  紧急故障响应手册")
tbl(
    ["故障等级", "触发条件", "立即操作", "后续处理"],
    [
        ["P0（服务完全中断）",
         "主节点宕机，从节点未自动提升",
         "1. 手动触发故障转移\n2. 检查 epoch 隔离日志\n3. 通知业务方切换连接",
         "事后分析故障原因；更新 HA 配置；添加监控告警"],
        ["P1（性能严重降级）",
         "P99 延迟 > 1s，OCC 冲突率 > 1000/min",
         "1. SHOW AIOPS ALERTS\n2. SHOW falcon.txn 找热点\n3. 临时降低并发",
         "索引优化；数据模型重设计；可能需要分片扩容"],
        ["P2（复制延迟告警）",
         "副本延迟 > 10000 LSN",
         "1. SHOW falcon.replication_stats\n2. 检查从节点负载\n3. 暂停从节点查询",
         "优化从节点资源；评估网络带宽；考虑升级硬件"],
        ["P3（磁盘空间告警）",
         "data_dir 使用率 > 80%",
         "1. 检查 WAL 段文件\n2. 执行 CHECKPOINT\n3. 清理旧备份",
         "扩容存储；评估 WAL 保留策略；容量规划"],
    ]
)
bl()

h2("19.6  有用的诊断 SQL 集合")
for line in [
    "-- 查看当前数据库版本和状态",
    "SELECT version();",
    "SHOW falcon.status;",
    "",
    "-- 查找占用内存最多的表",
    "SHOW falcon.table_stats;",
    "",
    "-- 查看最慢的查询指纹（按总 CPU 时间排序）",
    "SHOW AIOPS WORKLOAD;",
    "",
    "-- 查看有没有长时间运行的事务",
    "SHOW falcon.txn;",
    "",
    "-- 查看 GC 是否正常推进",
    "SHOW falcon.gc_stats;",
    "SHOW falcon.gc_safepoint;",
    "",
    "-- 查看连接分布（哪些 IP 连接最多）",
    "SHOW falcon.connections;",
    "",
    "-- 查看最近的安全事件",
    "SHOW falcon.audit_log;",
    "",
    "-- 查看 WAL 状态",
    "SHOW falcon.checkpoint_stats;",
    "",
    "-- 查看 AI 优化器状态",
    "SHOW AI STATS;",
]:
    cp(line)
pb()

# ══ 附录 ══
h1("附  录")
h2("A.  错误代码（SQLSTATE）参考")
tbl(
    ["SQLSTATE", "含义", "常见触发场景", "处理建议"],
    [
        ["40001", "serialization_failure（序列化失败）", "OCC 写-写冲突", "应用层重试事务（指数退避）"],
        ["23505", "unique_violation（唯一键冲突）", "INSERT 重复主键/唯一键", "使用 ON CONFLICT DO NOTHING/UPDATE"],
        ["23502", "not_null_violation（非空约束）", "INSERT/UPDATE NULL 到 NOT NULL 列", "检查应用逻辑，补充默认值"],
        ["23514", "check_violation（检查约束）", "值不满足 CHECK 条件", "检查业务逻辑和约束定义"],
        ["42P01", "undefined_table（表不存在）", "查询未建表", "确认表名拼写；检查所在 schema"],
        ["42703", "undefined_column（列不存在）", "查询不存在的列名", "检查列名拼写；表结构变更后更新 SQL"],
        ["0A000", "feature_not_supported", "使用了不支持的功能", "查看不支持功能列表；寻找替代方案"],
        ["08006", "connection_failure（连接失败）", "服务器宕机或网络中断", "重试连接；检查服务器状态"],
        ["57014", "query_canceled（查询取消）", "语句超时（statement_timeout）或手动取消", "优化查询性能；适当增大超时值"],
        ["53100", "disk_full（磁盘满）", "WAL 或数据目录磁盘空间不足", "立即清理磁盘；扩容存储"],
    ]
)
bl()

h2("B.  Crate 架构总览")
tbl(
    ["Crate 名称", "职责说明"],
    [
        ["falcon_storage", "多引擎存储（Rowstore/LSM/RocksDB/redb）、MVCC、二级索引、WAL、GC、USTM、TDE、CDC"],
        ["falcon_txn", "事务生命周期管理、OCC 验证、时间戳分配、死锁检测"],
        ["falcon_sql_frontend", "SQL 解析（sqlparser-rs）、绑定分析、类型检查"],
        ["falcon_planner", "逻辑/物理计划生成、CBO（代价优化）、AI 优化器（在线 SGD）、路由提示"],
        ["falcon_executor", "算子执行、表达式求值、全文搜索引擎、向量化列存 AGG"],
        ["falcon_protocol_pg", "PostgreSQL v3 线协议编解码、TCP 服务器"],
        ["falcon_cluster", "分片映射、WAL 复制、故障转移、scatter/gather、Epoch 隔离"],
        ["falcon_raft", "Raft 共识（单节点默认 / 多节点生产 raft_member）"],
        ["falcon_observability", "Prometheus 指标、结构化日志、pg_stat_statements"],
        ["falcon_enterprise", "控制平面 HA、TLS 轮换、备份恢复、多租户、资源隔离"],
        ["falcon_common", "共享类型、错误定义、配置、Schema、RBAC、数据类型（含 Decimal）"],
        ["falcon_server", "主二进制文件，组装所有组件，服务入口"],
        ["falcon_cli", "交互式 REPL CLI，集群管理命令，导入导出"],
        ["falcon_bench", "YCSB 风格基准测试工具"],
    ]
)
bl()

h2("C.  版本历史摘要")
tbl(
    ["版本", "主要新增功能"],
    [
        ["v0.1–v0.9", "OLTP 基础功能、WAL 持久化、主从复制（gRPC）、OCC 事务、SQL 解析器、安全认证、混沌测试加固"],
        ["v1.0", "LSM 磁盘引擎、企业功能（RLS/TDE/PITR/CDC）、分布式加固、完整 SQL 覆盖"],
        ["v1.0.1–v1.0.3", "零 Panic 加固、故障转移×事务矩阵测试、确定性与信任加固"],
        ["v1.1", "USTM 分层内存（Hot/Warm/Cold）、融合流式聚合、接近 PG 查询对等性"],
        ["v1.2（当前版本）", "Standard 版（RocksDB 默认）、PL/pgSQL 触发器、物化视图、ENUM 类型、AI 优化器（在线 SGD + SHOW AI STATS）、分片自动再均衡（47 项混沌测试）、AIOps 引擎"],
    ]
)
bl()

h2("D.  性能基准数据参考")
tbl(
    ["测试场景", "FalconDB", "PostgreSQL 16", "倍率"],
    [
        ["INSERT 吞吐（百万行批量插入）", "~340K 行/秒", "~185K 行/秒", "1.84x（FalconDB 更快）"],
        ["INSERT 阶段耗时（100 批 × 10K 行）", "~2.9 秒", "~5.4 秒", "0.54x（FalconDB 更快）"],
        ["单分片点查延迟（P99）", "< 1 ms", "~2-5 ms", "~2-5x（FalconDB 更快）"],
        ["跨分片查询（ORDER BY + COUNT）", "~2.8 秒", "~0.65 秒", "4.3x（PostgreSQL 更快）"],
        ["总计（DDL + INSERT + 查询）", "~6.4 秒", "~6.1 秒", "约相当"],
    ]
)
np("注：以上数据来自 FalconDB 内置基准测试（benchmarks/ 目录），硬件环境为标准云主机。实际性能取决于数据模型、查询复杂度和硬件配置。跨分片分析查询场景中 PostgreSQL 更快，FalconDB 优势在单分片 OLTP 场景。")
bl()

np("──────────────────────────────────────────────────────────────────────────")
np("文档结束  |  FalconDB v1.2  |  版权所有 © 2026 FalconDB Lab")
np("技术支持：通过 GitHub Issues 或官方文档网站获取帮助")
np("官方文档：https://github.com/falcondb-lab/falcondb")
np("──────────────────────────────────────────────────────────────────────────")

print("Part 3 完成，保存中...")
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
print("FalconDB 操作手册生成完毕！")
