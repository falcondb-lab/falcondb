# -*- coding: utf-8 -*-
"""Part 4a - 备份与监控（完整扩充）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn; from docx.oxml import OxmlElement

OUTPUT = r"c:\falcondb\falcondb-main\docs\FalconDB_操作手册.docx"
doc = Document(OUTPUT)

def h1(t): return doc.add_heading(t, 1)
def h2(t): return doc.add_heading(t, 2)
def h3(t): return doc.add_heading(t, 3)
def np(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.size = Pt(11)
    return p
def bp(t):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t); r.font.size = Pt(11); return p
def cp(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(t); r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0F0F0')
    pPr.append(shd); return p
def bl(): doc.add_paragraph()
def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        if c.paragraphs[0].runs:
            run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = c._tc; tcPr = tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        r = t.rows[ri+1]
        for ci,v in enumerate(row):
            r.cells[ci].text = str(v)
            if r.cells[ci].paragraphs[0].runs:
                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return t
def pb(): doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第十四章 备份与恢复（完整扩充版）
# ═══════════════════════════════════════════════════════════════
h1("第十四章  备份与恢复")

h2("14.1  备份策略设计原则")
np("数据库备份是生产运维的核心职责。良好的备份策略需满足三个目标：数据安全性（防止丢失）、快速恢复能力（最短 RTO）和成本效益（存储可控）。")
np("FalconDB 提供三种互补的备份机制，可按业务 RPO/RTO 要求组合使用。推荐生产组合：每日全量备份 + 持续 WAL 归档，实现任意时间点精确恢复。")
bl()
tbl(["备份类型","内容","RPO","RTO","典型频率"],
[["全量备份","所有数据文件+元数据+WAL历史","N/A（基线）","分钟至小时","每日或每周"],
 ["增量备份","上次备份后的 WAL / 变更日志","分钟级","全量时间+增量重放","每小时"],
 ["WAL 归档 (PITR)","持续实时归档所有 WAL 段","秒级","分钟级+WAL重放","实时连续"]])
bl()

h2("14.2  备份存储目标")
np("FalconDB 支持本地文件系统、Amazon S3、阿里云 OSS 三种存储目标。凭证均通过环境变量传入，不写入配置文件：")
for line in [
    "# 本地文件系统",
    "[backup]",
    'target = "local"',
    'path = "D:/backup/falcondb"',
    "retention_days = 30",
    "",
    "# Amazon S3",
    "[backup]",
    'target = "s3"',
    'bucket = "prod-backups"; prefix = "falcondb/"',
    'region = "ap-east-1"',
    "# 凭证: export AWS_ACCESS_KEY_ID=... AWS_SECRET_ACCESS_KEY=...",
    "",
    "# 阿里云 OSS",
    "[backup]",
    'target = "oss"',
    'bucket = "prod-backups"; prefix = "falcondb/"',
    'endpoint = "oss-cn-hangzhou.aliyuncs.com"',
    "# 凭证: export OSS_ACCESS_KEY_ID=... OSS_ACCESS_KEY_SECRET=...",
]: cp(line)
bl()

h2("14.3  Windows 手动备份操作")
h3("14.3.1  停服备份（强一致性保证，推荐）")
for line in [
    "# 1. 优雅停止服务",
    "falcon service stop",
    "",
    "# 2. 创建带时间戳的备份目录并复制",
    '$ts = Get-Date -Format "yyyyMMdd_HHmmss"',
    '$bdir = "D:\\backup\\falcon_$ts"',
    'New-Item -ItemType Directory $bdir | Out-Null',
    'Copy-Item -Recurse "C:\\ProgramData\\FalconDB\\data" "$bdir\\data"',
    'Copy-Item "C:\\ProgramData\\FalconDB\\conf\\falcon.toml" "$bdir\\"',
    "",
    "# 3. 记录备份元数据",
    '@{ timestamp=$ts; hostname=$env:COMPUTERNAME } | ConvertTo-Json > "$bdir\\meta.json"',
    "",
    "# 4. 重启服务并验证",
    "falcon service start",
    "psql -h 127.0.0.1 -p 5433 -U falcon -c 'SELECT version();'",
]: cp(line)
bl()

h3("14.3.2  热备份（CHECKPOINT 方式）")
for line in [
    "# 执行 CHECKPOINT 确保内存数据写入 WAL",
    "psql -h 127.0.0.1 -p 5433 -U falcon -c 'CHECKPOINT;'",
    "Start-Sleep -Seconds 2   # 等待 WAL 刷盘",
    '$ts = Get-Date -Format "yyyyMMdd_HHmmss"',
    'Copy-Item -Recurse "C:\\ProgramData\\FalconDB\\data" "D:\\backup\\falcon_hot_$ts"',
    "# 注意: 热备份期间的写入需要结合 WAL 归档才能恢复到一致状态",
]: cp(line)
bl()

h2("14.4  PITR 时间点恢复配置")
np("PITR 是应对误操作（意外 DELETE/DROP）最有效手段，允许恢复到任意历史时刻。")
h3("14.4.1  启用 WAL 归档")
for line in [
    "[pitr]",
    "enabled = true",
    'archive_dir = "D:/wal_archive"   # 建议独立磁盘',
    "retention_hours = 168             # 保留 7 天",
]: cp(line)
bl()

h3("14.4.2  执行 PITR 恢复到指定时间点")
np("场景：2026-03-12 14:30 执行了错误 DELETE，需要恢复到 14:25 的状态：")
for line in [
    "# 1. 停止服务",
    "falcon service stop",
    "",
    "# 2. 从最近全量备份恢复数据目录",
    'Remove-Item -Recurse "C:\\ProgramData\\FalconDB\\data" -Force',
    'Copy-Item -Recurse "D:\\backup\\falcon_20260312_020000\\data" "C:\\ProgramData\\FalconDB\\data"',
    "",
    "# 3. 在 falcon.toml 中添加恢复目标",
    "[recovery]",
    'target_time = "2026-03-12 14:25:00"',
    'archive_dir = "D:/wal_archive"',
    "",
    "# 4. 启动（自动重放 WAL 到目标时间后停止）",
    "falcon service start",
    "",
    "# 5. 确认恢复成功后，移除 [recovery] 段并正常重启",
    'Select-String "PITR" "C:\\ProgramData\\FalconDB\\logs\\falcon.log" | Select-Object -Last 20',
]: cp(line)
bl()

h2("14.5  恢复类型参考")
tbl(["恢复类型","说明","配置","使用场景"],
[["Latest（最新状态）","应用所有可用 WAL 到崩溃前最新状态","不设 target_time","主节点彻底损坏后全面恢复"],
 ["ToTimestamp（指定时间）","重放 WAL 到指定时间戳后停止","target_time = 'YYYY-MM-DD HH:MM:SS'","误操作后的精确 PITR"],
 ["ToLsn（指定 LSN）","重放 WAL 到指定日志序号","target_lsn = 12345678","已知 LSN 的精确恢复"],
 ["NewCluster（新集群）","恢复到全新独立实例","new_cluster = true","数据库迁移/克隆/灾备演练"]])
bl()

h2("14.6  备份验证最佳实践")
np("备份未经验证等同于没有备份。建议每月至少执行一次完整恢复演练：")
for line in [
    "# 在测试环境执行恢复验证",
    '$test = "C:\\falcondb_verify\\data"',
    'Copy-Item -Recurse "D:\\backup\\falcon_20260312\\data" $test',
    "",
    "# 使用不同端口启动验证实例",
    "falcon --data-dir C:\\falcondb_verify\\data --pg-addr 0.0.0.0:5444",
    "",
    "# 连接并验证关键数据完整性",
    "psql -h 127.0.0.1 -p 5444 -U falcon -c 'SELECT COUNT(*) FROM orders;'",
    "psql -h 127.0.0.1 -p 5444 -U falcon -c 'SELECT MAX(created_at) FROM orders;'",
    "psql -h 127.0.0.1 -p 5444 -U falcon -c 'SELECT COUNT(*) FROM users;'",
    "",
    "# 记录验证结果",
    '$result = @{ date=Get-Date; rows_orders=...; rows_users=... }',
    '$result | ConvertTo-Json >> "D:\\backup\\verification_log.json"',
]: cp(line)
bl()

h2("14.7  备份监控指标")
tbl(["监控项目","Prometheus 指标","告警条件"],
[["备份成功率","backups_completed / backups_started","< 100% 时立即告警"],
 ["最后备份时间","falcon_backup_last_success_timestamp","> 25 小时时告警（日备份）"],
 ["备份耗时","falcon_backup_last_duration_seconds","> 预期 2 倍时告警"],
 ["WAL 归档延迟","与 WAL 生成速率对比","落后 > 30 分钟时告警"],
 ["备份文件大小","与历史备份对比","偏差 > 20% 时告警（可能数据异常）"]])
for line in [
    "-- 查看备份审计记录",
    "SHOW falcon.audit_log;    -- 筛选 event_type = 'BACKUP_RESTORE'",
]: cp(line)
pb()

# ═══════════════════════════════════════════════════════════════
# 第十五章 监控与可观测性（完整扩充版）
# ═══════════════════════════════════════════════════════════════
h1("第十五章  监控与可观测性")

h2("15.1  可观测性三大支柱")
np("FalconDB 的可观测性体系围绕监控三大支柱（指标、日志、追踪）设计，三者结合形成完整的可观测性闭环：")
bp("指标（Metrics）：50+ SHOW 命令（SQL 接口）+ Prometheus 端点（/metrics）+ Grafana 仪表板，提供系统状态的数值快照")
bp("日志（Logs）：结构化 tracing 日志（text/JSON 格式）+ 慢查询日志 + 审计日志，提供事件的详细上下文")
bp("请求标识：每个请求携带 request_id（连接内单调递增）、session_id、query_id、txn_id、shard_id，贯穿所有输出")
bl()

h2("15.2  SHOW 命令完整参考")
tbl(["命令","说明","关键输出字段"],
[["SHOW falcon.status","服务器整体状态","version、role、uptime_secs、wal_lsn"],
 ["SHOW falcon.connections","当前连接详情","active_count、max_connections、per-session IP"],
 ["SHOW falcon.txn","活跃事务列表","txn_id、start_ts、type（Local/Global）、state"],
 ["SHOW falcon.txn_stats","事务统计汇总","committed、aborted、occ_conflicts、fast_path、slow_path"],
 ["SHOW falcon.txn_history","近期事务历史（最近 1000 条）","txn_id、duration_us、type、commit_lsn"],
 ["SHOW falcon.gc_stats","MVCC GC 统计","versions_reclaimed、bytes_reclaimed、safepoint_ts"],
 ["SHOW falcon.gc_safepoint","GC 安全点诊断","safepoint_ts、blocking_txn_id、blocking_txn_start_ts"],
 ["SHOW falcon.replication_stats","复制统计（主节点）","replica_id、applied_lsn、lag_lsn、lag_bytes、reconnects"],
 ["SHOW falcon.replica_stats","从节点状态（从节点执行）","primary_addr、applied_lsn、lag_ms、status"],
 ["SHOW falcon.slow_queries","慢查询日志（最近 200 条）","ts、duration_us、plan_type、sql（前 200 字符）"],
 ["SHOW falcon.table_stats","表统计信息","table_name、engine、estimated_rows、size_bytes"],
 ["SHOW falcon.checkpoint_stats","检查点统计","last_checkpoint_ts、wal_lag_bytes"],
 ["SHOW falcon.plan_cache","查询计划缓存","cache_entries、hit_rate、total_hits、total_misses"],
 ["SHOW falcon.security","安全态势","tls_enabled、tls_version、encryption_enabled、ip_allowlist"],
 ["SHOW falcon.security_audit","安全组件指标","auth_checks、auth_lockouts、firewall_checks、firewall_blocked"],
 ["SHOW falcon.audit_log","近期审计事件（最新 100 条）","event_type、ts_unix_ms、username、client_ip、detail"],
 ["SHOW falcon.scatter_stats","分布式查询统计","scatter_count、avg_scatter_us"],
 ["SHOW AI STATS","AI 优化器状态","enabled、samples_trained、model_ready、ema_mae_log2"],
 ["SHOW AIOPS STATS","AIOps 12 项汇总指标","slow_query_threshold_us、anomaly_alerts_total、index_advice_count"],
 ["SHOW AIOPS ALERTS","最近 50 条告警","alert_id、ts_unix_ms、severity、source、description"],
 ["SHOW AIOPS INDEX ADVICE","索引建议（按影响分降序）","table_name、full_scan_count、avg_scan_us、suggestion"],
 ["SHOW AIOPS WORKLOAD","Top-20 高耗时查询指纹","fingerprint、call_count、avg_us、p95_us、p99_us、error_count"],
 ["SHOW AIOPS SLOW QUERIES","最近 100 条慢查询","ts_unix_ms、duration_us、threshold_us、sql_text"],
 ["SELECT * FROM pg_stat_statements","语句级统计（PG 兼容视图）","query、calls、total_exec_time、mean_exec_time、rows"]])
bl()

h2("15.3  Prometheus 关键指标分组详表")
h3("15.3.1  事务指标")
tbl(["指标名称","类型","说明"],
[["falcon_txn_committed_total","Counter","总已提交事务数（含 Local + Global）"],
 ["falcon_txn_aborted_total","Counter","总回滚/中止事务数"],
 ["falcon_txn_fast_path_total","Counter","本地快速路径（OCC）提交数"],
 ["falcon_txn_slow_path_total","Counter","全局慢速路径（2PC）提交数"],
 ["falcon_txn_occ_conflicts_total","Counter","OCC 写-写冲突次数（需重试）"],
 ["falcon_txn_active","Gauge","当前进行中的事务数"],
 ["falcon_txn_commit_latency_us","Histogram","提交延迟分布（P50/P95/P99，单位微秒）"]])
bl()

h3("15.3.2  WAL 和存储指标")
tbl(["指标名称","类型","说明"],
[["falcon_wal_bytes_written_total","Counter","WAL 累计写入字节数"],
 ["falcon_wal_segments_total","Counter","WAL 段文件总创建数"],
 ["falcon_wal_flush_latency_us","Histogram","WAL fdatasync 延迟（关键指标，直接影响提交延迟）"],
 ["falcon_wal_backlog_bytes","Gauge","当前 WAL 未刷写积压字节数（> 0 说明写入压力大）"],
 ["falcon_wal_group_commit_batch_size","Histogram","每次组提交合并的事务数量（越大越高效）"],
 ["falcon_memory_used_bytes","Gauge","Rowstore 当前内存占用量"],
 ["falcon_memory_limit_bytes","Gauge","配置的内存预算上限（0 = 不限制）"],
 ["falcon_memory_pressure_total","Counter","触发内存背压事件的次数（应保持为 0）"],
 ["falcon_memory_evictions_total","Counter","GC / 淘汰事件次数"]])
bl()

h3("15.3.3  查询、连接和 GC 指标")
tbl(["指标名称","类型","说明"],
[["falcon_query_total","Counter","执行的查询总数"],
 ["falcon_query_latency_us","Histogram","端到端查询延迟（P50/P95/P99/P99.9）"],
 ["falcon_slow_query_total","Counter","超过慢查询阈值的查询数"],
 ["falcon_query_errors_total","Counter","查询错误数（user/retryable/transient/bug 分组）"],
 ["falcon_connections_active","Gauge","当前活跃连接数"],
 ["falcon_connections_rejected","Counter","因超过 max_connections 拒绝的连接数"],
 ["falcon_gc_versions_reclaimed_total","Counter","GC 回收的 MVCC 版本总数"],
 ["falcon_gc_bytes_reclaimed_total","Counter","GC 回收的总字节数"],
 ["falcon_gc_safepoint_ts","Gauge","当前 GC 安全点时间戳（应持续增长）"],
 ["falcon_replication_lag_lsn","Gauge（by replica_id）","副本与主节点的 LSN 差距（0 = 完全同步）"],
 ["falcon_replication_reconnects_total","Counter","副本重连次数（频繁说明网络不稳定）"],
 ["falcon_epoch_fence_rejections_total","Counter","Epoch 隔离拒绝的过期写入次数（非 0 说明发生过脑裂）"],
 ["falcon_crash_domain_panic_count","Gauge","捕获的 Panic 总数（非 0 需立即处理）"],
 ["falcon_admission_rejected_total","Counter","准入控制拒绝的请求数"],
 ["falcon_scheduler_high_active","Gauge","High 优先级（OLTP）当前活跃查询数"],
 ["falcon_scheduler_normal_wait_us","Gauge","Normal 队列累计等待微秒数"],
 ["falcon_scheduler_low_wait_us","Gauge","Low 队列（OLAP）累计等待微秒数"]])
bl()

h2("15.4  关键告警阈值建议")
tbl(["监控项","警告阈值","严重阈值","排查方向"],
[["P99 提交延迟","10 ms","50 ms","WAL 刷盘慢/复制延迟/同步副本模式"],
 ["副本复制延迟（LAG_LSN）","1,000","10,000","网络带宽/从节点负载/WAL 段大小"],
 ["OCC 冲突率（增速）","100/min","1,000/min","热点行争用/事务过长"],
 ["活跃连接数","80%×max","95%×max","连接池配置/慢查询阻塞连接"],
 ["内存使用率（Rowstore）","70%","90%","数据量增长/长事务阻塞 GC"],
 ["WAL 刷盘延迟 P99","5 ms","20 ms","磁盘 I/O 瓶颈/WAL 目录分区"],
 ["GC 安全点停滞","5 分钟","30 分钟","长事务/设置 statement_timeout"],
 ["Panic 计数","1","5","立即检查错误日志，联系技术支持"]])
bl()

h2("15.5  Prometheus Alerting Rules 示例")
for line in [
    "# alerts/falcondb.yml",
    "groups:",
    "  - name: falcondb_critical",
    "    rules:",
    "      - alert: HighCommitLatency",
    "        expr: histogram_quantile(0.99, rate(falcon_txn_commit_latency_us_bucket[5m])) > 50000",
    "        for: 2m",
    "        labels: { severity: critical }",
    "        annotations:",
    "          summary: 'P99 提交延迟 {{ $value }}μs 超过 50ms'",
    "",
    "      - alert: HighReplicationLag",
    "        expr: falcon_replication_lag_lsn > 10000",
    "        for: 5m",
    "        labels: { severity: warning }",
    "        annotations:",
    "          summary: '副本 {{ $labels.replica_id }} 延迟超过 10000 LSN'",
    "",
    "      - alert: MemoryPressure",
    "        expr: falcon_memory_used_bytes / falcon_memory_limit_bytes > 0.85",
    "        for: 5m",
    "        labels: { severity: warning }",
    "        annotations:",
    "          summary: 'Rowstore 内存使用率超过 85%，当前 {{ $value | humanizePercentage }}'",
    "",
    "      - alert: PanicDetected",
    "        expr: increase(falcon_crash_domain_panic_count[5m]) > 0",
    "        for: 0m",
    "        labels: { severity: critical }",
    "        annotations:",
    "          summary: '检测到 Panic，立即检查错误日志'",
    "",
    "      - alert: GCSafePointStalled",
    "        expr: changes(falcon_gc_safepoint_ts[30m]) == 0",
    "        for: 30m",
    "        labels: { severity: warning }",
    "        annotations:",
    "          summary: 'GC 安全点停滞超过 30 分钟，检查长事务'",
    "",
    "      - alert: OCCConflictSurge",
    "        expr: rate(falcon_txn_occ_conflicts_total[5m]) > 100",
    "        for: 3m",
    "        labels: { severity: warning }",
    "        annotations:",
    "          summary: 'OCC 冲突率 {{ $value }}/s，检查热点行竞争'",
]: cp(line)
bl()

h2("15.6  结构化日志解析（JSON 格式）")
np("生产环境建议使用 JSON 格式日志对接 ELK/Grafana Loki/CloudWatch，每条日志的关键字段：")
for line in [
    '{',
    '  "timestamp": "2026-03-12T14:23:45.123456Z",   // ISO 8601 UTC',
    '  "level": "info",                               // trace/debug/info/warn/error',
    '  "target": "falcon_storage::engine",            // Rust 模块路径',
    '  "message": "transaction committed",',
    '  "request_id": 12345,                           // 连接内单调递增',
    '  "session_id": 67890,                           // 会话 ID',
    '  "txn_id": 99,                                  // 事务 ID（0=autocommit）',
    '  "parse_us": 12, "plan_us": 8,                  // 各阶段延迟（微秒）',
    '  "execute_us": 45, "commit_us": 123,',
    '  "total_us": 188,                               // 端到端总耗时',
    '  "plan_type": "IndexScan",                      // 执行计划类型',
    '  "rows_affected": 1                             // 影响行数',
    '}',
]: cp(line)
bl()

h2("15.7  pg_stat_statements 兼容视图")
for line in [
    "-- Top-10 累计最耗时的查询",
    "SELECT",
    "    left(query, 100) AS sql,",
    "    calls,",
    "    round(total_exec_time::numeric, 2) AS total_ms,",
    "    round(mean_exec_time::numeric, 2) AS avg_ms,",
    "    round(stddev_exec_time::numeric, 2) AS stddev_ms,",
    "    rows",
    "FROM pg_stat_statements",
    "ORDER BY total_exec_time DESC",
    "LIMIT 10;",
    "",
    "-- 高频调用的查询（可能缺少缓存）",
    "SELECT left(query, 100) AS sql, calls, round(mean_exec_time::numeric,2) AS avg_ms",
    "FROM pg_stat_statements ORDER BY calls DESC LIMIT 10;",
    "",
    "-- 错误率最高的查询",
    "SELECT left(query, 100) AS sql, calls,",
    "    rows, shared_blks_hit, temp_blks_read",
    "FROM pg_stat_statements",
    "WHERE calls > 100",
    "ORDER BY (rows / NULLIF(calls, 0)) ASC",
    "LIMIT 10;",
    "",
    "-- 重置统计（需超级用户权限）",
    "SELECT pg_stat_statements_reset();",
]: cp(line)
pb()

print("Part 4a 完成，保存中...")
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
